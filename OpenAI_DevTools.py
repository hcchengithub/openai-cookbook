\ This file is .py good for Notepad++ to fold/unfold 

    fold all          alt-(shift)-0
    currnet level     alt-ctrl-(shift)-f
    Abolute level     alt-(shift)-1~8

\ 工作環境設定

    # These are my path envs. Magic 裡面可以引用 global variables by $varname or {varname} 
    jupyter   = %env jupyter   
    onedrive  = %env onedrive  
    github    = %env github    
    downloads = %env downloads 
    
    # 自動 %run forth.py 
    import glob, os
    for path_ in ["./", "../", "../../", "../../../", "../../../../", "../../../../../", "../../../../../../"]:
        pathname_ = glob.glob(
            os.path.join(path_, "forth.py"),
            recursive=True
            )
        if pathname_ : 
            # get_ipython().magic("run %s" % pathname_[0])
            get_ipython().run_line_magic("run", pathname_[0]) # python 3.8.1 之後建議新寫法
            break

    # 常用 files 

    %f s" {jupyter}\I study Pandas.ipynb" path\to/ . cr
    %f s" {github}\hubble2-nearest-neighbor\DataGrab_for_with_synonyms_and_chipset.ipynb" path\to/ . cr
    %f s" {github}\hubble2-nearest-neighbor\Hubble2NN_with_synonyms_and_chipset.ipynb" path\to/ . cr
    %f s" {github}\hubble2-nearest-neighbor\Hubble2NN_DevTools.py" path\to/ . cr
    %f s" {jupyter}\DevTools.py" path\to/ . cr
    %f s" {jupyter}\Matplotlib-學習筆記.ipynb" path\to/ . cr

    from icecream import ic
    ic.configureOutput(outputFunction=print) # 用正常顯示不要粉紅色的很礙眼
    ic('eep');
    
    from remote_pdb import RemotePdb
    RemotePdb('127.0.0.1',4444).set_trace()    
    
\ 如果本程式名列 left hand panel "Running Terminals and Kernels" 裡，可能已經有連結到主程式了。若從裡面關掉會把同一 kernel 的主程式也關掉。然而 Restart kernel 則不會破壞關聯。 

\ 查出 Console kernel 屬哪個程式的辦法: 
\   設定 global 變數 __file__, 兩邊一對照即知。

    __file__

\ 連結主程式：
\   1. right click the target tab > Create Console for Editor，這時候要選 kernel 連結。 
\   2. 從 kernels 裡重選主程式。 

\ datetime and pandas datetime 

    # 練習一下 type() 以及 is operator 的操作
    a = 1.2
    b = 3.4
    a is float --> False 這樣不行，要取 type() 才能 is  
    type(a) is float --> True 
        
    pd.Timestamp(2022,12,31,0,0,0) Pandas 有很豐富的日期時間功能 Timestamp, DateOffset, .. 等等
    day1 = pd.to_datetime("2022/11/3 12:34")
    day2 = pd.to_datetime("2022/11/9 2:51")
    pd.DateOffset(months=49) + day0
    pd.DateOffset(days=-49) + day0
    type(day1) is type(setup.now) --> False dateime 與 pandas 的日期時間不同 type 
    type(day1) == type(day2) --> True 
    type(day1) is pd.Timestamp --> True Bingo !!!
    type(day0) is datetime.datetime --> True Bingo !!!
    day1 is type(day2) --> True
    type(day0.date()) is datetime.date --> True Bingo !!!
    delta = day2 - day1 # 相減之後的 type 變成 pd.Timedelta 
    type(delta) is pd.Timedelta --> True Bingo !! where .days is an integer 
    delta.days --> 5 
    type(delta.days) --> int 

\ find a file under a folder --> copilot 正確推薦這個 function 真好用
    import os
    import fnmatch
    def find(pattern, path):
        result = []
        for root, dirs, files in os.walk(path):
            for name in files:
                if fnmatch.fnmatch(name, pattern):
                    result.append(os.path.join(root, name))
        return result
    # find('GRGX1-B140QAN05J+Functional+Spec+V01-0A_040821+(002).pdf_table.xlsx', 'c:\\Users\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\Vendor spec parsing\\Storage\\Panel spec extraction')
    # ['c:\\Users\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\Vendor spec parsing\\Storage\\Panel spec extraction\\lot2\\selected30\\GRGX1-B140QAN05J+Functional+Spec+V01-0A_040821+(002).pdf_table.xlsx']

    USERPROFILE=C:\Users\hcchen
    import glob,os
    glob.glob(
        os.path.join(os.environ['USERPROFILE']+"\\OneDrive", "**/forth.py"),
        recursive=True
        )

    glob.glob(
        os.path.join("./", "**/forth.py"),
        recursive=True
        )
    get_ipython().magic("run ../forth.py")

\ Azure CLI 
    os.system("az account set --subscription %s" % subscription_id[0])
    !az account set --subscription ed7216af-e52d-4c2d-9655-6fa3451e5ac0
    !az account set --subscription "azure subscription 1"
    !az account set --subscription R360WT

\ misc 
    from icecream import ic
    globals()
    subscription_id[0]
    resource_group
    name
    ws_name
    !chcp 950
    !cd
    get_ipython()
    dir(get_ipython())
    get_ipython().kernel
    dir(get_ipython().kernel)
    %f get_ipython :> () :> kernel.banner -->

    %f cr group_list            str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr storage_account_list  str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr keyvault_list         str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr vm_list               str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr webapp_list           str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr cosmosdb_list         str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    %f cr sql_server_list       str txt2json count --> dup
    group = peforth.pop()
    for i in group:
        %f i :> ['name'] -->
    np.floor(10.9)
    missing_rate = 0.75
    n_missing_samples = int(np.floor(data.shape[0] * missing_rate)) 
    missing_samples = np.hstack((np.zeros(data.shape[0] - n_missing_samples, dtype=np.bool), np.ones(n_missing_samples, dtype=np.bool)))
    %f missing_samples count --> nip
    rng = np.random.RandomState(0)
    rng.shuffle(missing_samples)
    missing_features = rng.randint(0, data.shape[1], n_missing_samples)
    data.iloc[np.where(missing_samples)[0], missing_features] = np.nan

    if not os.path.isdir('data'):
        os.mkdir('data')
    
\ Save the train data to a csv to be uploaded to the datastore
    pd.DataFrame(data).to_csv("data/train_data.csv", index=False)

    ds = ws.get_default_datastore()
    ds.upload(src_dir='./data', target_path='bankmarketing', overwrite=True, show_progress=True)

\ Upload the training data as a tabular dataset for access during training on remote compute
    
    train_data = Dataset.Tabular.from_delimited_files(path=ds.path('bankmarketing/train_data.csv')) # path 長這樣 $AZUREML_DATAREFERENCE_d8b680ae537443a688c6e4ca213bcf00
    label = "y"

    best_run
    %f best_run -->
    %f best_run dir -->
    %f best_run :> DELIM -->
    %f best_run -->
    %f best_run -->

\ 這個寫法沒有達到原目的，本來是想要對 missing_samples 的 rows 個別地把 missing_features 改成 NaN,
\ 可惜 .values() 不能用來做 assignment 而 .iloc[] .at[] 則不是個別，而是整 row 套用。
\ 所以只好改用 iloc 加上 for loop 了 

    data.iloc[np.where(missing_samples)[0], missing_features] = np.nan  # 改成 .iloc() 才對。

    np.where([False,  True,  True, True,  True, False, True])[0]
    %f missing_features count --> nip
    len(np.where(missing_samples)[0])

    data.iloc[[0,1,2], [7,7,7,7,9]] = ' ***** '  # 改成 .iloc() 才對。
    data.head()

    set(missing_features)


    data = pd.read_csv("https://automlsamplenotebookdata.blob.core.windows.net/automl-sample-notebook-data/bankmarketing_train.csv")
    rows = np.where(missing_samples)[0]
    for i in range(len(missing_features)):
        data.iloc[rows[i], missing_features[i]] = np.nan
    data

    %%time
    # Wait for the best model explanation run to complete
    from azureml.core.run import Run
    model_explainability_run_id = remote_run.id + "_" + "ModelExplain"
    print(model_explainability_run_id)
    model_explainability_run = Run(experiment=experiment, run_id=model_explainability_run_id)
    model_explainability_run.wait_for_completion()

    # Get the best run object
    best_run = remote_run.get_best_child()

    # Wall time: 27 s

    %%time
    import sys
    import json
    from azureml.automl.core.onnx_convert import OnnxConvertConstants
    from azureml.train.automl import constants

    from azureml.automl.runtime.onnx_convert import OnnxInferenceHelper

    def get_onnx_res(run):
        res_path = 'onnx_resource.json'
        run.download_file(name=constants.MODEL_RESOURCE_PATH_ONNX, output_file_path=res_path)
        with open(res_path) as f:
            result = json.load(f)
        return result

    if sys.version_info < OnnxConvertConstants.OnnxIncompatiblePythonVersion:
        test_df = test_dataset.to_pandas_dataframe()
        mdl_bytes = onnx_mdl.SerializeToString()
        onnx_result = get_onnx_res(best_run)

        onnxrt_helper = OnnxInferenceHelper(mdl_bytes, onnx_result)
        pred_onnx, pred_prob_onnx = onnxrt_helper.predict(test_df)

        print("pred_onnx",pred_onnx)
        print("pred_prob_onnx",pred_prob_onnx)
    else:
        print('Please use Python version 3.6 or 3.7 to run the inference helper.')
        
    # Wall time: 5.73 s

    %f pred_onnx -->
    %f pred_prob_onnx -->

    test_dataset
    validation_dataset
    test_df['y']
    test_df.drop(['y'], axis='columns')

\ Tibame class  
    from sklearn.metrics import confusion_matrix
    import numpy as np
    import itertools

    cf = confusion_matrix(test_df['y'].values, y_pred)
    plt.imshow(cf, cmap=plt.cm.Blues, interpolation="nearest")
    plt.colorbar()
    plt.title("Confusion Matrix")
    plt.xlabel("Predicted")
    plt.ylabel("Actual")
    class_labels = ["False", "True"]
    tick_marks = np.arange(len(class_labels))
    plt.xticks(tick_marks, class_labels)
    plt.yticks([-0.5, 0, 1, 1.5], ["", "False", "True", ""])
    # plotting text value inside cells
    thresh = cf.max() / 2.0
    for i, j in itertools.product(range(cf.shape[0]), range(cf.shape[1])):
        plt.text(
            j,
            i,
            format(cf[i, j], "d"),
            horizontalalignment="center",
            color="white" if cf[i, j] > thresh else "black",
        )
    plt.show()

    from sklearn.metrics import mean_squared_error, r2_score, auc, roc_auc_score, balanced_accuracy_score
    import sklearn
    %f sklearn :> metrics dir -->

    r"OneDrive\文件\Jupyter Notebooks\Tibame\All_AI_VI\L4\L4 課程範例檔\sklearn複習.ipynb".replace("\\","/")
    r"c:\Users\hcchen\OneDrive\文件\Jupyter Notebooks\MachineLearningNotebooks\tutorials\compute-instance-quickstarts\quickstart-azureml-automl\quickstart-azureml-automl.ipynb".replace("\\","/")      
    r"c:\Users\hcchen\OneDrive\文件\Jupyter Notebooks\MachineLearningNotebooks\tutorials\my-no-code-iris\script_run_notebook.ipynb".replace("\\","/") 
                                         
    pd.Categorical(test_df['y']).codes
    pd.Categorical(test_df['y_hat']).codes

    score = roc_auc_score(pd.Categorical(test_df['y_hat']).codes, pd.Categorical(test_df['y']).codes)
    score


    score = balanced_accuracy_score(pd.Categorical(test_df['y_hat']).codes, pd.Categorical(test_df['y']).codes)
    score

\ Azure 
    # Prepare a local subdirectory to have training, validation(option), and test(option) .csv files 
    if not os.path.isdir('data'):
        os.mkdir('data')
        
    # Save the train data to a csv to be uploaded to the datastore
    pd.DataFrame(training_data).to_csv("data/train_data.csv", index=False)
    pd.DataFrame(validation_data).to_csv("data/validation_data.csv", index=False)

    # Upload data to Azure cloud
    ds = ws.get_default_datastore()
    ds.upload(src_dir='./data', target_path='iris', overwrite=True, show_progress=True) # 將來改用 Dataset.File.upload_directory

    # Get train_data path 
    train_dataset = Dataset.Tabular.from_delimited_files(path=ds.path('iris/train_data.csv'))
    validation_dataset = Dataset.Tabular.from_delimited_files(path=ds.path('iris/validation_data.csv'))
    %f train_dataset -->
    %f validation_dataset -->

    # Wall time: 12.3 s


    import joblib

    # Load the fitted model from the script run.
    # Note that if training dependencies are not installed on the machine
    # this notebook is being run from, this step can fail.
    try:
        run.download_file("outputs/model.pkl", "model.pkl")
        model = joblib.load("model.pkl")
    except ImportError:
        print('Required dependencies are missing; please run pip install azureml-automl-runtime.')

    # --------------------------------------------------------
    hubble2 subject scoring datasets for Azure.ipynb
    # --------------------------------------------------------
    df
    %f df :> shape --> # 總共幾個
    df.shape[0] - df['tag'].sum() # 看看 tag==0 的有幾個 (品質不好的) 
    df.describe() # 基本統計數字沒啥用，features 都是 string 而 df.describe() 只看數字
    df['sn']
    df['subject'] # 中文亂碼嗎？ 先不管。這裡是 225 rows 的 training data 可能看不出來。

    # 確定多份 xlsx 檔都一樣，Yes. although file sizes are slightly different
    1 r"c:\Users\hcchen\OneDrive\文件\Jupyter Notebooks\Azure\hubble2 subject scoring\hubble2 subject qualiry train.xlsx" 
    2 r"c:\Users\hcchen\Downloads\hubble2 subject qualiry train.xlsx" 
      df = pd.read_excel(r"c:\Users\hcchen\Downloads\hubble2 subject qualiry train.xlsx")
    3 r"c:\Users\hcchen\Downloads\hubble2 subject qualiry train-compare.xlsx" 
      df = pd.read_excel(r"c:\Users\hcchen\Downloads\hubble2 subject qualiry train-compare.xlsx")

    # 既然都一樣，照用就好了，不用管 sn, Rand, text 這些了。挑好欄位 [tag, subject] 整個照順序送上去即可。含不含 text? 我評好壞時既然沒有看 text 就不要含。

    df_train_temp = df[['subject','tag']] # 引入 _temp 避免 SettingWithCopyWarning: A value is trying to be set on a copy of a slice from a DataFrame.
    df_train = df_train_temp.copy()
    df_train.shape


    df_train.info() # 查看一下 feature type 看對不對，不對！
        <class 'pandas.core.frame.DataFrame'>
        RangeIndex: 225 entries, 0 to 224
        Data columns (total 2 columns):
         #   Column   Non-Null Count  Dtype 
        ---  ------   --------------  ----- 
         0   subject  225 non-null    object <---- 就是 string 
         1   tag      225 non-null    int64 
        dtypes: int64(1), object(1)
        memory usage: 3.6+ KB

    df.dtypes # 查看所有 column 的 data type 
    df.astype({'a':float})  # 修改 column data type 

    df_train['len'] = df_train_temp['subject'].apply(len)
    df_train

    # 接下來把這個 dataset 送上 Azure cloud 

    %f ws :> experiments -->

    for experiment in ws.experiments:
        %f experiment --> # 只有一個，generator items 只知這樣取出

    %f run :> get_metrics() --> # 看起來是 best_model 的 metrics     

    # 想要自己比對 child_runs 的成績，
    i, children = 0, run.get_children()
    for child_run in children:
        i += 1
        %f i . space
        %f child_run :> display_name -->
        # 有 get_metrics() 

    # 列出各個 model 的 logloss 成績
    i, children = 0, run.get_children()
    for child_run in children:
        i += 1
        %f i . space
        %f child_run :> display_name -->
        %f child_run :> get_metrics() ==>


    # 結果發現 children 裡有很多別的東西，但這樣篩也不對 best_run 不再其中
    children = run.get_children()
    for child_run in children:
        if child_run.type == 'azureml.scriptrun':
            %f child_run -->
            break
    # 想要從 children 裡分離出 models 不成功
    children = run.get_children()
    for child_run in children:
        if child_run.type == 'azureml.scriptrun':
            %f child_run :> properties['iteration'] -->

    # 算了，就用 best_run 吧！
    best_run = run.get_best_child()


    import joblib
    try:
        run.download_file("outputs/model.pkl", "model.pkl")
        model = joblib.load("model.pkl")
    except ImportError:
        print('Required dependencies are missing; please run pip install azureml-automl-runtime.')
        raise

    ds = ws.get_default_datastore()
    Dataset.File.upload_directory(    
        src_dir="data",
        target=ds,
        overwrite=True,
        show_progress=True,
    )

    test_data_small = test_data.iloc[50000:50100]

    # Save the train data to a csv to be uploaded to the datastore
    pd.DataFrame(test_data_small).to_csv("data/test_data_small.csv", index=False, sep='|')

    ds.upload(src_dir='./data', target_path='data', overwrite=True, show_progress=True) # 將來改用 Dataset.File.upload_directory

    "this is a
    multiple line
    string" 

    df_train_temp['subject'] = df_train_temp['subject'].apply(preprocess_text)


    # cancel 之後短時間抓的 best_model 與等它跑更多出來之後的會 update 嗎？

    %f best_model dir -->
    %f best_model :> verbose -->
    %f best_model :> named_steps -->

    %f best_run dir -->
    %f best_run :> display_name -->

    json2file // ( json pathname -- ) Save json to text file 
    %f child_runs char child_runs.json json2file # 把 child_runs 保存成 .json 檔

    # 想去抓某個 model 的 confusion_matrix. Studio UI 上看起來是在該 child run 的 default storage root 
    child_run.download_file("confusion_matrix", "confusion_matrix.json") # 雲端，地端
    # "message": "File with path confusion_matrix was not found,
    # 可能是這個 child_run 不是 model 故沒有 confusion_matrix <--- 正確！

    %f child_run dir -->
    %f child_run :> log_confusion_matrix -->
    child_run.get_file_names()
        裡面就有 confusion_matrix
    child_run0.get_file_names() # 這個可以用來分辨，有 confusion_matrix 的才是 model 
    child_run0.get_environment() # 查這個 child_run 在 Azure cloud 上的 conda env. 慢。 
    %f child_run0 :> description --> # 沒用
    %f child_run0 :> get_secrets() --> # 沒用
    'confusion_matrix' in str(child_run0.get_file_names()) --> False 
    'confusion_matrix' in str(child_run4.get_file_names()) --> True 

    # 可以了，咱把他們掃過一輪，把所有 confusion matrix 都依序 download 下來加進 matrics dict 裡去。
    # 'confusion_matrix' in str(child_run.get_file_names())
    i, children = 0, run.get_children()
    for child_run in children:
        %f i -->
        if 'confusion_matrix' in str(child_run.get_file_names()): # 僅知之分辨這個 child_run 是否 model 的辦法
            child_run.download_file("confusion_matrix", "confusion_matrix.json") # 雲端，地端
            with open("confusion_matrix.json", "r") as f:
                child_runs[i]['confusion matrix'] = json.load(f)
        i += 1
    %f child_runs char child_runs.json json2file # 把 child_runs 保存成 .json 檔

    # 猜測一
    i = 7
    TN = child_runs[i]['confusion matrix']['data']['matrix'][0][0]
    FN = child_runs[i]['confusion matrix']['data']['matrix'][0][1] # <---- 可能相反
    FP = child_runs[i]['confusion matrix']['data']['matrix'][1][0] # <---- 可能相反
    TP = child_runs[i]['confusion matrix']['data']['matrix'][1][1]
    %f TN -->
    %f FN -->
    %f FP -->
    %f TP -->
    accuracy = (TP+TN)/(TN+FN+FP+TP)
    %f accuracy -->
    precision = (TP)/(FP+TP)
    %f precision -->
    recall = (TP)/(FN+TP)
    %f recall -->

    # 猜測二
    TN = child_runs[i]['confusion matrix']['data']['matrix'][0][0]
    FP = child_runs[i]['confusion matrix']['data']['matrix'][0][1] # <---- 可能相反
    FN = child_runs[i]['confusion matrix']['data']['matrix'][1][0] # <---- 可能相反
    TP = child_runs[i]['confusion matrix']['data']['matrix'][1][1]
    %f TN -->
    %f FN -->
    %f FP -->
    %f TP -->
    accuracy = (TP+TN)/(TN+FN+FP+TP)
    %f accuracy -->
    precision = (TP)/(FP+TP)
    %f precision -->
    recall = (TP)/(FN+TP)
    %f recall -->

    # 以上看不出哪個對，總之 TN 最高的為所求，看看哪個？
    for i in range(len(child_runs)):
        try:
            TN = child_runs[i]['confusion matrix']['data']['matrix'][0][0]
            %f i . space 
            %f TN -->
        except:
            pass

        
    # 取得這三個 model 的 Hyperlink 方便上 Studio UI 去細看    
    i, children = 0, run.get_children()
    for child_run in children:
        %f i -->
        if i in [4,10,23]:
            display(child_run)
        i += 1    

    # 把 model 的 display_name 加進去
    i, children = 0, run.get_children()
    for child_run in children:
        %f i -->
        if i in [4,10,23]:
            child_runs[i]['display_name'] = child_run.display_name
            print(child_runs[i]['display_name'])
        i += 1    
    %f child_runs char child_runs.json json2file # 把 child_runs 保存成 .json 檔

    %f child_runs :>  [4]['display_name'] --> # 眼看這個最好，抓最多 0 的出來。有待證實。
    %f child_runs :> [10]['display_name'] -->
    %f child_runs :> [23]['display_name'] -->
    child_runs :>  [4]['display_name'] --> bright_stomach_98911r6f (<class 'str'>)
    child_runs :> [10]['display_name'] --> crimson_muscle_d3mn553k (<class 'str'>)
    child_runs :> [23]['display_name'] --> tough_root_29q6vqgw (<class 'str'>)

    # 把這三個 model.pkl download 下來
    i, children = 0, run.get_children()
    for child_run in children:
        %f i -->
        if i in [4,10,23]:
            child_run.download_file("outputs/model.pkl", "model_%d.pkl" % i) # 雲端，地端
        i += 1    
        
    # 把 .json 讀回來
    %f char child_runs.json readTextFile py> eval(pop()) \ eval() is more reliable than txt2json
    child_runs = peforth.pop()


    # 上 Studio UI 看看第二好的這一批 
    i, children = 0, run.get_children()
    for child_run in children:
        if i in [6,7,8,9,12,24]:
            %f i -->
            display(child_run)
        i += 1    

    #7 #8 MaxAbsScaler, ExtremeRandomTrees 
    #9 SparseNormalizer, RandomForest
    #24 MaxAbsScaler, LightGBM 

    # 把這 4 個 model.pkl download 下來
    i, children = 0, run.get_children()
    for child_run in children:
        if i in [7,8,9,24]:
            %f i -->
            child_run.download_file("outputs/model.pkl", "model_%d.pkl" % i) # 雲端，地端
        i += 1    

    !ren model_7.pkl    model_7_ExtremeRandomTrees.pkl
    !ren model_8.pkl    model_8_ExtremeRandomTrees.pkl
    !ren model_9.pkl    model_9_RandomForest.pkl
    !ren model_24.pkl   model_24_LightGBM.pkl

    i = 0
    for child_run in child_runs:
        try:
            print(i, child_run['display_name'])
        except:
            pass
        i += 1    

    i, children = 0, run.get_children()
    for child_run in children:
        %f i -->
        child_runs[i]['id'] = child_run.id  # 把 child_run.id 加進去
        print(child_runs[i]['id'])
        i += 1    

    _dict = {
        'sn' : i,
        'display_name' : child_run.display_name,
        'run_object' : child_run,
        'model_fname' : fname,
    }

    child_run.download_file("outputs/model.pkl", fname)
    child_runs.append(_dict)

    child_runs[0]['run_object']    

    for child_run in child_runs:
        try:
            model = joblib.load(child_run['model_fname'])
            %f child_run :> ['sn'] -->
        except:
            pass

        
    data_sample = PandasParameterType(pd.DataFrame({"random0": pd.Series([0.0], dtype="float64"), "random1": pd.Series([0.0], dtype="float64"), "random2": pd.Series([0.0], dtype="float64"), "random4": pd.Series([0.0], dtype="float64")}))
    input_sample = StandardPythonParameterType({'data': data_sample})
    method_sample = StandardPythonParameterType("predict")
    sample_global_params = StandardPythonParameterType({"method": method_sample})

    result_sample = NumpyParameterType(np.array([False]))
    output_sample = StandardPythonParameterType({'Results':result_sample}) 

    # Save child_runs to a .json file 啊！不行，其中有 run object 
    %f child_runs char child_runs.json json2file # 把 child_runs 保存成 .json 檔
    # Restore child_runs
    %f char child_runs.json readTextFile py> eval(pop()) \ eval() is more reliable than txt2json
    child_runs = peforth.pop()

    y = test_data_orig['yy'].values

    %%time
    i = 0 
    for child_run in child_runs:
        %f i . space child_run :> ['display_name'] -->
        %f child_run :> ['run_object'].get_metrics()
        child_run['metrics'] = peforth.pop()
        # child_runs[0]['metrics']['accuracy']
        i += 1
    # Wall time: 32min 7s


    subscription_id
    resource_group
    my_name
    ws_name

    ws.compute_targets
    %f ws dir -->
    compute_name = '-'.join("{}-ci-{}".format(my_name,ws._workspace_id).split('-')[:3])
    compute_name
    %f instance -->

    %f ws :> compute_targets -->
    skip_creating_compute_instance = False
    for compute_target in ws.compute_targets.values():
        skip_creating_compute_instance = True
        %f compute_target :> name -->
        %f compute_target :> type -->
        %f compute_target :> id --> 


    %f ws :> compute_targets.keys() list :> [0] dup --> >x # compute name 
    %f ws :> compute_targets ==>
    %f x@ ws :> compute_targets[pop()] dir -->
    %f x@ ws :> compute_targets[pop()].type -->


    # 稍微防呆一下
    if subscription_id[:4] == "c27e":
        resource_group='%s_AutoML_resource_group' % my_name
    elif subscription_id[:4] == "ed72":      # R360WT 公司扣款
        resource_group='r360devteam'         # 公司帳號固定扣款群組
    else:
        assert False, "unknown resource-group! 你填對了嗎？" 

    x_train = np.array(range(9)).reshape(3,3);x_train
    x_train = np.expand_dims(x_train, -1) # 有意義的 arg2:[-3,-2,-1,0,1,2] 要嘛給每個 cell 變成 [cell]；要嘛給每個 row 變成 [[1,2,3]]
    x_train


    # Create workspace 
    ws = Workspace.create(name=ws_name,
        subscription_id=subscription_id[0],
        resource_group=resource_group,
        create_resource_group=True,
        tags = {"owner":my_name},
        location='eastus2'
    )
    # Save the workspace configuration file to local computer 
    # ws.write_config(path="./", file_name="ws_config.json")
    ws.write_config() # 不帶 path nor filename 讓 ws.write_config() 自己決定。
    print("---- Workspace.create ----")

    ws.compute_targets['hcchen-ci-d198fa61']
    list(ws.compute_targets.keys())[0]  <------ 取得 compute name 拿去填 config 
    list(ws.compute_targets.values())[0]
    dir(ws.compute_targets.keys())

    # 變出 一列 random numbers 
    np.random.rand(9)
    # 變出 一列 random numbers 轉置為 3x3 matrix 
    np.random.rand(9).reshape(3,3)


    y = np.random.rand(9)
    y_hat = y + np.random.rand(9)/10


    for i in range(len(child_runs)):
        try:
            # 可能有些 model 無法這樣 load 下來用
            model_from_file = joblib.load(child_runs[i]['model_fname'])
        except:
            print("joblib can not handle this model: %s" % child_runs[i]['model_fname'])
            continue
        y_hat = model_from_file.predict(test_data) # 可以直接接受 DataFrame
        accuracy = accuracy_score(y_test, y_hat)
        # con_matrix = confusion_matrix(y_test,y_hat)
        print(child_runs[i]['model_fname'], "\tActual/Metric accuracy: %0.02f/%0.02f \tDelta: %0.02f" % (accuracy, child_runs[i]['metrics']['accuracy'], accuracy-child_runs[i]['metrics']['accuracy']))

    TypeError: Cannot convert sklearn.neighbors._dist_metrics.ManhattanDistance to sklearn.metrics._dist_metrics.DistanceMetric

    問題：
    On Azure ws compute instance STANDARD_DS3_V2 
    ----> 2 run = experiment.submit(automl_config, show_output=True)
    ConfigException: ConfigException:
        Message: Input of type '<class 'pandas.core.frame.DataFrame'>' is not supported. Supported types: [azureml.data.tabular_dataset.TabularDataset]Please refer to documentation for converting to Supported types: https://docs.microsoft.com/en-us/python/api/azureml-core/azureml.core.dataset.dataset?view=azure-ml-py
    Local run 直接送 DataFrame 我記得可以，這個 VM 的 configuration 又不行了。一律轉成 Azure dataset 也好。


    問題：
    On Azure ws compute instance STANDARD_DS3_V2 
    joblib.__version__
    '0.14.1'
    ----> 2     model_from_file = joblib.load(child_runs[i]['model_fname'])
    TypeError: Cannot convert sklearn.neighbors._dist_metrics.ManhattanDistance to sklearn.metrics._dist_metrics.DistanceMetric

    我猜是 joblib version 的問題，於是進 terminla 
    pip install joblib==1.1.0 最新版

    結果變成
    ----> 2     model_from_file = joblib.load(child_runs[i]['model_fname'])

\ 如何處理指定 directory 以下的所有 excel fiels
\ ===========================================
    https://www.geeksforgeeks.org/how-to-iterate-over-files-in-directory-using-python/

    # so whare are the excel files ? It's certainly a list.
    datapath = r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing" 


    # 1. 這個方法只找出 files at the specified directory 
    # os.listdir()
    import os

    # iterate over files in
    # that directory
    for filename in os.listdir(datapath):
        f = os.path.join(datapath, filename)
        # checking if it is a file
        if os.path.isfile(f):
            print(f)


    # 2. 這個方法只找出 files at the specified directory 
    # os.scandir()
    import os

    # iterate over files in
    # that directory
    for filename in os.scandir(datapath):
        if filename.is_file():
            print(filename.path)


    # 3. 這個方法只找出 files at the specified directory 
    # Use pathlib library

    from pathlib import Path
    # iterate over files in
    # that directory
    files = Path(datapath).glob('*.*')
    for file in files:
        print(file)


    # 4. 這個方法就對了，the specified directory 以下的所有 files 
    # os.walk() 他是個 generator
    # 指定的 directory 以下每到了一層就傳回 dirs and files 以及這一層本身 root 
    # 所以每個 file 的 full path 就是 root + filename
    import os
    for root, dirs, files in os.walk(datapath):
        for filename in files:
            print(os.path.join(root, filename))


    # 5. 這個方法有看到 subdirectories 但是沒有展開進去。
    # use glob library 
    import glob

    # iterate over files in
    # that directory
    for filename in glob.iglob(f'{datapath}/*'):
        print(filename)

\ How to iterate excel worksheets
\ ================================
    https://www.thiscodeworks.com/python-loop-through-excel-sheets-place-into-one-df-stack-overflow-python/5ff5c97426c684001453b627

    %f be*
    import pandas as pd
    sheets_dict = pd.read_excel(doc_list[0]['pathname'], sheet_name=None)
    full_table = pd.DataFrame()
    for name, sheet in sheets_dict.items():
        peforth.bp(11,locals())
        sheet['sheet'] = name
        sheet = sheet.rename(columns=lambda x: x.split('\n')[-1])
        full_table = full_table.append(sheet)
    full_table.reset_index(inplace=True, drop=True)
    print(full_table)

    sheets_dict type --> # <class 'dict'>
    sheets_dict :> keys() --> # dict_keys(['0', 'p2_0', 'p3_0', 'p3_1'...])
    sheets_dict :> values() list :> [0] --> # DataFrame

    dataset = [dict] # worksheet {filename, worksheetname,text,label}

    dataset.iloc[0]['sheetname']
    %f sheet_df :> index list -->
    %f sheet_df :> columns list -->
    ' '.join(sheet_df.columns)
    ' '.join(list(sheet_df.index))

    sheet_df.astype({'rows':str}) # int, float, str or pd.StringDtype() 

    sheet_df['rows'] = sheet_df.index.astype(str)
    sheet_df.iloc[0]['rows'] # 檢查無誤
    sheet_df.index.dtype


    ' '.join(sheet_df.index.astype(str))
    %f doc_list count --> nip

    df = dataset.query('y==1')
    print("df.shape",df.shape)
    df

    {
        'filename':xl_list[i]['filename'],
        'sheetname':sheetname,
        'feature_columns':' '.join(sheet_df.columns),
        'feature_indexs':' '.join(sheet_df.index.astype(str)),
        'y':label_fun(sheetname),
    }
    dataset['y'].sum()
    [i for i in dataset['feature_indexs']]

    xl_list[i]['filename']
    df.T.head(1).T


    ' '.join(sheet_df[sheet_df.columns[0]].astype(str)),

\ Azure 
    ws = Workspace.from_config("config.json")
    print("---- Workspace.from_config ----")

    ws = Workspace.get(
        name=ws_name,
        subscription_id=subscription_id[0],
        resource_group=resource_group,
        )    
    print("---- Workspace.get ----")

    ws = Workspace.create(name=ws_name,
        subscription_id=subscription_id,
        resource_group=resource_group,
        create_resource_group=True,
        tags = {"owner":"hcchen"},
        location='eastus2'
    )
    ws.write_config("config.json")

    %f ws --> 

    # Wall time: 2min 32s   Create new workspace 
    # [ ] 奇怪，Azure CLI clone.ipynb 可以，本 cell 就不行？！ copy 整個 cell 過來試試看。。。。。。

    df = pd.read_csv(r"data\labeled5.csv")
    df


    best_run, fitted_model = run.get_output()


    "best_run = automl_run.get_best_child()"
    "best_dnn_run.download_file('outputs/model.pkl', model_dir + '/model.pkl')"
    "y_pred = best_model.predict(X_test_df)\n",

    y_pred = fitted_model.predict(test_df)

    test_dataset

    df.append() 用來給 dataframe 增添一 row 的，要 deprecated 了。照建議改用 pd.concat([df0, df1], ignore_index=True)
    https://www.adamsmith.haus/python/answers/how-to-insert-a-row-into-a-pandas-dataframe
    df0 = pd.DataFrame([[3, 4], [5, 6]])
    row = pd.DataFrame([pd.Series([11, 22])]) # 把直的 Series 轉成一 row 的 dataframe 
    df = pd.concat([df0, row], ignore_index=True) # 把 row 從 df0 下面加上去
    # 上面只是方便理解。 Series 也可以是橫的，這個才有用。 
    row = pd.DataFrame([pd.Series({'a':11,'b':22,'c':33})])

    s = ""
    for row in df0.iloc[:,[0]].itertuples(): # 只抓 leftmost column 整條串成一 string 
        s += str(row[1]) if row[1] else ''
    s

    # 整豎 leftmost column 改成 string 容忍常有 None 在其中
    ' '.join([str(row[1]) if row[1] else '' for row in df0.iloc[:,[0]].itertuples()])

    print(df)

    training_df

    datapath=r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing\labeled"
    xl_list = [] # excel file pathname list 
    for root, dirs, files in os.walk(datapath):
        for filename in files:
            pathname = os.path.join(root, filename)
            if pathname.endswith('xlsx'): 
                xl_list.append({"pathname":pathname,"path":root,"filename":filename})
    %f xl_list count --> nip
    dataset = pd.DataFrame({
        'filename': pd.Series(dtype='str'),
        'sheetname': pd.Series(dtype='str'),
        'feature_columns': pd.Series(dtype='str'),
        'feature_leftmost': pd.Series(dtype='str'),
        'feature_values': pd.Series(dtype='str'),
        'y': pd.Series(dtype='int')
    })

    ' '.join(sheet_df[sheet_df.columns[0]].astype(str)) if sheet_df.shape[0] else ''

    s = ""
    for r in range(sheet_df.values.shape[0]):
        for c in range(sheet_df.values.shape[1]):
            s += str(sheet_df.values[r][c])
    s    
    preprocess_text(s)

    r = requests.get('https://wistron.mobagel.com/')
    %f r :> text[:200] --> # 顯示為 text 
    %f requests :> codes.ok -->  # 這兩個比較即知成敗 
    %f r :> status_code -->      # 這兩個比較即知成敗 
    %f r :> url --> # 如果是複雜的 request 這個就是實際發出去的 url 

    r = requests.get("https://wistron.mobagel.com/v1/experiment/62625dc59d1db19b67f49fa8")
    response = requests.get('https://www.cmoney.tw/app/')


    r = requests.get("https://wistron.mobagel.com/v1/experiment/62625dc59d1db19b67f49fa8", headers = request_headers_dict)

    dfs[6]['RATING'][0][3] # '\uf0b0' 應該是 '\u00b0' or '°' 故 print('\uF0b0') 顯現為亂碼 。
    dfs[6]['RATING'][7][5] # '±' 這個就對了
    print('\uF0b0') # 
    ord('±') # 177 == 0xB1 
    "%X" % ord('') # 'F0B0' 
    print(chr(ord('\u00b0') + i*256) ) # '°'
    print(chr(ord('\u00b1') + i*256) ) # '±'


    逐張 df 檢查 df.columns 以及 df.values 都改成 string 然後掃描每個 char 的 ord('a') 凡是有超過 255 的就把整個 df 印出來。

    def search_pdf(serch_String, pdf_pathname):
        pdf_object = PyPDF2.PdfFileReader(pdf_pathname)
        # Get number of pages
        NumPages = pdf_object.getNumPages()
        # Enter code here
       
        # Extract text and do the search
        match_page = -1 # -1 代表 not found
        for i in range(0, NumPages):
            PageObj = pdf_object.getPage(i)
            Text = PageObj.extractText()
            if re.search(serch_String, Text): 
               match_page = i
        peforth.bp(11,locals())
        return match_page 

    search_pdf("soho",r"c:\Users\hcchen\Downloads\cv211_luc_crash_cart_adapter_ds_tc.pdf")
    re.search(peforth.execute('serch_String').pop(), peforth.execute('Text').pop())
    re.search("ten", peforth.execute('Text').pop(), flags=re.M|re.I)

    def Pdf_to_excel(pdf_pathname):
        table_filename = pdf_pathname+'_table.xlsx'
        assert not os.path.isfile(table_filename), "Abort! I don't overwrite your existing file : " + table_filename
        pdf = pdfplumber.open(pdf_pathname) 
        for page_num in range(len(pdf.pages)):   
            page = pdf.pages[page_num]       
            tables = page.extract_tables() # 自動讀取表格資訊，返回列表
            if len(tables) >= 0:
               for i in range(len(tables)):
                   table = tables[i]
                   for j in range(len(table)):
                       for k in range(len(table[j])):
                           table[j][k] = str(table[j][k]).replace('None','')
                           table[j][k] = str(table[j][k]).replace('\n','')           
                   if not os.path.isfile(table_filename): #新增xlsx檔
                          #將列表轉為df 
                          table_df = pd.DataFrame(table[1:],columns=table[0]) 
                          # 儲存excel
                          table_df.to_excel(table_filename, index=False, sheet_name= str(i))          
                   else:
                          with pd.ExcelWriter(table_filename, mode='a', engine='openpyxl') as writer:
                          #將列表轉為df
                               table_df = pd.DataFrame(table[1:],columns=table[0])
                          # 儲存excel
                               table_df.to_excel(writer, index=False, sheet_name= "p"+str(page_num+1)+"_"+str(i))
        peforth.bp(11,locals())
        pdf.close()
        
        return table_filename 

    Pdf_to_text(r"c:\Users\hcchen\Downloads\Intel-one-page.pdf")
    Pdf_to_text(r"c:\Users\hcchen\Downloads\cv211_luc_crash_cart_adapter_ds_tc.pdf")
    Pdf_to_text(r"c:\Users\hcchen\Downloads\(page 1202) 630094_ADL_P_PCH_EDS_Vol2_Rev1p3.pdf")

    Pdf_to_excel(r"c:\Users\hcchen\Downloads\(page 1202) 630094_ADL_P_PCH_EDS_Vol2_Rev1p3.pdf")
    Pdf_to_excel(r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing\General from teams\non-tagged\619503_ADL_EDS_Vol2b_Rev0p71.pdf")
    Pdf_to_excel(r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing\General from teams\non-tagged\630094_ADL_P_PCH_EDS_Vol2_Rev1p3.pdf")

    pdf_path = r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing\General from teams\non-tagged\630094_ADL_P_PCH_EDS_Vol2_Rev1p3.pdf"
    dfs = tabula.read_pdf(pdf_path, pages=1278) # stream=True 表示來源是 URL, pages="all" or 1 or [1,2,3]
    %f dfs -->

    for i in range(1,10):
        dfs = tabula.read_pdf(pdf_path, pages=i) 
        print(i)
        display(dfs)
        print()

    tables = tabula.read_pdf(r"c:\Users\hcchen\Downloads\(page 1202) 630094_ADL_P_PCH_EDS_Vol2_Rev1p3.pdf", pages=1)
    len(tables)

    %f mask_paths count nip -->
    %f train_gen dir -->
    train_gen.batch_size
    train_gen.aug
    train_gen.folder_path
    %f train_gen :> img_paths count nip --> # 1098 
    %f train_gen :> mask_paths count nip --> # 1098 
    train_gen.indexes --> array([924, 742, 432, ..., 238, 453, 485])
    %f train_gen :> indexes count nip --> # 1098
    %f train_gen :> indexes type --> # ndarray
    plt.hist(train_gen.indexes) --> train_gen.indexes 就是 0~1097 的 shuffled
    train_gen.indexes.min() --> 0
    train_gen.indexes.max() --> 1097
    train_gen.indexes.sort() 這下真的 inplace 了
    %f train_gen :> indexes[111] -->
    %f train_gen :> indexes[222] -->
    %f train_gen :> indexes[333] -->

\ mitosheet 

    import mitosheet
    mitosheet.sheet()

    # save DataFrame to .feather                              
    df.to_feather(path_name + ".feather")                             

    pd.value_counts(df_X.dtypes)
    pd.value_counts([1,1,1,2,3,3])
    pd.value_counts(pd.Series({"a":1,"b":1,"c":1,"d":2,"e":3,"f":3,}))
    pd.value_counts(df_X['Sex_of_Driver'])

    # Pandas Series 本具 visualizing 的能力                             
    unique_values.plot.bar(logy=True, figsize=(15, 4),
                           title="Unique values per feature");
    # 'area', 'bar', 'barh', 'box', 'density', 'hexbin', 'hist', 'kde', 'line', 'pie', 'scatter'
      

    # This cell creates %%csv cell magic. See http://ipython-books.github.io/14-creating-an-ipython-extension-with-custom-magic-commands
    from io import StringIO
    from IPython.core.magic import (register_cell_magic)
    @register_cell_magic
    def csv(line, cell):
        return pd.read_csv(
            StringIO(cell), # cell 是除了第一行之外的整個 cell 當作 input file
            sep=line,       # 第一行指定 seperator 也可以直接用 \s+ 把所有 white space 都當成 seperator
            engine='python',  # or 'C' 速度快但功能可能少點
            skip_blank_lines=True,  # default 就是 True
            dtype={'a':'str'},
            converters={
                'b':lambda x: np.where(int(x)>5, 'Big', 'Small') # 這時候 x 都當成 str 即使最後會自動轉成最佳的 type 
            })

                                 
    %%csv \s+
    a b c d
    1 2 3 a
    4 5 # b
    7 2 3 a                             
    8 5 # c
    9 2 3 a                             
    2 2 3 a                             

    df = _
                                 
    df.drop(labels=["a"], axis=1).duplicated()

    TESTDATA=""";index;col1;col2;
    ;1;4.4;99;
    a;2;4.5;200;
    b;3;4.7;65;
    ;4;3.2;140;
    """
    df = pd.read_csv(io.StringIO(TESTDATA)  , sep=";")
    df # 表中的 index column 是個普通 column 不是 index  (row 的 label)
    training_df
    %f help json                             

\ Vendor spec parsing 
    pkl_files_raw = r'''
        Documents\Jupyter Notebooks\Vendor spec parsing\model_#46_ivory_cartoon_n2ytg30w.pkl
        Documents\Jupyter Notebooks\Vendor spec parsing\model_#47_amusing_vase_ps9dvyqr.pkl
        Documents\Jupyter Notebooks\Vendor spec parsing\model_#48_placid_gold_q0gr9p7n.pkl
        Documents\Jupyter Notebooks\Vendor spec parsing\model_#49_polite_vase_vtkw2zy4.pkl
        ... snip ...
        '''

    pkl_files = [onedrive_path + i.strip() for i in pkl_files_raw.split('\n') if i.strip()]
    pkl_files
        
    %f child_runs count --> nip
                                 
    %f child_runs py> str(pop()) --> \ char c:\Users\hcchen\Downloads\1.json json2file
    %f child_runs py> str(pop()) txt2json -->
    %f child_runs :> [0].keys() -->

    models[i] = 
    {'display_name': 'affable_camel_27c6chlr',
     'model_fname': 'model_#25_affable_camel_27c6chlr.pkl',
     'metrics': . . . }
     
    models = []
    for child_run in child_runs:
        model = {k:v for k,v in child_run.items() if k not in ["sn","run_object"]}
        models.append(model)
    models[0].keys()

    for i in range(len(models)):
        print(i, models[i]['metrics']['f1_score_weighted'])
    models[22]  affable_camel_27c6chlr 'f1_score_weighted': 0.9875987143191639,  

    取得 model.pkl 的 path 
    modelpath = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing" + "\\"
    modelpath + models[22]['model_fname']

    選用這個 f1 score weighted 最高分的來跑跑看。。。
    'c:\\Users\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\Vendor spec parsing\\model_#25_affable_camel_27c6chlr.pkl'            
    onedrive_path = "c:\\Users\\hcchen\\OneDrive\\"
    onedrive_path + r"Documents\Jupyter Notebooks" 
    model_from_file = joblib.load('c:\\Users\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\Vendor spec parsing\\model_#25_affable_camel_27c6chlr.pkl')

    光 load 這個 model 都不行
    model_from_file = joblib.load('c:\\Users\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\Vendor spec parsing\\model_#25_affable_camel_27c6chlr.pkl')
    ImportError: cannot import name 'HoltWintersResultsWrapper' from 'statsmodels.tsa.holtwinters' (c:\Users\hcchen\AppData\Local\MyMLenv\lib\site-packages\statsmodels\tsa\holtwinters\__init__.py)

    整個都 load load 看。。。
    models_from_pkl = [] 
    for f in pkl_files:
        print(f, end=" ")
        try:
            m = joblib.load(f)
            print("ok")
        except Exception as e:
            m = e
            print("failed")
        models_from_pkl.append(m)
    結果發現：全部都不行！應該是前面啥 Azure 東西漏 import 了？
    No, the pkl that woked yesterday same failure too. So probably
    1. .yml venv or requirements version issue  <---- 應該是
    2. try LBB2 see see . . .  OK !!! 
    3. create venv through .yml on LPM2 --> [ ] 

    [ ] 可以用好幾個 model 聯合起來，得到勿枉勿縱的精確結果。

    # peforth.execute('results').pop()[0] # a model 
    # peforth.execute('results').pop()[0]['mst'][0]['list'] # a pdf's table list 
    # 還原 all results (all models run all pdf files get results)
    %f dropall s" d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\gray_music_3m4cx2tl\models all results_working.txt" readTextFile py> eval(pop())
    results = peforth.pop()
    for model in results:
        for pdf in model['mst']:
            pdf['list'] = eval(pdf['list']) # String 轉回成 list 
            # %f pdf -->
            
    # 一搜發現其實 pythhon pickle 就是用來 save object 的了，不用靠 json    
    # https://stackoverflow.com/questions/11218477/how-can-i-use-pickle-to-save-a-dict-or-any-other-python-object/22691344#22691344
    import pickle
    # save 
    with open(r"c:\Users\8304018\Downloads\all_results.pkl","wb") as f:
        pickle.dump(results, f) # 動用 protocal 使檔案稍小一點，但是相容性恐有問題，別用了。

    # restore Load data (deserialize)
    with open(r"c:\Users\8304018\Downloads\all_results.pkl", 'rb') as fi:
        all_results2 = pickle.load(fi)
        
    all_results2 == results # 靠！真的是 True 

    # 上 Azure 去直接 save .pkl 吧。。。 來不及了， compute-instance stop 就關機了。
        
    %f results count --> nip 
    %f results type -->
    len(results)


    new_child_runs = []
    for child_run in child_runs:
        new_child_runs.append({k:v for k,v in child_run.items() if k != 'run_object'})

    # How to save - restore python objects with .pkl file    
        import pickle
        with open(r"c:\Users\8304018\Downloads\child_runs.pkl","wb") as f:
            pickle.dump(new_child_runs, f)
            # save object to .pkl 

        # Load data (deserialize)
        with open(r"c:\Users\8304018\Downloads\child_runs.pkl", 'rb') as fi:
            child_run_read = pickle.load(fi)
            # restore object from .pkl 

        child_run_read == new_child_runs # 靠！真的是 True 

    # 舊的，分 array 不好，不如就一個大表
    all_pdfs = [] # 20 個 {"pdf name":..., "model":..., "metrics":..., "list":... }
    for idx in all_df.index:
        # print( all_df.loc[idx]['mst'])
        df = pd.DataFrame(columns=["model","metrics","list"])
        for mst in all_df.loc[idx]['mst']:
            row = {
                'model':all_df.loc[idx]['model'],
                'metrics':all_df.loc[idx]['metrics'],
                'list':mst['list'],
            }
            row = pd.DataFrame([pd.Series(row)])
            # dataset = dataset.append(row, ignore_index = True)  deprecate warning 改用 .concat
            df = pd.concat([df,row], ignore_index = True)
        all_pdfs.append(df)
    %f all_pdfs count nip -->
    all_pdfs[0]

    # 整張大表好
    # 20 個 {"pdf_name":..., "model":..., "metrics":..., "list":... }
    all_pdf_df = pd.DataFrame(columns=["pdf_name", "model", "metrics", "list"])
    for idx in all_df.index:
        for mst in all_df.loc[idx]['mst']:
            row = {
                'pdf_name':mst['pdf_file_pathname'],
                'model':all_df.loc[idx]['model'],
                'metrics':all_df.loc[idx]['metrics'],
                'list':mst['list'],
            }
            row = pd.DataFrame([pd.Series(row)])
            # dataset = dataset.append(row, ignore_index = True)  deprecate warning 改用 .concat
            all_pdf_df = pd.concat([all_pdf_df,row], ignore_index = True)
    %f all_pdf_df count nip -->
    all_pdf_df

    "c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\gray_music_3m4cx2tl\all_df.pkl" 

    # 按 metrics 展開成個別 columns
    metric_names = [k for k in all_pdf_df.iloc[0].metrics]
    for m in metric_names:
        col = [row.metrics[m] for row in all_pdf_df.itertuples()]
        all_pdf_df[m] = col
        
        
    %f col count --> nip
    %f col -->
    df.to_excel(r"c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\gray_music_3m4cx2tl\vcc.xlsx" )


    batch_x[batch_i].shape
    %f outputs count --> nip
    chr(20013) chr(ord('\u6587')) \u99ac\u8036\u901a(1)
    ord('中') == ord('\u4e2d')
    %f words ascii
    %f words char
    training_table_id

    Mobegal 把中文 (udf-8) "中文馬耶通" 記錄成 "\u4e2d\u6587\u99ac\u8036\u901a" 
    ord('中') == ord('\u4e2d') --> True
    chr(ord('\u6587')) --> '文'

    咱有自己的 mobegal user id 
        "_id": "62623daecd294e9b3681e137",
        "username": "H.C Chen"

    mobegal 的 'table' object 就像 Orange3 也有自己的 table object.
    %f response_json -->

    %f response type -->
    %f response dir -->
    response.json()
    %f response_json -->
    mlist = response_json['model_list']
    %f mlist count --> nip # Get List (Model) url = f"{host}/v1/experiment/{experiment_id}/model/getlist?projectId={project_id}"
    %f mlist :> [0].keys() -->
    %f mlist :> [0]['name'] -->


    %f experiment_response_json type -->
    %f experiment_response_json :> keys() -->
    experiments = experiment_response_json['experiments']
    %f experiments type -->
    %f experiments count --> nip 
    experiments[0]['name']
    experiments[1]['name']

    response_json['table']['status']

    experiment['status']
    experiment['experiment']["recommendations"][0]['model_id']
    experiment['experiment']["recommendations"][0]['name']

    experiment.keys()
    experiment['experiment'].keys()


    models
    %f models :> ['model_list'] count --> nip

    models[1]['name']
    models[2]['name']
    models[3]['name']
    models[4]['name']
    %f model_id_dict count --> nip

    df = pd.DataFrame(columns=['id','name','metrics'])
    for m in model_list:
        model_id = m['_id']
        row = {
            'id': model_id,
            'name': m['name'],
            'metrics': m['attributes']['auc']['cv_average']['th_max_by'],
        }
        row = pd.DataFrame([pd.Series(row)])
        df = pd.concat([df,row], ignore_index = True)

    # 把 DataFrame 某一 dict column 展開成 keys 為 title 的 multiple columns  
    # 確定所有 models 的 cv_average metrics 項目都一樣 --> Yes 
    metric_names = df.iloc[0]['metrics'].keys()    
    for i in range(1,13):
        print(df.iloc[i]['metrics'].keys() == metric_names)

    # 按 metrics 展開成個別 columns
    for m in metric_names:  
        col = [row.metrics[m] for row in df.itertuples()]
        df[m] = col
    #
    # 直接畫出 csv 的結構
    csv = """
        index;col1;col2
        1;4.4;99
        2;4.5;200
        3;4.7;65
        4;3.2;140
        """
    csv = '\n'.join([s.strip() for s in csv.split('\n') if s.strip()])
    print(csv)


    if __name__ == "__main__" :
        datapath = onedrive_path + r"Storage\Vendor Spec parsing\亂碼" 
        pdf_list = [] # pdf file pathname list 

        for root, dirs, files in os.walk(datapath):
            for filename in files:
                pathname = os.path.join(root, filename)
                if pathname.endswith('pdf'): 
                    pdf_list.append({"pathname":pathname,"path":root,"filename":filename})

        model_from_file = joblib.load(onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\gray_music_3m4cx2tl\outputs\model.pkl")
        for i in range(len(pdf_list)):
            pathname = pdf_list[i]['pathname'] # 一個 pdf 檔
            print("\n%s" % pathname)
            dfs = pdf_to_dfs(pathname)
            for page_num_df in dfs:
                print(".",end=" ")
                y_hat = model_from_file.predict(table2row(page_num_df[2])) 
                if y_hat[0]:
                    print(f"\npage:{page_num_df[0]+1}, table:{page_num_df[1]}")
                    display(page_num_df[2])
        print()

    import time
    time.localtime()
        time.struct_time(tm_year=2022, tm_mon=5, tm_mday=12, tm_hour=12, tm_min=50, tm_sec=37, tm_wday=3, tm_yday=132, tm_isdst=0)
    time.clock()
        DeprecationWarning: use time.perf_counter or time.process_time instead
        336.5210761
    time.ctime()
        'Thu May 12 12:53:59 2022'
    time.gmtime()
        time.struct_time(tm_year=2022, tm_mon=5, tm_mday=12, tm_hour=4, tm_min=54, tm_sec=30, tm_wday=3, tm_yday=132, tm_isdst=0)
    time.monotonic()
        776349.203
    time.monotonic_ns()
        776367203000000
    time.monotonic()
        776384.39 以秒為單位好像 T0 比較近好像
    time.time()
        1652331396.675766
    time.time()
        1652331408.5268917 以秒為單位 T0 比較久以前好像
    time.time()
        1652331419.295778

    # 以後 text preprocess 都用 %run 的，因為是 common code 集中管理
        try:
            del preprocess_text # 先殺掉，再驗證
        except:
            pass
        %run "../../Vendor spec parsing/Stage1/text-preprocess.py"
        preprocess_text # 確定有這個 function 產生

    #

    Stage2 
    [ ] 訓練 Mobegal 分辨 df 是有東西的還是沒東西的。以整個表的 shape[0], shape[1], df.columns 字數， df.leftmost 字數， 整表字數, 當作 X. 
    # 改寫成讀出所有的 table > 經 Mobegal 過濾 > 檢查 excel > df 是否 df.columns 可靠？會不會有例外？方便 Angus 加上 'Mark' column. 
    def path2df(datapath=r"c:\Users\hcchen\OneDrive\Storage\Vendor Spec parsing\labeled"):
        
        # 把整個 path 之下的所有 excel files 轉成一個特殊安排的 DataFrame.
        # 這些 excel files 都是由 EE 零件手冊 pdf 轉成 excel 而來。手工在 worksheet name 上後綴 _y 表示 labeled as True.
        # 轉出來的 DataFrame 有 [filename, sheetname, feature_columns, feature_leftmost, y] 這些 features.
        
        xl_list = [] # excel file pathname list 
        for root, dirs, files in os.walk(datapath):
            for filename in files:
                pathname = os.path.join(root, filename)
                if pathname.endswith('xlsx'): 
                    xl_list.append({"pathname":pathname,"path":root,"filename":filename})
        # an empty df 
        dataset = pd.DataFrame({
        'filename': pd.Series(dtype='str'),
        'sheetname': pd.Series(dtype='str'),
        'feature_columns': pd.Series(dtype='str'),
        'feature_leftmost': pd.Series(dtype='str'),
        'feature_values': pd.Series(dtype='str'),
        'y': pd.Series(dtype='int')
        })

        for i in range(len(xl_list)):
            pathname = xl_list[i]['pathname'] # 一個 excel 檔
            sheets_dict = pd.read_excel(pathname, sheet_name=None, engine="openpyxl") # 整個 excel 檔的 worksheets dict
            for sheetname, sheet_df in sheets_dict.items(): 
                s = ""
                for r in range(sheet_df.values.shape[0]):
                    for c in range(sheet_df.values.shape[1]):
                        s += str(sheet_df.values[r][c])
                _row = {
                    'filename':xl_list[i]['filename'],
                    'sheetname':sheetname,
                    'feature_columns':preprocess_text(' '.join(sheet_df.columns)),
                    'feature_leftmost':preprocess_text(' '.join(sheet_df[sheet_df.columns[0]].astype(str)) if sheet_df.shape[0] else ''),
                    'feature_values': preprocess_text(s),
                    'y':label_fun(sheetname)
                }
                row = pd.DataFrame([pd.Series(_row)])
                # dataset = dataset.append(row, ignore_index = True)  deprecate warning 改用 .concat
                dataset = pd.concat([dataset,row], ignore_index = True)
                # peforth.bp(11,locals())
                #if xl_list[i]['filename'] == "CardReader_RTS5242-GR_datasheet20140730-1_071.05242.0003.pdf_table_vcc.xlsx":
                #    peforth.bp(11,locals())
        return dataset

    path2df(datapath=onedrive_path + "Storage\\Panel spec extraction\\")
    row py: display(pop())
    sheet_df py: display(pop())

    # 有瑕疵的先用手工深入研究
    讀進 c:\Users\hcchen\OneDrive\Storage\Panel spec extraction\test\B156HAN02.9(HW_0A)+Pre-functional+spec+V0.1+-+2020-03-30+for+Dell_C9PFN.pdf
    檢查 df [1] p20_0_y  
    import pdfplumber # pip install pdfplumber
    pdf_pathname = r"c:\Users\hcchen\OneDrive\Storage\Panel spec extraction\test\B156ZAN03+8_HW_0A_+Pre+functional+spec+V0.1+(X00)+-2020-03-27+for+Dell_47R3H.pdf" # [6] p18_1_y 
    pdf = pdfplumber.open(pdf_pathname) 
    for page_num in range(len(pdf.pages)):   
        page = pdf.pages[page_num]       
        tables = page.extract_tables() # 自動讀取表格資訊，返回列表
        if page_num == 17 :  # p18-1
            break
    # 以上成功，得到 page (page_num + 1) 頁上的 tables 了
    tables type is list, count is table 個數 on the page. 
    %f tables count --> nip
    %f tables :> [0] type --> # list 其實是 matrix 
    df = pd.DataFrame(tables[1]) # 哈! 可以直接餵成 DataFrame 好看多了，將來也許用這招好做事有轉機
    df # 確定讀到了
    # 觀察獨到的 matrix 
    tables[1]


    df = row.sheet_df
    get_titles_str(df)
    get_leftmost_str(df)
    get_entire_table_str(df)

    %f row :> filename dir -->
    row.filename[:row.filename.find('.pdf')+4] # 這就是 pathname of the .pdf file 了
    # 但是 path2df 要的是 excel 所以直接可用
    df = None
    for i,row in enumerate(results.itertuples()):
        %f i . space row :> filename -->
        temp = excel2df(row.filename)
        if type(df) != "NoneType":
            %f i --> 
            df = pd.concat([df,temp], ignore_index = True)
        else:
            %f i --> # first one
            df = temp # the first one 
    df

    # Hubble2 RegEx.py 
    ands
    ors

    all_tokens = []
    for ors in ands:
        tokens += [t for t in ors]
        
    all_patterns


    pathname = r"c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Vendor Spec parsing\pdfs\normal\ALC3254_DataSheet_1.11-1_071.03254.M001.pdf_table.xlsx"
    wb = openpyxl.load_workbook(pathname) #path to the Excel file
    ws = wb.worksheets[0]
    cell = ws['b5']
    label_fun(cell)
    worksheet2df(ws)
    for ws in wb.worksheets:
        %f ws :> title -->
        break
        
    sheetname = wb.worksheets[0].title
    excel2df(pathname)    
    path2df(r"c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Vendor Spec parsing\pdfs\smaller\AngusLabel")
        
    dir(rows.filename)
    [value[113:] for value in rows.filename.values] # 113 try and error got 113 

    df.big_set[0].split(' ')

    import re
    set1 = set((re.split(r'\s+','we are splitting    the words')))
    set1
    set2 = set(re.split(r'\s+','aa bb cc'))
    set2
    set1 = set1.union(set2)
    set1

    z = re.match("^[a-z]+$","asdf s")
    z

    z = re.search("[a-z]+","asdfs33")
    z

    for row in df.itertuples():
        if re.match("^[a-z]+$",row.big_set):
            df.drop(row)
        
        
        
    # Pandas delete row 
    a = pd.DataFrame([[1,2],[3.4]])
    for row in a.itertuples():
        break
    a.drop(row.Index)

    # 先把 alphabet 數字加小數點 都用 space 與前後 split 開
    s = re.sub('([a-z])([^a-z])', '\\1 \\2', s)
    s = re.sub('([^a-z])([a-z])', '\\1 \\2', s)
    s = re.sub('([0-9.])([^0-9.])', '\\1 \\2', s)
    s = re.sub('([^0-9.])([0-9.])', '\\1 \\2', s)
    # 把 alphabet 數字加小數點 以外的都剔除 
    s = re.sub('[^a-z0-9.]', ' ', s)
    # 最後整理一下
    s = ' '.join(re.split('\\s+', s))
    %f s count swap --> . cr
    df3.filename = df2.filename
    df3['filename'] = df2.filename
    df3['sheetname'] = df2.sheetname
    print('\n'.join(list(df3.filename+" "+df3.sheetname)))
    print('\n'.join(list(df1.filename+" "+df1.sheetname)))

    row_test_data4['y_predicted']    = row_test_data4_predicted.y_predicted
    row_test_data4['y_1_confidence'] = row_test_data4_predicted['1_predicted_proba']
    row_test_data4['y_0_confidence'] = row_test_data4_predicted['0_predicted_proba']

    %f be*
    df = path2df(datapath=working_path)  # 得到的 df 是 training dataset 總表了

    <py>
    import random
    l = [1,2,3,4,5,6,7,8,9]
    print(l)
    random.shuffle(l)
    push(l)
    </py> -->

    <py>
    import random
    push(random.sample(range(3,30-3),5))
    </py> -->

    %f be*
    df = path2df(datapath=working_path)  # 得到的 df 是 training dataset 總表了

    pd sheetname get_titles_str get_leftmost_str get_entire_table_str pathname filename df sheet_df 
    <py>
        pd, sheetname, get_titles_str, get_leftmost_str, get_entire_table_str, pathname, filename, df, sheet_df = pop(8), pop(7), pop(6), pop(5), pop(4), pop(3), pop(2), pop(1), pop(0)
        # Augmentation by cutting the target table 5 times
        # 任取一數 15, 前 4 後 4 還剩 7 個，比 5 大夠多了。每個夠大的 table 切五刀，每刀都把它變成兩個 table，共 10 個新的。
        if sheet_df.shape[0] >= 15: 
            cuts = random.sample(range(3,sheet_df.shape[0]-3),5) # 切五刀
        print(cuts)
        print(sheet_df.iloc[0:cuts[0]].shape) # correct
        print(sheet_df.iloc[cuts[0]:].shape) # correct 
        for cut in cuts: # 切五刀變十個
            cut_dfs = [] # 切一刀變兩個
            cut_dfs.append(sheet_df.iloc[0:cut])
            cut_dfs.append(sheet_df.iloc[cut:].copy())
            # 上半段
            if True:
                cut_df = cut_dfs[0]
                _row = {
                    'filename': filename,
                    'sheetname':sheetname,
                    'feature_columns':  get_titles_str(cut_df),
                    'feature_leftmost': get_leftmost_str(cut_df),
                    'feature_values':   get_entire_table_str(cut_df),
                    'y':1,
                }
                row = pd.DataFrame([pd.Series(_row)])
                df = pd.concat([df,row], ignore_index = True)
            # 下半段
            if True:
                cut_df = cut_dfs[1]
                cut_df.columns = cut_df.iloc[0] # 模仿 2nd part of a cut'ed table 把 title 改掉
                cut_df.reset_index(inplace=True, drop=True)
                cut_df.drop(0, axis="rows", inplace=True) # 丟棄變成 title 的 row 
                _row = {
                    'filename': filename,
                    'sheetname':sheetname,
                    'feature_columns':  get_titles_str(cut_df),
                    'feature_leftmost': get_leftmost_str(cut_df),
                    'feature_values':   get_entire_table_str(cut_df),
                    'y':1,
                }
                row = pd.DataFrame([pd.Series(_row)])
                df = pd.concat([df,row], ignore_index = True)
            
        push(df)
    </py> constant extended_df // ( -- df ) bigger df 
    extended_df -->

    df.columns = df.iloc[0]
    df.drop(0, axis="rows", inplace=True)
    [f"{str(cell)}_{str(i)}" for i,cell in enumerate(df.iloc[0])]
        %f col type -->

    df.iloc[0] = 
    columns = pd.Series([f"{str(cell)}_{str(i)}" for i,cell in enumerate(df.iloc[0])])
    df.rename(columns=columns, inplace = True)
    df

    df.rename(columns=pd.Series("abcde"), inplace = True)
    df.columns
    dtale.show(df,open_browser=True)
    help(dtale.show)

    # Message box 
    # https://stackoverflow.com/questions/2963263/how-can-i-create-a-simple-message-box-in-python
    import ctypes  # An included library with Python install.   
    ctypes.windll.user32.MessageBoxW(0, "Continue?", "Pause", 0) # click ok to continue

    pd.read_csv(io.StringIO(csv_in_memory)  , sep="\\s")

    list(df1['code'])
    sum(df1['code'])

    dfa.query('y==1').filename.values v.s. dfb.Filename.values
    比對無誤
    df1.columns == dfa.columns

    df = pd.read_csv(onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\table_flaw_training_data_v2.csv")
    df.iloc[0].filename
    df.filename
    'B140QAN05J Functional Spec V01-0A_060321.pdf_table.xlsx' in list(df.filename)

    'B140QAN05J Functional Spec V01-0A_060321.pdf' in list(df.filename)
    'p20_0_flaw_y' in list(df.sheetname)

    __name__ == "__main__"
    datapath = onedrive_path + "Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction" 
    xl_list = [] # excel file pathname list 
    for root, dirs, files in os.walk(datapath):
        for filename in files:
            pathname = os.path.join(root, filename)
            if filename in list(df.filename):
                %f filename -->
                pathname = os.path.join(root, filename)
                xl_list.append({"pathname":pathname,"path":root,"filename":filename})
                
    df.shape
    pathname

    table_flaw_training_data_v2_df = pd.read_csv(onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\table_flaw_training_data_v2.csv")
    table_flaw_training_data_v2_df.query('filename == "%s"' % "B156ZAN03+8_HW_0A_+Pre+functional+spec+V0.1+(X00)+-2020-03-27+for+Dell_47R3H.pdf_table.xlsx").sheetname

        for i in range(len(xl_list)): # 逐一 excel 表
            pathname = xl_list[i]['pathname'] # 一個 excel 檔
            filename = xl_list[i]['filename']
            sheets_dict = pd.read_excel(pathname, sheet_name=None, engine="openpyxl") # 整個 excel 檔的 worksheets dict
            sheetnames = table_flaw_training_data_v2_df.query('filename == "%s"' % filename).sheetname # 這個 xlsx 檔裡的 target tables (sheetnames)
            for sheetname, sheet_df in sheets_dict.items(): # 這個 xlsx 檔裡的所有 worksheets 逐一
                if sheetname in list(sheetnames):
                    %f filename . space sheetname . cr
                    row = table_flaw_training_data_v2_df \
                        .query('filename == "%s"' % filename) \
                        .query('sheetname == "%s"' % sheetname) #  a row of the training data v2, OK 者要 aug 變多
                    raise;
                    
                    
    table_flaw_training_data_v2_df \
        .query('filename == "%s"' % filename) \
        .query('sheetname == "%s"' % 'p15_1').shape

    sheet_df.shape[0]
    import random
    random.sample(range(3,sheet_df.shape[0]-3),2)
    inst.target_table.values[0]
    import re

    xl_list

    for i in range(len(xl_list)):
        pathname = xl_list[i]['pathname'] # 一個 excel 檔
        sheets_dict = pd.read_excel(pathname, sheet_name=None, engine="openpyxl") # 整個 excel 檔的 worksheets dict
        for sheetname, sheet_df in sheets_dict.items(): 
            %f sheetname -->

    dfa = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\lot2_test_23files_target_tables.csv"  
    dfa = pd.read_csv(dfa, encoding='utf-8')

    dfa.iloc[0].filename

    xl_list

    import pandas as pd
    import openpyxl

    pathname = r"c:\Users\8304018\Downloads\【更新】參訓名單_v5.xlsx" 

    # 直接 access excel if you like to do it this way refer to openpyxl documents
    wb = openpyxl.load_workbook(pathname) #path to the Excel file


    # 先轉成 Pandas DataFrame 功能比較完備
    sheets_dict = pd.read_excel(pathname, sheet_name=None, engine="openpyxl") # 整個 excel 檔的 worksheets dict
    for sheetname, sheet_df in sheets_dict.items():
        for r in range(sheet_df.values.shape[0]):
            for c in range(sheet_df.values.shape[1]):
                print(sheet_df.values[r][c])

    # excel 讀進來問題很多，最好用視覺化 jupyterlab web 環境 e.g. mitosheet。以下舉例處裡掉 title 的問題。
    sheet_df.columns = list(sheet_df.iloc[0])
    sheet_df.drop(0, axis="rows", inplace=True)

    import nltk
    nltk.download('stopwords')
    from nltk.corpus import stopwords
    name = tokenize_column(airbnb["name"].ascharacter(), stopwords.words('english'))
    nltk.download()
    text airbnb["name"].ascharacter()
    %f name -->
    %f name :> shape -->
    %f airbnb -->

    train["is_train"] = 1  # 增加一欄 "is_train" 
    test["is_train"] = 0

    drift_data = train.rbind(test)
    drift_data["is_train"] = drift_data["is_train"].asfactor()

    x = list(set(train.col_names) - set(["price", "is_train", "id", "host_name"]))

    # https://googleapis.dev/python/aiplatform/latest/index.html
    aiplatform.init(
        # your Google Cloud Project ID or number
        # environment default used is not set
        project="782786629111",

        # the Vertex AI region you will use
        # defaults to us-central1
        location="asia-east1",

        # Google Cloud Storage bucket in same region as location
        # used to stage artifacts
        staging_bucket='gs://cloud-ai-platform-bce4682d-7df8-47d2-ba92-5522264ffdfe',

        # custom google.auth.credentials.Credentials
        # environment default creds used if not set
        # credentials=GOOGLE_APPLICATION_CREDENTIALS,

        # customer managed encryption key resource name
        # will be applied to all Vertex AI resources if set
        # encryption_spec_key_name=my_encryption_key_name,

        # the name of the experiment to use to track
        # logged metrics and parameters
        experiment='Panel spec extraction experiment 1',

        # description of the experiment above
        experiment_description='try predict Panel spec extraction'
    )

\ GCP

    path = "d:\\hcchen\\OneDrive\\Documents\\Jupyter Notebooks\\GCP\\artful-fortress-353205-707cd94101e5.json"
           "d:\hcchen\OneDrive\Documents\Jupyter Notebooks\GCP\artful-fortress-353205-707cd94101e5.json" 
    !dir "d:\hcchen\OneDrive\Documents\Jupyter Notebooks\GCP\artful-fortress-353205-707cd94101e5.json" 
    gcs_source

    C:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\GCP\vertex-ai-samples-main\notebooks\community\migration\UJ4 AutoML for structured data with Vertex AI Regression.ipynb

    %env GOOGLE_APPLICATION_CREDENTIALS artful-fortress-353205-707cd94101e5.json
    os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    my_gcp_project_id = "artful-fortress-353205"
    my_gcp_cloud_storage_bucket_name = "cloud-ai-platform-bce4682d-7df8-47d2-ba92-5522264ffdfe"
    my_gcp_cloud_storage_location = "asia-east1" #  (Taiwan)
    path_to_my_gcp_credentials = r"c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\GCP\artful-fortress-353205-707cd94101e5.json"
    my_credentials = "artful-fortress-353205-707cd94101e5.json"

    %f aiplatform :> TabularDataset.list()[0] dir -->
    %f aiplatform :> TabularDataset.list()[0].display_name -->
    %f aiplatform :> TabularDataset.list()[0].to_dict() -->

    %f aiplatform :> Model.list()[0].to_dict() --> 


    peforth.bp(33,locals())
    print("----Key-value pairs found in document----")
    for kv_pair in result.key_value_pairs:
        peforth.bp(44,locals())
        if kv_pair.key:
            print(
                    "Key '{}' found within '{}' bounding regions".format(
                        kv_pair.key.content,
                        format_bounding_region(kv_pair.key.bounding_regions),
                    )
                )
        if kv_pair.value:
            print(
                    "Value '{}' found within '{}' bounding regions\n".format(
                        kv_pair.value.content,
                        format_bounding_region(kv_pair.value.bounding_regions),
                    )
                )
    peforth.bp(55,locals())
    def format_bounding_region(bounding_regions):
        if not bounding_regions:
            return "N/A"
        peforth.bp(66,locals())
        return ", ".join("Page #{}: {}".format(region.page_number, format_bounding_box(region.bounding_box)) for region in bounding_regions)


    kv_pair :> to_dict() -->
    kv_pair :> key -->
    kv_pair :> value -->
    kv_pair :> key.bounding_regions -->
    kv_pair :> key.content -->

    format_bounding_region result <py>
    format_bounding_region, result = pop(1), pop(0)
    for entity in result.key_value_pairs: # result.entities:
        # print("Entity of category '{}' with sub-category '{}'".format(entity.category, entity.sub_category))
        # print("...has content '{}'".format(entity.content))
        print("...within '{}' bounding regions".format(format_bounding_region(entity.key.bounding_regions)))
        print("...with confidence {}\n".format(entity.confidence))
    </py>

        
    result dir --> .entities
    result :> to_dict() -->

    document_analysis_client = DocumentAnalysisClient(endpoint=endpoint, credential=AzureKeyCredential(key))
    document_analysis_client.begin_analyze_document_from_url?

    %f df :> shape -->
    %f DIF count nip -->
    %f DEA count nip -->
    %f DIF -->
    df['dif'] = DIF
    df['dea'] = DEA

    # 直接想要撈出 2021-01-01 不行，可能不存在
    # d0_row = df[df.date>=d0_str][0:1] # 這才是第一天
    # 不用管 d0 直接用 index 0 

    # DI
    df['DI'] = (df['max'] + df['min'] + 2*df['close'])/4
    # EMA12
    ema12 = np.zeros(df.shape[0])
    base = ema12[11] = df.DI[0:12].mean()
    for i in range(12,df.shape[0]):
        base = ema12[i] = (base*11+df.DI[i]*2)/13
    df['EMA12'] = ema12 

    # EMA26
    ema26 = np.zeros(df.shape[0])
    base = ema26[25] = df.DI[0:26].mean()
    for i in range(26,df.shape[0]):
        base = ema26[i] = (base*25+df.DI[i]*2)/27
    df['EMA26'] = ema26

    # dif1226
    dif1226 = df.EMA12-df.EMA26
    dif1226[0:25] = 0
    df['DIF1226'] = dif1226

    # macd9 
    macd9 = np.zeros(df.shape[0])
    base = macd9[33] = df.DIF1226[33+1-9:33+1].mean()
    for i in range(34,df.shape[0]):
        base = macd9[i] = (base*(9-1) + df.DIF1226[i]*2)/(9+1)
    df['MACD9'] = macd9

    /data/twse_2330_2000-01-04_2022-06-24.pkl

    df.query('date >= "%s"' % d0_str)
    df.query('date >= "%s" and date <= "%s"' % (d0_str, end_date_str))
    d0.strftime('%Y-%m-%d')
    start_date.strftime('%Y-%m-%d')
    obj.tax

    # Annual return
    Annual_return = (df.query('signal != 0').iloc[-1].total_profit/trader_fund)**(365/(end_date - d0).days) - 1
    Annual_return *= 100

    trade_total
    datetime.datetime(2021,1,1) -datetime.datetime(2017,9,23) # 1009
    ((1372317.54/1000000)**(365/1115) - 1)*100 # 10.14
    start_date_str
    end_date_str


    from PyPDF2 import PdfFileReader, PdfFileWriter
    pdf_reader = PdfFileReader(r"c:\Users\hcchen\Downloads\290-0922417555_11105_notification.pdf")
    pdf_reader.decrypt("F122172789")
    pdf_writer = PdfFileWriter()
    for page in range(pdf_reader.getNumPages()):
        # page is [0,1,2... ] page numbers
        pdf_writer.addPage(pdf_reader.getPage(page))
    with open(r"c:\Users\hcchen\Downloads\out.pdf", "wb") as out:
        pdf_writer.write(out)

    %f pdf_reader dir -->
    pdf_reader.getPage(0)


    from PyPDF2 import PdfFileReader, PdfFileWriter
    pdf_reader = PdfFileReader(r"c:\Users\hcchen\Downloads\290-0922417555_11105_notification.pdf")
    pdf_reader.decrypt("F122172789")
    with Timer():
        jsons = []
        for page in pdf_reader.pages:
            
            # page to temp.pdf
            pdf_writer = PdfFileWriter() # new writter object
            pdf_writer.addPage(page) # load with the current page 
            with open(r"temp.pdf", "wb") as out: # save page to temp.pdf 
                pdf_writer.write(out)

            # temp.pdf to file.io
            fhandle = !curl -F "file=@temp.pdf" https://file.io 2> nul:
            %f fhandle :> [0] txt2json 
            docurl = peforth.pop()['link']
            # file.io to Azure Form Recognizer 
            # result in json
            # get jsons[] 
            jsons.append(form_recognize(docurl))
    # Wall time: 54751.35350227356 ms
    # jsons[] to excel by openpyxl seems be easy 

\ FinMind 
    %f jsons count nip -->
    %f jsons :> [2].to_dict() char jsons.txt json2file \ 帳單第三頁
    tables = jsons[2].tables
    %f tables count --> nip  # 第三頁裡有 2 張 table 
    tables[0].row_count
    tables[0].column_count
    tables[0].cells[0]
    tables[0].cells[23]
    tables[1].cells[8]
    tables[1].cells[16]
    tables[1].cells[24]
    tables[1].cells[32]
    tables[1].cells[8]
    tables[1].row_count
    tables[1].column_count

    for i in range(0,80):
        print(i,tables[1].cells[i].content)
        if (i+1) % 8 == 0: print()
        
    command = index_stock_price.query('date == "%s"' % date).iloc[0].cmd
    1000000 - 249.5*4000*(1+fee) # trade_detail 2018/3/7 signal, 3/8 open, traider_fund   577.84
    1000000 - 248.0*4000*(1+fee) #        df_bt 2018/3/7 signal, 3/7 open,      balance  6586.4
    1000000 - 247.0*4000*(1+fee) #        df_bt 2018/3/7 signal, 3/7 close,     balance 10592.1

    all_stock_price.query('date == "%s"' % "2018-03-07").iloc[0].open
    all_stock_price.query('date == "%s"' % "2018-03-07").open.values[0]
    all_stock_price.query('date == "%s"' % "2018-03-07").open.iloc[0]
    all_stock_price.query('date == "%s"' % "2018-03-07").open.[0] 不能這樣寫，因為第一個 index 是 4519 而非 0 !!!
    all_stock_price.query('date == "%s"' % "2018-03-07").close.iloc[0]
    index_stock_price.query('date == "%s"' % "2018-03-07").cmd.iloc[0]
    trade_detail.query('date == "%s"' % "2018-03-09").signal.iloc[0]

    idx = df_open.query('signal != 0').index[-1]+1 # 最後一個 command 的下一天
    total_profit = df_open.iloc[idx].total_profit

    idx = df_close.query('signal != 0').index[-1]+1 # 最後一個 command 的下一天
    total_profit = df_close.iloc[idx].total_profit

\ access windows registry 
    import winreg
    %f winreg dir -->
    key = winreg.OpenKey(
        winreg.HKEY_CURRENT_USER,
        r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders"
    )
    dir, type = winreg.QueryValueEx(key, "AppData")

\ get environment variable PATH 
    with winreg.ConnectRegistry(None, winreg.HKEY_CURRENT_USER) as root:
        with winreg.OpenKey(root, "Environment", 0, winreg.KEY_ALL_ACCESS) as key:
            # winreg.SetValueEx(key, "PATH", 0, winreg.REG_EXPAND_SZ, value)
            value,type = winreg.QueryValueEx(key, "PATH")

\ get environment variable PATH 
    with winreg.ConnectRegistry(None, winreg.HKEY_CURRENT_USER) as root:
        with winreg.OpenKey(root, "Environment", 0, winreg.KEY_ALL_ACCESS) as key:
            # winreg.SetValueEx(key, "PATH", 0, winreg.REG_EXPAND_SZ, value)
            value,type = winreg.QueryValueEx(key, "PATH")
        
\ Access registry key 
    import winreg as reg
    reg_path = r'SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\chrome.exe'
    reg_key = reg.OpenKey(reg.HKEY_LOCAL_MACHINE, reg_path, 0, reg.KEY_READ)
    value      = reg.QueryValue  (reg_key, None)
    value,type = reg.QueryValueEx(reg_key, "Path")  # 注意！ 用 reg.QueryValueEx() 否則出錯。

\ Access environment variables 
    %env path
    # C:\Users\hcchen
    homepath = %env USERPROFILE

\ Show message box on Jupyternotebook
    說穿了就是 call 用 win32, user32, windll 等，那等於 Windows 所有的功能了！重點 ctypes library 是 standard library ;-) 
    # Message box #_msgbox_ #_message box_ #_python_ #_jupyter_ #_jupyternotebook_ #_jupyterlab_ #_win32_ #_windll_  #_user32_  
    # https://stackoverflow.com/questions/2963263/how-can-i-create-a-simple-message-box-in-python
    import ctypes  # An included library with Python install.   

    # ctypes.windll.user32.MessageBoxW(0, "Your text", "Your title", 1)
    ctypes.windll.user32.MessageBoxW(0, "Continue or abort", "Hochi's strategy", 1)

\ 取得 FinMin Annual Return 
    obj :> final_stats.AnnualReturnPer --> -5.69 (<class 'numpy.float64'>)

    2330 2001-1-1 2022-6-24 300 balance 1.0 2.0 3.0 

\ check a year has how many trading days? 52 weeks * 5 week days = 260 that's it. 
\ 從第 36 天依序掃描到最後一天，不用 random 亂跳。好看得出來景氣變化的影響。

    import random
    for i in range(100):
        print(random.randint(0,all_stock_price.shape[0]-260))
        
    all_stock_price.index.values[0:100]
    total how many days in all_stock_price.shape --> (5571,10)

    df = all_stock_price.iloc[start:start+300].copy()
    datetime.datetime.strptime("2000-1-1", "%Y-%m-%d")
    (datetime.datetime.strptime(df.date.iloc[299], "%Y-%m-%d") - datetime.datetime.strptime(df.date.iloc[0], "%Y-%m-%d")).days

    Stop running https://stackoverflow.com/questions/73663/how-do-i-terminate-a-script
        # 中止 running cells 有很多種方式 
        # 溫和，不會使 jupyternotebook kernel restart 但都會觸發如 SystemExit 的紅字 
        raise SystemExit 
        sys.exit()
        assert False
        # 兇猛，會使 jupyternotebook kernel restart 
        import Ipython; app = IPython.Application.instance(); app.kernel.do_shutdown(True);
        quit()
        exit()
        os._exit(0)

    target_tables = [(row.filename,row.sheetname) for row in df[df['PredictedLabel'] == 1].itertuples()]

\ read excel 
    import openpyxl
    pathname = r"c:\Users\8304018\Downloads\charlenge_training3.xlsx"
    wb = openpyxl.load_workbook(pathname) #path to the Excel file
    %f wb dir -->
    %f wb :> ["工作表1"] -->
    %f wb :> worksheets -->
    %f wb :> worksheets[0] -->
    ws = wb.worksheets[0]

    %f openpyxl :> workbook dir -->

    list(zip(df.date[df.signal != 0],df.signal[df.signal != 0]))

    df[df['PredictedLabel'] == 1][['filename','sheetname']]

\ 自動抓 OneDrive path 以及 working directory 
    onedrive = %env ONEDRIVE
    onedrive += "\\"
    working_path = !cd
    working_path = working_path[0]

    for filename in set(target_tables.filename): # get filenames
        target_sheets = set(target_tables[target_tables['filename']==filename].sheetname)
        wb = openpyxl.load_workbook(find(filename, datapath)[0])
        for s in wb.sheetnames:
            # 把 target_sheets 以外的都殺掉
            if s not in target_sheets:  
                wb.remove(wb[s])
        wb.save("filtered_" + filename)

\ 抽 pdf 某幾頁

    from PyPDF2 import PdfFileReader, PdfFileWriter
    for filename in set(target_tables.filename): # get filenames
        target_sheets = tuple(target_tables[target_tables['filename']==filename].sheetname) # target sheet names (str)
        target_sheets = [tuple(map(int,s[1:].split('_'))) for s in target_sheets] # sheet name string 轉成 (page#,table#)

        pdf_filename = filename.split("_table.xlsx")[0]
        pdf_reader = PdfFileReader(find(pdf_filename, datapath)[0])
        pdf_writer = PdfFileWriter()

        for page_num,_ in target_sheets:
            pdf_writer.addPage(pdf_reader.getPage(page_num-1))
            
        with open(datapath + "filtered_" + pdf_filename, "wb") as out:
            pdf_writer.write(out)
    
\ FinMind
    df_close.iloc[0].date
    df_open.iloc[0].date
    df_close.iloc[0].date
    df_open.iloc[0].close

    index_stock_price.iloc[217]    
        
        
    stock_price :> ['date'][i] >= "2003-09-19":    
    %f i -->
    stock_price :> ['date'][32] -->
    be 22 33 44
    dropall stock_price :> ['open'][33] hold_price :> [32] .s
    all_stock_price.iloc[start:start+300]
    min(all_stock_price.shape[0]-1,start+300)
    min(all_stock_price.shape[0]-1,i+299)

    df = add_ma_column(all_stock_price, 5)
    index_stock_price['5-60'] = index_stock_price.ma5 - index_stock_price.ma60
    index_stock_price['5-10'] = index_stock_price.ma5 - index_stock_price.ma10

    def add_sig_columns(dfi: pd.DataFrame) -> pd.DataFrame:
        df = dfi.copy() # make a copy for safe 
        sig10 = np.zeros(df.shape[0])
        sig60 = np.zeros(df.shape[0])
        for i in range(df.shape[0]):
            if df['5-60'][i-1] > 0 and df['5-60'][i] <= 0 :
                sig60[i] = -1 # downward 
            if df['5-60'][i-1] < 0 and df['5-60'][i] >= 0 :
                sig60[i] = +1 # upward 
            if df['5-10'][i-1] > 0 and df['5-10'][i] <= 0 :
                sig10[i] = -1 # downward 
            if df['5-10'][i-1] < 0 and df['5-10'][i] >= 0 :
                sig10[i] = +1 # upward 
        df['sig10'] = sig10
        df['sig60'] = sig60
        return df 

# --- review hubble2 NN -------------------------------
    df = scan_begin_end("2022-1-1","2022-1-10")

    all_columns

    from elasticsearch import Elasticsearch
    from elasticsearch_dsl import Search
    client = Elasticsearch("http://10.30.87.15:9200")

    begin = "2022-01-01"
    end = "2022-01-10"
    all_columns = ['bug_reproduce_procedure',
         'shortdescription',
         'bug_solution',
         'last_updated',
         'bug_description',
         'bug_update_comment',
         'rootcause',
         'customersystem',
         'general_componentType',
         'project_code',
         'testprocedure',
         'component',
         'affected_component',
         'longdescription',
         'bug_history',
         'bug_id',
         'bug_subject',
         'root_cause',
         'id',
         'comment',
         'description',
         'reproducesteps',
         'title',
         'comments',
         'communication',
         'history',
         'howtorecover',
         'longsummary',
         'reproduceprocedure',
         'suggestion',
         'summary',
         'technicalrootcause',
         'testcasenumber',
         'caseid',
         'softwarerootcause',
         'patestcasenumber',
         'remarks',
         'symptom',
         'correctiveaction',
         'issuediscoverymethod',
         'preventiveaction',
         'stepstoreproduce',
         'componenttype',
         'bugsolution']

    s = Search(using=client, index="hubble2") \
        .query("exists", field="bug_id") \
        .filter("range",last_updated = {"gte": begin,"lt": end}) \
        .extra(track_total_hits=True) \
        .source(all_columns)
    hits = s.scan()

    df = pd.DataFrame()
    for hit in s.scan():
        d = { k:hit[k] for k in hit if k!="meta"}
        d["id"] = hit.meta['id'] # 唯一不重複的 id 在 Kibana 上的 meta columns 內 
        df = df.append(d, ignore_index=True)

\ ----- converting pdf to word docx -----------------
\ the study is mainly on jeforth.3hta and the 

    %%time 
    if __name__ == '__main__':    

        # 把整個 path 之下的所有 PDF files 轉成 .docx files.
        # 這些 PDF files 都是 EE 零件手冊.

        import os 
        %env path d:\GitHub\jeforth\jeforth.3hta
        
        onedrive_path = %env ONEDRIVE
        onedrive_path += "\\"
        working_path = !cd
        working_path = working_path[0]
        datapath = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction" + "\\"
        
        pdf_list = [] # pdf file pathname list 
        for root, dirs, files in os.walk(datapath):
            for filename in files:
                pathname = os.path.join(root, filename).strip()
                if pathname.lower().endswith('pdf'): 
                    pdf_list.append({"pathname":pathname,"path":root,"filename":filename})

        for d in pdf_list:
            pathname = d['pathname'] # 一個 pdf 檔，以下供 command line $pathname 展開要用到。
            print("Converting : %s" % pathname)
            # !start /wait jeforth.hta include word.f pdf2docx $pathname
            !start /wait 3hta.bat include word.f pdf2docx $pathname

        print("*"*80)
        print("Done " * 16)
        print("*"*80)    
        
# PDF to Word 
    %env path d:\GitHub\jeforth\jeforth.3hta
    !start /wait 3hta.bat include word.f pdf2docx d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction\lot1\test\MDL_133WUXGA+MLOC_Lenovo_2D_210628.pdf

    !start /wait 3hta.bat include word.f pdf2docx D:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction\lot1\training\1. Lynx Serval_CS21 Product_Spec-Rev D_20201125-20201126 sign by NEC Sugimoto-san.pdf 

    %env path d:\GitHub\jeforth\jeforth.3hta
    !start /wait 3hta.bat include word.f pdf2docx d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction\lot2\NCPD_LM156LFDL02_Spec_AV1.0_20200302_VNCT2.pdf

    import win32com.client # !pip install pywin32 or even better conda install pywin32 if ImportError: DLL load failed while importing win32api
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.visible=True
    %f word_app _dir_ --> # 沒啥用
    %f word_app :> Documents --> \ <COMObject <unknown>> (<class 'win32com.client.CDispatch'>)
    %f word_app :> documents --> \ 故意用小寫 <COMObject <unknown>> (<class 'win32com.client.CDispatch'>)
    %f word_app :> Documents.Count --> # 查現在有幾個 word 檔 open 著
    %f word_app :> documents.count --> # 故意用小寫 OK，但是 .Open() .Close() 就必須講究。
    %f pathname word_app :: documents.Open(FileName=pop()) # Open 就必須大寫，小寫不行！

    %f word_app :> FileConverters.count --> 4 (<class 'int'>)
    %f word_app :> FileConverters(1).Path --> C:\Program Files\Common Files\Microsoft Shared\TEXTCONV (<class 'str'>)
    %f word_app :> FileConverters(2).Path --> C:\Program Files\Common Files\Microsoft Shared\TEXTCONV (<class 'str'>)
    %f word_app :> FileConverters(3).Path --> C:\Program Files\Common Files\Microsoft Shared\TEXTCONV (<class 'str'>)
    %f word_app :> FileConverters(4).Path -->       # {A5C79653-FC73-46ee-AD3E-B64C01268DAA} (<class 'str'>)
    %f word_app :> FileConverters(1).FormatName --> # 復原任何檔案的文字 (<class 'str'>)
    %f word_app :> FileConverters(1).classname -->  # Recover (<class 'str'>)
    %f word_app :> FileConverters(1).Extensions --> # * (<class 'str'>)
    %f word_app :> FileConverters(2).FormatName --> # WordPerfect 6.x (<class 'str'>)
    %f word_app :> FileConverters(2).classname -->  # WordPerfect6x (<class 'str'>)
    %f word_app :> FileConverters(2).Extensions --> # wpd doc (<class 'str'>)
    %f word_app :> FileConverters(3).FormatName --> # WordPerfect 5.x (<class 'str'>)
    %f word_app :> FileConverters(3).classname -->  # WrdPrfctDos (<class 'str'>)
    %f word_app :> FileConverters(3).Extensions --> # doc (<class 'str'>)
    %f word_app :> FileConverters(4).FormatName --> # PDF Files (<class 'str'>)
    %f word_app :> FileConverters(4).classname -->  # IFDP (<class 'str'>)
    %f word_app :> FileConverters(4).Extensions --> # pdf (<class 'str'>)
    %f word_app :> UserName --> \ H.C. Chen/WHQ/Wistron (<class 'str'>) # on LBB2 OA

    %f word_app :> ActiveDocument --> \ # Failed in </py> (compiling=False): (-2147352567, '發生例外狀況。', (0, 'Microsoft Word', '因為沒有開啟文件，所以無法使用這個指令。', 'wdmain11.chm', 37016, -2146824040), None)
    %f word_app :> ActiveDocument --> \ NCPD_LM156LFDL02_Spec_AV1.0_20200302_VNCT2.docx (<class 'win32com.client.CDispatch'>)
    %f word_app :> ActiveDocument.Close() # 傳回 None (沒有傳回值)

    \ 故意 manually 開了兩個 files  
    %f word_app :> Documents(1).name --> YTXJK.docx (<class 'str'>)
    %f word_app :> Documents(2).name --> NCPD_LM156LFDL02_Spec_AV1.0_20200302_VNCT2.docx (<class 'str'>)
    %f word_app :> Documents(3).name --> 
    %f word_app :: Documents(3).Activate() --> None (<class 'NoneType'>)
    %f word_app :: Documents(2).Activate() --> None (<class 'NoneType'>)
    %f word_app :> ActiveDocument.name --> # 切來切去無誤

    pathname = r"d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp\filtered_Cyborg XC0MJ_B160QAN01_HW_0A_ functional spec V01 (A00) - 2021-04-01 for Dell_Final.pdf"
    %f pathname word_app :: documents.Open(FileName=pop()) # 沒有 ConfirmConversions=bool 就是自動決定，果然 pdf 沒有跳出來要求 confirm format 種類。
    %f word_app :: documents.Open(FileName=pop(),ConfirmConversions=False) # Failed in </py> (compiling=False): (-2147352571, '類型不符。', None, 2)
    %f word_app :: documents.Open(FileName=pop(),ConfirmConversions=True) # 跳出選單要求 user 選擇 file format 種類 (pdf, word perfect 等)
    %f pathname word_app :: documents.Open(FileName=pop(),ConfirmConversions=False,NoEncodingDialog=True,Revert=True) # 都不要問，順暢執行。
    %f word_app :: Documents(1).SaveAs2(FileName="1.docx") # no return value, saved to: (%env USERPROFILE) + "\\Documents\\1.docx" which is c:\Users\8304018\Documents\1.docx 
    %f word_app :: ActiveDocument.Close() # .Close() method doesn't return. Close 必須大寫。

    \ 這跟 winword 無關，只是嘗試怎麼取出 pathname 去除 .ext --> [ ] 應該 officially 採用 pathlib library https://docs.python.org/zh-tw/3/library/pathlib.html
    ss = r"      d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp\filtered_Cyborg XC0MJ_B160QAN01_HW_0A_ functional spec V01 (A00) - 2021-04-01 for Dell_Final.pdf     "
    %f ss :> strip().lower().rindex('.pdf') --> 165 (<class 'int'>)
    %f ss :> strip()[:165] --> d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp\filtered_Cyborg XC0MJ_B160QAN01_HW_0A_ functional spec V01 (A00) - 2021-04-01 for Dell_Final (<class 'str'>)

    %f word_app :> ActiveDocument value doc // ( -- Word Document ) 
    %f doc --> \ NCPD_LM156LFDL02_Spec_AV1.0_20200302_VNCT2.docx (<class 'win32com.client.CDispatch'>)
    %f doc :> tables --> \ <COMObject <unknown>> (<class 'win32com.client.CDispatch'>)
    %f doc :> tables.count --> \ 17 (<class 'int'>)
    %f doc :> tables(1).columns.count --> 5 (<class 'int'>) \ correct !
    %f doc :> tables(1).rows.count --> 29 (<class 'int'>) \ super correct !!!! 因分頁被切割的 table 也能正確視為一個！
    %f doc :> paragraphs.count --> 2492 (<class 'int'>)

    %f doc :> paragraphs(3) --> 印出 printable 的東西 (<class 'win32com.client.CDispatch'>)
    %f doc :> Paragraphs(1).Range.Start -->  0
    %f doc :> Paragraphs(1).Range.End   --> 23
    %f doc :> Paragraphs(2).Range.start --> 23(number)
    %f doc :> Paragraphs(2).Range.End --> 49(number)
    %f doc :> Range(0,23) --> \ 查看 paragraph 1 

    for i in range(800,1000):
        print(i,peforth.execute("doc").pop().paragraphs(i)) 
        # 印 paragraph 出來看看，有東西。

    %f 868 >x ." ----- " cr 
    %f x@ doc :> Paragraphs(pop()).Range.Start
    %f x@ doc :> Paragraphs(pop()).Range.End
    %f ( start end ) doc :> Range(pop(1),pop(0)) --> # doc.Range() 尾巴有個 CR 但是缺 LF 所以印出來會怪怪的。
    %f xdrop ." ====== " cr

    # 列印出所有 sentences 很花時間
        doc = peforth.execute("doc").pop()
        %%time 
        for i in range(2000,doc.sentences.count):
            print(i, doc.sentences(i))
        # Wall time: 19.9 s

        %%time 
        s = ""
        for i in range(2000,doc.sentences.count):
            s += "%d %s\n" % (i, doc.sentences(i))
        print(s)    
        # Wall time: 19 s 很驚訝，快沒多少。表示問題不在 print() 

    %f \ This snippet works! copy target table from one document to another 
    %f     word_app :> Documents(2).name --> filtered_B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx(string)
    %f     word_app :> Documents(1).name --> 文件1(string)
    %f     > word_app :> Documents(1) value newdoc
    %f     newdoc :> ActiveWindow.Panes.count --> 1(number)
    %f     \ 手工點選 winword.exe table 左上角全選，然後來 jeforth 執行這兩行，真的就把 table copy 過去了
    %f        doc :> ActiveWindow.Panes(1).Selection.Copy()  --> undefined(undefined) 沒有傳回值
    %f     newdoc :> ActiveWindow.Panes(1).Selection.paste()  --> undefined(undefined) 沒有傳回值

    \ Selection 
        %f word_app :: ActiveDocument.Select() \ 把整個 doc 都 mark 起來，沒有傳回值。
        %f doc :> ActiveWindow.Selection.Cut() \ mark 好一段文字再來執行，成功！傳回 None 表示沒有傳回值。
        %f doc :> ActiveWindow.Selection.Cut() \ 若無 selec 則 com_error: 此方法或屬性無法使用，因為物件是空的。

        %f word_app :> selection.type --> 2(number) \ 當前 active doc 選中一小段 text 或跨多種 types 
        %f word_app :> selection.type --> 1(number) \ 當前 active doc 沒有 selection 
        %f word_app :> selection.type --> 5(number) \ 當前 active doc 選中整個 table 
        %f word_app :> selection.type --> 4(number) \ 當前 active doc 選中 table 內幾個 cells 

        %f word_app :> selection value selection // ( -- obj ) Word Selection object 跟定當前 active document 
        %f selection :> type --> # 這個 selection object 動態隨時反映當時的 selection 但跟定某 document 
        
    \ selection.Collapse()
        %f word_app :> selection.Collapse()  \ cursor 跳到 selection 的開頭
        %f word_app :> selection.Collapse(1) \ cursor 跳到 selection 的開頭
        %f word_app :> selection.Collapse(0) \ cursor 跳到 selection 的結尾
        %f word_app :> selection.Collapse(Dddirection=0) \ cursor 跳到 selection 的結尾
        %f                                ^^^^^^^^^^^^ 在 jeforth.hta 亂寫都無所謂，只認位置；在 python 必須大小寫完全正確。
        %f word_app :> ActiveDocument.select() word_app :> selection.Collapse(0) \ 跳到整篇最後面

    \ 打算建新 doc 然後把 table copy 過去, 藉此清除 table 以外的東西
    \ https://www.thespreadsheetguru.com/blog/2014/5/22/copy-paste-an-excel-table-into-microsoft-word-with-vba

        %f word_app :> Documents.Add() --> 文件1 (<class 'win32com.client.CDispatch'>)
        %f word_app :> Documents['文件1'] --> 文件1 (<class 'win32com.client.CDispatch'>) 試試看果然可以 ;-)    
        %f word_app :> Documents['文件1'] value newDoc // ( -- obj ) my new word document in memory

    \ 成功把 clipboard 內容貼到文件最後面去
        %f newDoc :> Range().end
        %f newDoc :> Range(pop()-1).Paste()

    \ 取得 table object 
        %f word_app :> ActiveDocument.tables.count --> 1(number)
        %f word_app :> ActiveDocument.tables(1).range --> PIN NO. ... (object)

    \ This experiment gets Good news!! simply put new table after the last table merges them into one table!!!
        %f doc :> tables.count --> 17 (<class 'int'>) 
        %f doc :> tables(8).Range.Copy() \ copy a table from the spec
        %f newDoc :> Range().end ( int ) \ get the position number of the end of the document 
        %f newDoc :> Range(pop()-1).Paste() \ paste the above table to end of the document. DevTools.py 也成功！
        %f doc :> tables.count --> 17 (<class 'int'>) \ source no change
        %f doc :> tables(13).Range.Copy() \ copy a table from the spec 再來一個，完全不一樣的 table 
        %f newDoc :> Range().end ( int ) \ get the position number of the end of the document 
        %f newDoc :> Range(pop()-1).Paste() \ paste the above table to end of the document. DevTools.py 也成功！
        %f newDoc :> tables.count --> 1 (<class 'int'>)  \ 有趣的來了 only one table on the target document! why? Because merged into one!!
        %f newDoc :> tables(1).Delete() \ try to delete the ONE table and it's true the multiple tables in one is deleted at once. 

    \ This experiment cut() the table > clear the doc > paste() the table back so as to keep only the table. 很成功！ But this way handles only one table unsuitable for multiple tables.
        %f newDoc :> tables(1).Range.Cut() \ cut the table to clipboard
        %f newDoc :> Range().Delete() \ delete the entire document 
        %f newDoc :> Range().Paste() \ restore the table so as to clean the document only keeps the table

    \ 改寫成 python, paste something to the end of the document
        newDoc = peforth.execute("newDoc").pop()
        newDoc.Range(newDoc.Range().end-1).Paste() # paste the above table to end of the document. DevTools.py 也成功！

    \ 改寫 Sentence 
        %f word_app :> ActiveDocument.Sentences.count ( count )
        %f word_app :: ActiveDocument.Sentences(pop()-1).Text="hello_world!"  \ 注意用 :: 不能用 :> 

    \ 插入 a string 到 word 後面
        %f word_app :> ActiveDocument.Words(1).Select() \ select 1st word 
        %f     word_app :> selection.Collapse(0) \ 跳到 word 結尾。
        %f     s" (This is a test.) "  word_app :: ActiveDocument.Words(1).Text=pop(1) \ 白跳了，整個 word 都被改掉。
        %f     s" (This is another test.) "  word_app :> selection :: InsertAfter(pop()) \ 用 selection.InsertAfter() 就對了
        %f     s" (This is a test.) "  word_app :: ActiveDocument.Words(1).InsertAfter(pop()) \ 這一行抵上面全部

    \ copy multiple tables 到新檔，separated by paragraph to avoid unexpected table merge  

        %f newDoc :> tables.count --> # 0
        newDoc.Range(newDoc.Range().end-1).InsertParagraphAfter()
        %f doc :: tables(8).Range.Copy() \ copy a table from the spec
        newDoc.Range(newDoc.Range().end-1).Paste()
        
        %f newDoc :> tables.count --> # 1
        newDoc.Range(newDoc.Range().end-1).InsertParagraphAfter()
        %f doc :: tables(9).Range.Copy() \ copy a table from the spec
        newDoc.Range(newDoc.Range().end-1).Paste()
        
        %f newDoc :> tables.count --> # 2
        newDoc.Range(newDoc.Range().end-1).InsertParagraphAfter()
        %f doc :: tables(11).Range.Copy() \ copy a table from the spec
        newDoc.Range(newDoc.Range().end-1).Paste()

        %f newDoc :> tables.count --> # 3

    \ Use the above methods 1. clean the newDoc 2. copy all tables from doc to newDoc. 3. Close doc. 4. newDoc save with prefix filename
    \ 成功了，but why bother? why not read DataFrames from .docx tables directory?

        tables = []
        for i in range(1,doc.tables.count+1):
            tables.append(doc.tables(i))
        
        newDoc.Range().Delete() # delete all clean the new doc
        for t in tables:
            newDoc.Range(newDoc.Range().end-1).InsertParagraphAfter()
            t.Range.Copy() # 
            newDoc.Range(newDoc.Range().end-1).Paste()

    \ Read tables from docx file to pandas DataFrames        
    \ https://medium.com/@karthikeyan.eaganathan/read-tables-from-docx-file-to-pandas-dataframes-f7e409401370

        from docx import Document # !pip install python-docx
        document = Document(r"c:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\docx\ST133FN133AKF SPEC_V1.0_20210703.docx")

        import pandas as pd
        from docx import Document
        def read_docx_table(document,table_num=1,nheader=1):
            table = document.tables[table_num-1]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df = pd.DataFrame(data)
            if nheader == 1:
                df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
            elif nheader == 2:
                outside_col, inside_col = df.iloc[0], df.iloc[1]
                hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col,inside_col)))
                df = pd.DataFrame(data,columns=hier_index).drop(df.index[[9,1]]).reset_index(drop=True)
            elif nheader > 2:
                print("More than two headers not currently supported")
                df = pd.DataFrame()
            return df

    for i in range(len(document.tables)):
        print(i)
        display(read_docx_table(document, table_num=i))
        
    table = document.tables[15]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    df

    import win32com.client
    word_app = win32com.client.Dispatch("Word.Application")

    onedrive_path = %env ONEDRIVE
    onedrive_path += "\\"
    pathname = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Storage\Panel spec extraction\lot1\test\ST133FN133AKF SPEC_V1.0_20210703.pdf" 
    !notepad $pathname              # pathname 中含有怪字，用 notepad 證實該 pathname 仍有效。
    word_app = win32com.client.Dispatch("Word.Application")
    def pdf2docx(pathname):
        pathname.lower().rindex(".pdf")
        o_pathname = "%s.docx" % pathname[:pathname.lower().rindex(".pdf")]
        print("Converting %s . . . " % pathname)
        word_app.Documents.Open(
            FileName = pathname, 
            ConfirmConversions = False,
            NoEncodingDialog = True,
            Revert = True
            ) # 都不要問，順暢執行。
        word_app.ActiveDocument.SaveAs2(FileName = o_pathname)
        print("to %s" % o_pathname)

\ I learn python-docx   
\ https://ithelp.ithome.com.tw/articles/10225127 
\ python-docx 用來把 .docx tables 轉成 DataFrame 以便於 cook 出 feature columns 供 AutoML 的 training dataset 
\ and test dataset 之用。雖然 Word 會把 pdf 中跨頁的 tables 合併成一個，但仍有未被合併者而被視為兩個 tables。
\ python-docx 不適合用來合併這些上下 tables word_app 才適合，放在一起即是。因此以下的練習沒有實際應用到。

    import docx
    datapath = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp" + "\\"
    !dir "C:\Users\hcchen\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp\"
    try to read tables from one doc and copy to another 

    doc1 = docx.Document(pathname)

    doc1 = docx.Document(datapath + "SD11G97709+M133NW4J+R3+Spec+HW1+2（V3KS）.docx")
    doc2 = docx.Document()    
    paragraph = doc2.add_paragraph()
    paragraph._p.addprevious(doc1.tables[15]._tbl)
    paragraph._p.addprevious(doc1.tables[16]._tbl)
    paragraph._p.addnext(doc1.tables[5]._tbl) # 很奇怪，用 .addnext() 會少 row !! .addprevious() 則 OK 
    paragraph._p.addnext(doc1.tables[5]._tbl) # 很奇怪，用 .addnext() 會少 row !! .addprevious() 則 OK 
    doc2.save(r'c:\Users\hcchen\Downloads\demo.docx')
    # 效果很怪，多做幾次，每次都不一樣！！ 因為 doc1.tables[4]._tbl 看起來是 list 其實又像是 generater 每抓一次就會
    next 到下一個去，而且 len(doc1.tables) 還會減一！

    %f paragraph dir -->
    %f doc1 :> core_properties.author -->
    %f doc1 :> core_properties.category -->
    %f doc1 :> core_properties.comments -->
    %f doc1 :> core_properties.content_status -->
    %f doc1 :> core_properties.created -->
    %f doc1 :> core_properties.identifier -->
    %f doc1 :> core_properties.keywords -->
    %f doc1 :> core_properties.language -->
    %f doc1 :> core_properties.last_modified_by -->
    %f doc1 :> core_properties.last_printed -->
    %f doc1 :> core_properties.modified -->
    %f doc1 :> core_properties.revision -->
    %f doc1 :> core_properties.subject -->
    %f doc1 :> core_properties.title -->
    %f doc1 :> core_properties.version -->

\ MsgBox message box 
    import ctypes
    answer = ctypes.windll.user32.MessageBoxW(0, "Continue?", "Panel spec extraction", 1) # 1:OK 2:Abort 


\ 研究 "B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx" table# 錯離的問題
    doc --> \ B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx (<class 'win32com.client.CDispatch'>)
    page --> 19 (<class 'int'>)
    文件7  竟然 page 19-1 還是不對！ 應該是 table 17 才對。檢查 dataset 可能 .docx 裡的 table 數與 dataset 不符！
    檢查 B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx 在 dataset temp.csv 當中有幾個 table? <--- 會變！慢慢增加到穩定才能用。

    pathname = find("B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx", datapath)[0]
    pathname = find("1. Lynx & Serval_CS21 Product_Spec-Rev D_20201125-20201126 sign by NEC Sugimoto-san.docx", datapath)[0]
    doc_lib = docx.Document(pathname)
    doc_com = word_app.Documents.Open(FileName = pathname, ConfirmConversions = False,NoEncodingDialog = True,Revert = True)
    
    len(doc_lib.tables)  # 26
    doc_com.tables.count # 26

    \ 奇怪了，這樣一查，19-1 沒錯哇？
    t = doc_com.tables[19-1]
    newDoc = word_app.Documents.Add()
    t.Range.Copy()
    newDoc.Range(newDoc.Range().end-1).Paste()

    t = doc.tables[19]
    newDoc = word_app.Documents.Add()
    t.Range.Copy()
    newDoc.Range(newDoc.Range().end-1).Paste()
    
    newDoc.Range(newDoc.Range().end-1).Paste()

    read_docx_table(doc_lib,table_num=1)

\ 探測 doc_com 要花多少時間 ready? 某些 .docx 檔要幾十秒才 ready !!
\ 結果 tested two .docx files one has the problem the other has not.  
    pathname = find("B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx", datapath)[0]
    pathname = find("1. Lynx & Serval_CS21 Product_Spec-Rev D_20201125-20201126 sign by NEC Sugimoto-san.docx", datapath)[0]

    pathname = r"d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\docx\3F7D0.docx" 
    pathname = r"d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\docx\HGT09.docx" 
    

    import time 
    import win32com.client
    word_app = win32com.client.Dispatch("Word.Application")
    doc = word_app.Documents.Open(FileName = pathname, ConfirmConversions = False,NoEncodingDialog = True,Revert = True)
    for i in range(40):
        print("%d:%d " % (i, doc.tables.count), end="")
        time.sleep(1)
    doc.Close()

    \ 探測 doc_com 要花多少時間 ready? 
    \ same experiment but scan entire directory 

    import time, docx, win32com.client
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.visible=True
    
    file_list = [] # excel file pathname list 
    for root, dirs, files in os.walk(r"d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\docx"):
        for filename in files:
            pathname = os.path.join(root, filename)
            if pathname.endswith('docx'): 
                file_list.append({"pathname":pathname,"path":root,"filename":filename})
                                  
    for i in range(len(file_list)): # 逐一
        pathname = file_list[i]['pathname'] # 一個 excel 檔
        filename = file_list[i]['filename']
        doc_lib = docx.Document(pathname)
        doc_com = word_app.Documents.Open(FileName = pathname, ConfirmConversions = False,
                                          NoEncodingDialog = True,Revert = True)
        print("lib:%d com:%d %d %s" % 
              (len(doc_lib.tables), doc_com.tables.count, 
               len(doc_lib.tables)==doc_com.tables.count, filename))

    \ 77 個 .docx 當中找到了這些 files 是有問題的。可能長達幾十秒才 ready! 用 python-docx 比對可偵測出問題。 
    lib:12 com:11 0 2M5HF-B140UAN03.1+HW-0A_Preliminary-Functional+Spec_V0.1_0407Y21+(002).docx
    lib:22 com:21 0 4RRP5+B156XTN08.2+(HW_0A)+Pre-functional+spec.+-+2018-05-24+for+DELL.docx
    lib:26 com:24 0 B133UAN01.2 Functional Spec_1104Y20_Lenovo.docx
    lib:22 com:21 0 B133UAN01.3 (0A) Functional Spec _0904Y20.docx
    lib:22 com:21 0 B140HAN04.D_HP_0A_ Final Spec_V1.0_0311Y20_HP.docx
    lib:12 com:11 0 B140UAN03.1 HW-0A_Preliminary-Functional Spec_V0.1_0407Y21.docx
    lib:42 com:40 0 Fenrir+B152770W1.docx
    lib:43 com:42 0 Fenrir+B152805W1.docx
    lib:51 com:50 0 HGT09.docx
    lib:26 com:25 0 HHYCY.docx
    lib:52 com:46 0 LP140WFA-SPMA_Final_CAS_Ver_1 0_HP_200221.docx
    lib:58 com:57 0 M32FY.docx
    lib:18 com:15 0 SD11H76425+LN.VB102H001_CS22+Product_Spec-Rev+A_20211208_update.docx
    lib:50 com:47 0 TDVNC-LP140WU2-SPM1_Pre+CAS_Dell_v0.1_210315_v0.1+(002).docx

    \ This is my workaround
    \ https://stackoverflow.com/questions/72968289/python-pywin32-win32com-client-responses-wrong-table-before-microsoft-word-doc/72974238#72974238

    import time, docx, win32com.client
    word_app = win32com.client.Dispatch("Word.Application")

    doc_lib = docx.Document(pathname)
    doc_com = word_app.Documents.Open(FileName = pathname)
    while len(doc_lib.tables) != doc_com.tables.count:
        # need to wait probably a minute for doc_com to be really ready
        time.sleep(1)
                                      
    import win32com.client as win32
    xl = win32.gencache.EnsureDispatch('Excel.Application')
    xl.Visible = 1

    wb = excel_app.Workbooks.Open("%s.xlsx" % pathname[:pathname.lower().rindex(".docx")])  # New Excel file that is blank
    wb = excel_app.Workbooks.Add()
    workbook.SaveAs(new_file_path+'/UpdatedSheet.xls')

    xl.DisplayAlerts = False
    ws = wb.Worksheets('Sheet1')
    ws = wb.Worksheets(1)
    ws.name --> '工作表1'
    dest_cell = ws.Range('A1')
    obj = ws.OLEObjects()

\ 用 temp_labeled.csv 加上 word_app 來把 target table 從 .docx 中取出來，考慮有上下兩頁的情形。

    datapath = onedrive_path + r"Documents\Jupyter Notebooks\Vendor spec parsing\Storage\temp" + "\\"

    import docx,time,ctypes
    import win32com.client # !pip install pywin32 
        # or even better conda install pywin32 if ImportError: DLL load failed while importing win32api
    word_app  = win32com.client.Dispatch("Word.Application")
    excel_app = win32com.client.Dispatch("Excel.Application")
    word_app.Visible = excel_app.Visible = True # 正式版關掉
    
    # datapath 裡的 docx files cook 成 csv 讓 AI 打好了 label 得 temp_labeled.csv
    target_tables = pd.read_csv(
        r"d:\OneDrive\Documents\Jupyter Notebooks\Vendor spec parsing\Panel spec extraction\data\temp_labeled.csv", 
        encoding='utf-8')
    target_tables = target_tables[target_tables['PredictedLabel'] != 0] # 篩出 target tables 

    for i,filename in enumerate(set(target_tables.filename)): # get filenames 逐一
        target_sheets = set(target_tables[target_tables['filename']==filename].sheetname) # 固定屬於某 filename 的所有 worksheets 
        pathname = find(filename, datapath)[0]
        print("%d %s" % (i, filename))
        
        doc_lib = docx.Document(pathname) # 用來取得 table 個數當參考，必須在 word_app Open() 之前否則不明原因會失敗。
                                          # PackageNotFoundError: Package not found at ...
        time.sleep(1) # 隔遠一點
        doc = word_app.Documents.Open(
            FileName = pathname, 
            ConfirmConversions = False,
            NoEncodingDialog = True,
            Revert = True
            ) # 都不要問，順暢執行。
        
        # doc 可能尚未 ready 需要等待
        while len(doc_lib.tables) != doc.tables.count:
            print("Wait probably a minute for doc to be really ready.")
            time.sleep(1)
        
        # Create an empty new doc to merge target tables they are probably separated 
        newDoc = word_app.Documents.Add(); time.sleep(1)
        for page in target_sheets:
            t  = doc.tables[page]
            t.Range.Copy()
            newDoc.Range(newDoc.Range().end-1).Paste()
        newDoc.tables(1).Range.Copy() # Copy the one table
            
        # Create an empty excel to paste the table 
        wb = excel_app.Workbooks.Add()
        ws = wb.Worksheets(1)
        flag = True
        while flag:
            try:
                ws.Paste(Destination=ws.Range('A1'))
                flag = False
            except:
                time.sleep(1)
                pass
        

        # save the new doc and excel 
        #newDoc.SaveAs2(FileName = "%s.xlsx" % pathname[:pathname.lower().rindex(".docx")])
        wb.SaveAs("%s.xlsx" % pathname[:pathname.lower().rindex(".docx")])

        # 等一等
        answer = ctypes.windll.user32.MessageBoxW(0, "Continue?", "Panel spec extraction", 1) # 1:OK 2:Abort 
        
        # Close opened files proceed to next one 
        newDoc.Close(SaveChanges=False)
        wb.Close()
        doc.Close()
    print("Done !!!")
    #


\ Top Python Tricks That Will Boost Your Skills | by Haider Imtiaz | Python in Plain English
\ https://python.plainenglish.io/top-python-tricks-that-will-boost-your-skills-c55e8268ca5f
    👉1. Most Frequent Value in a List
        #example codemy
        mylist = [2, 2, 4, 5, 6, 4, 2, 3, 9, 3, 2, 4, 4, 7, 0, 1]
        freq = max(set(mylist), key = mylist.count) # mylist.count(3) --> 3 有出現 2 次。 
        print(freq) 
        #Output
        #2

    👉2. Printing String N Times  咱常用
    👉3. Get-Module Path          直接 display(peforth) 即見之
    👉4. Reverse String in Cool Way
        mystring = "this is a test"
        ''.join(reversed(mystring))
        ''.join([c for c in mystring[::-1]])

        #output
        'tset a si siht'
        'tset a si siht'


    👉5. Use of Enums in Python
    Enums is a class in Python 超簡單

        class MyEnums:
            a,b,c = 1, 2, 3

        print(MyEnums.a) # 1
        print(MyEnums.b) # 2
        print(MyEnums.c) # 3

    👉6. Get Memory Usage of an Object
    We will use sys built-in method getsizeof().

        #example code
        import sys
        var = 400
        var2 = "400"
        print(sys.getsizeof(var)) # 28
        print(sys.getsizeof(var2)) # 52

    👉7. Reversed Iteration
        mylist = [1, 2, 3, 4, 5]
        for x in reversed(mylist): # or mylist[::-1] samething 
            print(x)

        #output
        # 5
        # 4
        # 3
        # 2
        # 1

    👉8. Get Current Date and Time
        #example code
        import datetime
        print(datetime.datetime.now())
        #output
        # 2021-08-22 18:22:10.791426

    👉9. Chain Assignment
        #example code
        a = b = c = 2
        print(a, b, c) # 2 2 2

    👉10. Merging Dictionaries
        One way to do that is by unpacking them in a new dictionary.
        #example code
        x = {"a": 1}
        y = {"b": 2}
        z = {**x, **y}

        print(z) # {'a': 1, 'b': 2}

\ 整理 project-code vs chipsets lookup table 
\ 本來是按 project code 列出對應的 chipset. 轉製成按 chipset 列出對應的 project codes 

    # Sample target lookup table looks like this: 
    '''
    chipset2projectcode = {
        "Intel Alder Lake-P":{'QRQY00012429','4PD0RE01E001'},  # 用 chipset 帶出 project code 的子集合
        'AMD Barcelo-FP6':{'4PD0MM01B001','3PD0EU010001','3PD0E501B001','3PD0EA01B001','4PD0MX010001',}
        ... snip ....
        }
    '''
    
    df = pd.read_csv(r"c:\Users\8304018\Downloads\hubble2\output_chipsets 202208051152.csv")
     
   
    projectcode2chipsets = {}
    for i,row in enumerate(df.itertuples()):
        chipsets = eval(row.chipsets) # list of chipset names
        projectcode2chipsets[row.project_code] = projectcode2chipsets.get(row.project_code,set()).union(set(chipsets))

    chipset2projectcode = {} # 目標 dict 
    for i,row in enumerate(df.itertuples()):
        chipsets = projectcode2chipsets[row.project_code]
        for chipset in chipsets:
            set_pjcodes = chipset2projectcode.get(chipset,set())
            set_pjcodes.add(row.project_code)
        chipset2projectcode[chipset] = set_pjcodes
        
    # 雙向的 lookup table 都成功了。
    
    lookup_table_chipsets = chipset2projectcode.keys()
    %f lookup_table_chipsets count --> nip
    lookup_table_pjcodes = projectcode2chipsets.keys()
    %f lookup_table_pjcodes count --> nip # project vs chipset 對照表中的 project code 個數

\ 檢查 hubble2 project code vs chipset 對照表

    missing = set()
    existing = set()
    for i,s in enumerate(df.project_code): # 三年來的 TRs 
        pjcodes = {c.strip() for c in s.split(',') if c.strip()} # 一條 TR 可能有多個 project code 
        for pj in pjcodes:
            if pj not in lookup_table_pjcodes:
                missing.add(pj)
            else:
                existing.add(pj)
    %f missing count --> nip
    %f \ missing -->
    %f existing count --> nip
    %f \ existing -->

    # 再確定一遍， project code count 是 621 嗎？為何 '4PD0SD01A001' 找不到？顯然漏了。
    all_projects = set()
    for i,row in enumerate(df.itertuples()):
        pjcodes = set(row.project_code.split(','))
        all_projects = all_projects.union(pjcodes)
    len(all_projects) # 538 個，三年來只有 538 個。
    pj in all_projects # True 
    pj in projectcode2chipsets.keys() # False !!! 怎麼會這樣？ 檢查一遍看看
    
    i = 0 
    for p in all_projects:
        if p not in projectcode2chipsets.keys():
            print(p)
            i += 1
    print(i) # 確定缺了 176 個無誤，這下怎麼辦？ 缺的就當作吻合好了。怎麼做到？
        
    
\ hubble2 撈取 raw data from Elasticsearch directly 

    all_columns = ['bug_reproduce_procedure', 'shortdescription', 'bug_solution', 'last_updated', 'bug_description', 
        'bug_update_comment', 'rootcause', 'customersystem', 'general_componentType', 'project_code', 'testprocedure', 
        'component', 'affected_component', 'longdescription', 'bug_history', 'bug_id', 'bug_subject', 'root_cause', 'id', 
        'comment', 'description', 'reproducesteps', 'title', 'comments', 'communication', 'history', 'howtorecover', 
        'longsummary', 'reproduceprocedure', 'suggestion', 'summary', 'technicalrootcause', 'testcasenumber', 'caseid', 
        'softwarerootcause', 'patestcasenumber', 'remarks', 'symptom', 'correctiveaction', 'issuediscoverymethod', 
        'preventiveaction', 'stepstoreproduce', 'componenttype', 'bugsolution', 'originator', 'creator_employee_id', 'owner_employee_id','owner']

    begin = datetime.datetime(2019,8,1)
    end   = datetime.datetime(2022,8,1)

    s = Search(using=client, index="hubble2") \
        .query("exists", field="bug_id") \
        .filter("range",last_updated = {"gte": begin,"lt": end}) \
        .extra(track_total_hits=True) \
        .source(all_columns)

    df = pd.DataFrame()
    for i,hit in enumerate(s.scan()):
        if not (i % 1000):
            %f i --> 
        d = { k:hit[k] for k in hit if k!="meta"}
        d["id"] = hit.meta['id'] # 唯一不重複的 id 在 Kibana 上的 meta columns 內 
        df = df.append(d, ignore_index=True)

    fname = "20190801-20220801-raw-data.feather"
    df.reset_index().to_feather(path + fname )  # feather 有時候會要求先 reset_index

\ 證實一條 TR 對應多個 project codes 
    # 列出所有含多 project_code 的 TR 
    for i,row in enumerate(df.itertuples()):
        pjcodes = set(row.project_code.split(','))
        if len(pjcodes) > 1:
            print(row.bug_id)
    df.query('bug_id == "1774732"').project_code.iloc[0]  # 印一個來看看
    # 印出來看，確定了多 project code 的 TR 大量存在。該等 project code 欄位裡放的是逗點分開的 project codes 的 string 
    for id in BugIds:
        pjcodes = df.query('bug_id == "%s"' % id).project_code.iloc[0]
        %f pjcodes -->
    
\ df 只留下 target chipset
    # 專挑有多重 project code 的 TR 來實驗
    BugIdss = ''' 
        211054
        1739138
        BITS423492
        1776824
        BITS429159
        209136
        1759036
        1773928
        1773098
        1776948
        209587
        BITS429659
        BITS429655
        BITS429208
        1774172
        1775607
        205396
        '''
    BugIds = [i.strip() for i in BugIdss.split('\n') if i.strip() ]

\ 當該 TR 的 chipset 是空的可以考慮一律入選，這些缺 chipset info 的 TR 可以事先準備好。
    %%time
    wildcards = set() # index list 缺 chipset info 的 TRs
    for i,row in enumerate(df.itertuples()):
        tr_pjcodes = set(row.project_code.split(','))  # 這個 TR 的 project code(s)
        for pj in tr_pjcodes: # 每個 project code 的 chipsets 先都連吉起來。供後續檢查是否屬於 target chipset? 
            # 如果這個 project code 沒有 chipset info 就直接入選，一個 project code 入選就整條 TR 入選。
            if pj not in lookup_table_pjcodes:
                wildcards.add(row.Index)
    %f wildcards count --> nip # 這些是沒有 chipset info 的 TRs (index of df), 約 25%. 
    # sum([df.iloc[i].project_code in lookup_table_pjcodes for i in wildcards]) # check it out indeed 23K TRs w/o chipset info 

\ 求出單一 TR 的 chipsets 集合，與 target_chipsets 交集，有東西的 TR 入選。
    %%time
    df_chipset_index = []
    for i,row in enumerate(df.itertuples()):
        tr_selected_flag = False # 先假設本 row (TR) 沒有入選
        # 以下針對 row 這條 TR 檢查，看它的 chipset 屬不屬於 target chipsets 
        tr_pjcodes = set(row.project_code.split(','))  # 這個 TR 的 project code(s)
        tr_pjchips = set() # 這條 TR 的 chipsets 集合。可能有多個 project code 每個 project code 又有多個 chipset. 
        for pj in tr_pjcodes: # 每個 project code 的 chipsets 先都連吉起來。供後續檢查是否屬於 target chipset? 
            # 如果這個 project code 沒有 chipset info 就直接入選，一個 project code 入選就整條 TR 入選。
            tr_pjchips = tr_pjchips.union(projectcode2chipsets.get(pj,set()))  # 缺 chipset info 的 project 就當作空集合
        # 得到 tr_pjchips 
        if target_chipsets & tr_pjchips: 
            df_chipset_index.append(row.Index)
    %f df_chipset_index count --> nip
    # 合併 chipset screen 與 wildcards 的方法
    # df.iloc[[*wildcards,*df_chipset_index]].shape

\ 自動找出 TR 的 chipsets     
    target_tr = BugIds[0]
    # target_tr = 'PIMS-153777'
    # target_tr = '201239'
    if target_tr in wildcards_bugid: 
        print("Warning: Bug Id=%s has no chipset info." % target_tr)
    row = df.query('bug_id == "%s"' % target_tr)
    # 找出這條 TR 的 chipsets
    tr_pjcodes = set(row.project_code.values[0].split(','))  # 這個 TR 的 project code(s)
    tr_pjchips = set() # 這條 TR 的 chipsets 集合。可能有多個 project code 每個 project code 又有多個 chipset. 
    for pj in tr_pjcodes: # 每個 project code 的 chipsets 先都連吉起來。供後續檢查是否屬於 target chipset? 
        tr_pjchips = tr_pjchips.union(projectcode2chipsets.get(pj,set()))  # 缺 chipset info 的 project 就當作空集合
    %f tr_pjchips -->

\ I learn TiyaroAI

    import requests

    url = "https://api.tiyaro.ai/v1/ent/huggingface/1/facebook/bart-large-xsum"

    payload = {"input": "Machu Picchu was believed (by Richard L. Burger, professor of anthropology at Yale University) to have been built in the 1450s.However, a 2021 study led by Burger used radiocarbon dating (specifically, AMS) to reveal that Machu Picchu may have been occupied from around 1420-1530 AD. Construction appears to date from two great Inca rulers, Pachacutec Inca Yupanqui (1438–1471) and Túpac Inca Yupanqui (1472–1493).  There is a consensus among archeologists that Pachacutec ordered the construction of the royal estate for his use as a retreat, most likely after a successful military campaign. Although Machu Picchu is considered to be a royal estate, it would not have been passed down in the line of succession. Rather it was used for 80 years before being abandoned, seemingly because of the Spanish conquests in other parts of the Inca Empire. It is possible that most of its inhabitants died from smallpox introduced by travelers before the Spanish conquistadors arrived in the area."}
    headers = {
        "accept": "*/*",
        "content-type": "application/json",
        "authorization": "Bearer d1ntRU5BbNDkXlkzrBv5"
    }

    response = requests.request("POST", url, json=payload, headers=headers)

    print(response.text)

    %env TIYARO_API_KEY 1L6mGSSr00wooUsYi2ZQ.gCZxlGxCBaPc4NhiDQtb2p3PrNdK98Fn5ZrHikO3
    apiKey = os.getenv("TIYARO_API_KEY")

class Supermarket:
    product = "Milk"  # class attribute
    
    def __init__(self, product, best_before):
        self.best_before = best_before  # instance attribute
        self.product = product
    # @staticmethod    
    def normalize_product_name(self, product):
        product = product.capitalize().strip()        
        peforth.bp(11,locals())
        return product
        
    norm_product = Supermarket.normalize_product_name("milk  ")
    'Milk'
    >>> obj = Supermarket("Bread", "2022-05-18")
    >>> obj.normalize_product_name("milk  ")
    'Milk'    
        
    df0.to_feather(r"c:\Users\8304018\Downloads\hubble2\hubble2-df0.feather")

\ 分析 hubble2 components in df.affected_component 
    components_list = list(df.affected_component)
    %f components_list type -->
    %f components_list count --> nip
    components = []
    for line in components_list:
        pieces = [ piece.strip() for piece in line.split('-') if piece]
        components.append(pieces)
    %f components count --> nip
    component_set = {c[0] for c in components if c}
    %f component_set count --> nip # 645

    remove (affected_component)
        [HW] [ME] 'rd_ce' 'rd_hw.*' [KC]* 
        (Owned by Lenovo) 'option_*','option_Dongle', 'option_U3Dock', 'option_U3dock', 
        'option_USB', 'option_Ultraportable_BT_Speaker', 'option_hub', 'option_monitor_stander',
        'option_speaker', TA_* (Lenovo Dock) TV_* (Lenovo ThinkView) 
        (Owned by Acer) SpatialLabsDisplayService 
    保留

\ 'creator_employee_id', 'owner_employee_id' [ ] are all emails? 
    df0[['bug_id','customersystem', 'creator_employee_id', 'owner_employee_id', 'affected_component','general_componentType']]
    # 用 mito 來查
    
jason = dict()
for BugId,customers in results.items():
    for customer,bugs in customers.items():
        jason[customer] = []
        for i,bug in enumerate(bugs.itertuples()):  # 遍歷所有的 rows of a DataFrame 
            if bug.tcount < tcount_threshold[bug.customersystem] : continue # 字數太少的不要，很多都是草稿或 ME、EE的。
            subject = bug.bug_subject if bug.bug_subject else bug.shortdescription # 研究過了，主要是 bug_subject 否則就用 shortdescription.
            if not subject.strip(): continue # 有些缺 Subject 的
            bug_dict = {} # dict of this bug
            bug_dict['bug_id']=bug.bug_id
            bug_dict['subject']=subject
            bug_dict['id']=bug.id
            jason[customer].append(bug_dict)
            if i >= 99 : break
jason['query'] = {
        'BugIds': BugIds,
        'target_chipsets' : list(target_chipsets),
        'auto_chipset' :  auto_chipset,
        'include_chipset_unknown' :  include_chipset_unknown,
        'ands' :  ands
    }


    for k in results['2085170'].keys():
        %f k results :> ['2085170'][pop()].shape k . --> 
        
    for cus in jason.keys():
        %f cus jason :> [pop()] count cus . space --> nip 
        
    jason[cus]

\ setup
    with open("setup.py","r",encoding="utf-8") as f:
        exec(f.read())

\ input 
    with open("input.py","r",encoding="utf-8") as f:
        exec(f.read())    

\ webserver web api and flask

    from flask import *
    import json, time

    app = Flask(__name__)

    @app.route('/', methods = ['GET'])
    def home_page():
        data_set = {'Pageee':'Homeeee', 'Messageeee':'Successfully loaded the Home pageeee', 'Timestamp': time.time()}
        json_dump = json.dumps(data_set)
        return json_dump

    @app.route('/user/', methods = ['GET'])
    def request_page():
        user_query = str(request.args.get('user')) # /user/?user=STELLnaTalie
        data_set = {'Pagxxe':'Requestxxx', 'Messagexxx': f'Successfully got the request for {user_query}xxxx', 'Timestamp': time.time()}
        json_dump = json.dumps(data_set)
        return json_dump

    if __name__ == '__main__':
        app.run(port=7777)

\ Hubble2 用 creator 來研究 components        
    df = df0[[
        'bug_id','owner_employee_id','creator_employee_id',
        'affected_component','general_componentType','general_status','id'
        ]]

    %%time
    # 列出所有 creators 的集合
    creators = set(df.creator_employee_id)
    %f creators count --> nip
    creator_components = {}
    for creator in creators:
        creator_components[creator] = {}
        creator_df = df.query('creator_employee_id == "%s"' % creator)
        creator_components[creator]['components'] = components = set(creator_df.affected_component)
        creator_components[creator]['count'] = len(components)
    # Wall time: 6.91 s
    
    %%time
    # 列出所有 owner 的集合
    owners = set(df.owner_employee_id)
    %f owners count --> nip
    owner_components = {}
    for owner in owners:
        owner_components[owner] = {}
        owner_df = df.query('owner_employee_id == "%s"' % owner)
        owner_components[owner]['components'] = components = set(owner_df.affected_component)
        owner_components[owner]['count'] = len(components)
    # Wall time: 6.91 s
    
    # 查看 owners 的 component count 
    for i,owner in enumerate(owners):
        print(i, owner, owner_components[owner]['count'])
        
    # 沒有 owner 的 TR 應該可以忽略，慎重起見查一下他們的 status 
    idxs = [i for i,TR in enumerate(df.itertuples()) if TR.owner_employee_id == '']
    %f idxs count --> # 23922 無誤，與 mito 一致。   
    df.iloc[idxs]

\ debug DFT184886  text preprocess 的 bug 
\ DFT184886 逮出我程式裡一個大問題！text preprocess 意外地把 TR 的文字濾掉了一些，正好被你發現了！神！
    找到了 ： dfs[3].iloc[1100]

\ fix DFT184886 改了 text preprocessor 這下 tcount 要重訂 shreshold 了。這得從 corpus 從頭做起，很花時間。
\ 這段 snippet 算出所有 TR 的 tcount 9W 條要花 5 分鐘。

    %%time
    df = df0.copy()  # 方便反覆 debug 
    data = pandas_to_orange(df[['bug_id','Text','subject_weighted']])
    corpus = Corpus.from_table(data.domain,data)
    corpus.set_text_features([StringVariable('Text'), StringVariable('subject_weighted')]) # 指定 text columns 否則他亂設的不對。
    bow = BowVectorizer()
    bow = BowVectorizer(norm=bow.L2, wlocal=bow.COUNT, wglobal=bow.SMOOTH)
    vectors = bow.transform(corpus)
    vectors = Table.from_table(Domain(vectors.domain.attributes[1:], None),vectors) # 濾掉第一個 column which is named '?' that means NaN.
    X = vectors.X   # vectors.X_df if you want to view it
    tcount = [x.getnnz() for x in X] # vectors.X 橫向的 token 個數，越大字越多可能品質越好，太小鐵定不好。
    df['tcount'] = tcount # 增加一欄 
    # Wall time: 5min 27s

    # 保存好不容易製作好的 9W 條 TR 的 df with tcount 
    df.to_feather(r"c:\Users\8304018\Downloads\hubble2\tcounts 2022-8-26.feather" )

    # 本想把帶有 tcount 的 hubble2 df 存放在 file.io 長期保存，結果時間蠻短的。
    fileio1py = !curl -F "file=@c:\Users\8304018\Downloads\tcounts 2022-8-26.zip" https://file.io 2> nul: 
    %f fileio1py :> [0] txt2json ==>
    {'success': True, 
     'status': 200, 
     'id': 'e05a8170-24e2-11ed-a484-4de0b89a5b11', 
     'key': 'GPavXe0tpX5C', 
     'path': '/', 
     'nodeType': 'file', 
     'name': 'tcounts 2022-8-26.zip', 
     'title': None, 
     'description': None, 
     'link': 'https://file.io/GPavXe0tpX5C', 
     'private': False, 
     'expires': '2022-09-09T02:00:36.343Z', 
     'downloads': 0, 
     'maxDownloads': 1, 
     'autoDelete': True, 
     'size': 172753528, 
     'mimeType': 'application/octet-stream', 
     'screeningStatus': 'pending', 
     'created': '2022-08-26T02:00:36.343Z', 
     'modified': '2022-08-26T02:00:36.343Z'
    } (<class 'dict'>)

\ 沒有 owner 的 tcount 如何？
    df = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\tcounts 2022-8-26.feather")

    mito column formula
    if(or(owner_employee_id=="",owner_employee_id==" "),None,owner_employee_id)
    
    結果是不一定。

\ affected_component 有 4千多種，做成 BoW 來嘗試 clustering 分類
    components = {c.strip() for c in components} # components 來自執行 components.txt 第一行
    len(components) # 有 4392 之多！
    with open(r"c:\Users\8304018\Downloads\component.csv", "w", encoding='utf8') as f:
        for c in components:
            f.write('"%s"\n' % c)
    df = pd.read_csv(r"c:\Users\8304018\Downloads\component.csv", encoding='utf8')

\ Advanced exploratory data analysis (EDA)
\ hcchen/OneDrive/Documents/Jupyter Notebooks/miykael.ai/03_advanced_eda/nb_advanced_eda.ipynb

    from sklearn.datasets import fetch_openml
    # Download the dataset from openml
    dataset = fetch_openml(data_id=42803, as_frame=True)

    # Extract feature matrix X and show 5 random samples
    df_X = dataset["frame"]
    df_X.sample(5)
    
    檢查發現異常, 有個 id (Accident_Index) "201543P296025" 高達 1332 rows 應該有很多重複的。
    咱抄一下 hubble2NN 的 code 來給它算算有多少重複的。。。 
    see https://www.machinelearningplus.com/pandas/pandas-duplicated/
    
    df_X.duplicated().sum() # 竟然是零！
    
\ read .json file 
    import json
    from setup import path 
    with open(path + "input.json") as f:  # path from setup.py
        inputs = json.load(f)

\ watchdog  monitors a folder and take actions 
    import sys
    import logging
    from icecream import ic
    from watchdog.observers import Observer
    from watchdog.events import LoggingEventHandler
    import peforth
    if __name__ == "__main__":
        logging.basicConfig(level=logging.INFO,
                            format='%(asctime)s - %(message)s',
                            datefmt='%Y-%m-%d %H:%M:%S')
        path = r"c:\Users\8304018\Downloads\hubble2"
        logger = LoggingEventHandler()
        class event_handler:
            def dispatch(event):
                logger.dispatch(event)
                observer.stop() # so as to avoid multiple triggered modify events

    for i in (1,2,3):
        observer = Observer() # 這個 object 就是作者稱作 thread 的東西， .stop() 之後就不能再用了，只能 create 新的。
        observer.schedule(event_handler, path, recursive=True)
        observer.start()
        assert observer.is_alive(), "Fatal error, watchdog observer is not alive!!"
        observer.join() # infinit waiting 
        observer.stop()
    %f # The End !!
    
\ read .feather   
    import pandas as pd
    df = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\2021-6-1.feather")
        
\ logging 
    import logging    
    logfile_pathname = r"c:\Users\8304018\Downloads\hubble2.log" 
    FORMAT = '%(levelname)s %(asctime)s %(pathname)s Line:%(lineno)d %(message)s'
    def reset_logging(level=logging.DEBUG, pathname=r"c:\Users\8304018\Downloads\hubble2.log"):
        # Clear the log file 
        for i in logging.root.handlers: logging.root.removeHandler(i) # 一定要先把舊的清空，basicConfig 才有效。
        logging.basicConfig(level=level, filename=logfile_pathname if PROD else None, filemode='w', format=FORMAT)
    reset_logging()
    logging.info('End of setup.py')

\ 其實沒有 peforth Python 自己就可以變出 repl  (from peforth log.txt)
        # https://dev.to/amal/building-the-python-repl-3468
        def repl() -> None:
            try:
                while True:
                    try:
                        _in = input(">> ")
                        try:
                            print(eval(_in))
                        except:
                            out = exec(_in)
                            if out != None:
                                print(out)
                    except Exception as e:
                        print(f"Error: {e}")
            except KeyboardInterrupt as e:
                print("\nExiting...")

\ Scratchpad debugging (from peforth log.txt also OneNote2022>2022 )

    從 function 裡面可以把程式停掉來檢查 context，甚至在 peforth.bp() 裡面也可以做得到，命令是 
    
        raise SystemExit("Stop right there!") 
        
    配合 magic %tb and %debug 很方便查看當時的現況。peforth 的 _locals_ 也可以整個帶出來到 globals() 方法是：

        for k,v in peforth.execute("_locals_").pop().items():
            globals()[k] = v

    因此配合 DevTools.py 可以很方便地做各種實驗。範例如下：

        __file__ = "DevTools.py"
        import pandas as pd  # 故意看看 df 能不能 SystemExit 之後藉由 peforth 從 function 內部被帶出來到 globals()? 
        def thefunction():
            a = 123
            b = 456
            print(a+b)
            df = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\2021-6-1.feather")
            peforth.bp(11,locals()) # raise SystemExit("Stop right there!")

        # 執行之後被 peforth 中斷，在斷點中除了 peforth 的功能之外還可以 raise SystemExit 把程式停掉。
        thefunction()

        # raise SystemExit 停掉程式之後可以把當時的 context 帶到 globals() 來查看，甚至用 DevTools.py 來做更複雜的實驗。
        %tb
        for k,v in peforth.execute("_locals_").pop().items():
            globals()[k] = v

\ pandas.f
    %f cd 查出 working directory of "I study Pandas.ipynb" 在 OneDrive > jupyter Notebook/ 把 pandas.f mklink 過去:
    mklink /H "d:\GitHub\peforth\playground\pandas.f"  "d:\OneDrive\Documents\Jupyter Notebooks\pandas.f" 
    
    %f include pandas.f 
    %f include d:\GitHub\peforth\playground\pandas.f 
    %f words \ 果然有 dummy_df
    %f help dummy-df
    %f dummy-df -->
    %f dummy-df
    df = peforth.pop()
    %f help see-dummy-df
    %f see-dummy-df

\ 繼續研究 hubble2 component 

    # read df0 from .feather 
    df0 = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\hubble2-2022-10-4.feather")
    df0.info()
    df0.describe()
    df0.customersystem.unique() # array(['LilyLBG', 'UTS', 'Annie', 'LilyTBG', 'Astro', 'ROSA']
    df0[df0.customersystem == 'LilyLBG'].shape
    df0[df0.customersystem == 'Annie'].shape
    df0[df0.customersystem == 'Astro'].shape
    df0[df0.customersystem == 'LilyLBG'].shape
    df0[df0.customersystem == 'LilyTBG'].shape
    df0[df0.customersystem == 'ROSA'].shape
    df0[df0.customersystem == 'UTS'].shape
    df0[['customersystem', 'id']].groupby(['customersystem']).count()  # 查看各 OEM 條數
    df_uts = df0[df0.customersystem == 'UTS']
    df_uts.info()

    # 成功從 UTS 1030 條中分析出 sw column 所用的 mito column formula 
    # https://docs.trymito.io/how-to/interacting-with-your-data 
        df_uts = df0[df0.customersystem == 'UTS']
        or(
            FIND(lower(affected_component+component+general_componentType),'sw'),
            FIND(lower(affected_component+component+general_componentType),'bios'),
            FIND(lower(affected_component+component+general_componentType),'application'),
            FIND(lower(affected_component+component+general_componentType),'software'),
            FIND(lower(affected_component+component+general_componentType),'driver'),
            FIND(lower(affected_component+component+general_componentType),'fw'),
            FIND(lower(affected_component+component+general_componentType),'f/w'),
            FIND(lower(affected_component+component+general_componentType),'firmware'),
            FIND(lower(affected_component+component+general_componentType),'whql'),
            FIND(lower(affected_component+component+general_componentType),'win10'),
            FIND(lower(affected_component+component+general_componentType),'win11'),
        )

    # I am thinking the component heirachy is [rd sw fw bios] or [rd hw ce] or [rd hw kc vga] or [pm] or [rd me] 
    # 先分析出 sw 那麼 rd 也確定了
    
    # 只分出 sw 就很好用了，量夠可以試驗 project 的品質指標。先比較跨 project 的出 bug 曲線。改看最大量的 ROSA.
    
        df_rosa = df0[df0.customersystem == 'ROSA']
        
        # 結果發現篩出 sw column 的 column formula 同 UTS! 而且篩出來的結果，就 sw column 而言，與 general_componentType 一致。
        # 可見 hubble2 team 的分類對 ROSA 而言很可靠。
    
    # 好，利用 ROSA 來訓練 DataRobot 分辨 sw 出來 --> 錯，以上 (單分出 SW) 既然這麼簡單，何不全部六個 OEM 都用 mito 做好來？
        df_uts = df0[df0.customersystem == 'UTS']
        df0[df0.customersystem == 'LilyLBG'].shape
        df0[df0.customersystem == 'Annie'].shape
        df0[df0.customersystem == 'Astro'].shape
        df0[df0.customersystem == 'LilyLBG'].shape
        df0[df0.customersystem == 'LilyTBG'].shape
        df0[df0.customersystem == 'ROSA'].shape
        df0[df0.customersystem == 'UTS'].shape
    
    # mito export df0 沒完沒了，改用 python 好了
        %%time
        df0 = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\hubble2-2022-9-12.feather")
        df0['comp'] = df0.affected_component + ' ' + df0.component + ' ' + df0.general_componentType
        df0['sw'] = [bool(sum([token in row_r.comp.lower() for token in {'sw', 'bios', 'application', 
                    'software', 'driver', 'fw', 'f/w', 'firmware', 'whql', 'win10', 'win11'}])) for i,row_r in df0.iterrows()]
        # Wall time: 11.8 s
        df0.to_csv(r"c:\Users\8304018\Downloads\df0_sw.csv") # 高達 965M, zip 過還有 121M!!
        
        # [X] 送這個 .zip 上 DataRobot 可以嗎？ 不行！ Dataset with size 921 MB exceeds the limit of 200 MB. 
            Task ID=c2186b98-e706-4d73-88e3-94fd4ae564e2 Project ID=63202ce8a8ada389bdb5e5e6
    [X] 所以要分 OEM 先 sample sw T/F 各 100 共 200, 200 * 6 共 1k2 一個 epoch 來多做幾次看看結果
    
        subject = [ row.bug_subject if row.bug_subject else row.shortdescription
                     for row in df0.itertuples() ]
        df0['subject'] = subject
        
        selected_columns = ['id' ,'bug_id' ,'subject', 'Text' ,'component' ,'general_componentType' 
            ,'affected_component' ,'originator' ,'owner' ,'creator_employee_id' ,'owner_employee_id']
        sw_set = {'sw', 'bios', 'application', 'software', 'driver', 'fw', 'f/w', 'firmware', 'whql', 'win10', 'win11'}
            # Examined from feature general_component 仔細核對過，確定沒有漏的，而且跨 OEM 就這一套有效。
        dfs = []
        for oem in df0.customersystem.unique():
            %f oem -->
            df = df0[df0.customersystem == oem][selected_columns].copy()
            df['comp'] = df.affected_component + ' ' + df.component + ' ' + df.general_componentType
            df['sw'] = [bool(sum([token in row_r.comp.lower() for token in sw_set])) for i,row_r in df.iterrows()]
            df.drop('comp', axis='columns', inplace=True)
            print(df.groupby(['sw'])['sw'].count()) # 查看 sw column 的 true,false 各幾個，與 mito 比較 --- 一致！
            print()
            dfs.append(df[df.sw ==  True].sample(100))
            dfs.append(df[df.sw == False].sample(100))
        df = pd.concat(dfs)
        df.info()
        df.to_csv(r"c:\Users\8304018\Downloads\hubble2_sample_1200_sw_eval.csv", index=False)

    [X] 同上，加上 tcount 的考量，品質不好的即非 sw 從 1200 中丟棄。
        import numpy as np
        import pandas as pd
        from collections import OrderedDict
        import Orange
        from Orange.data import Table, Domain, ContinuousVariable, DiscreteVariable, StringVariable
        from orangecontrib.text import Corpus
        from orangecontrib.text.vectorization.bagofwords import BowVectorizer
        import setup
    
        def pandas_to_orange(df):
            domain, attributes, metas = construct_domain(df)
            orange_table = Orange.data.Table.from_numpy(domain = domain, X = df[attributes].values, Y = None, metas = df[metas].values, W = None)
            return orange_table

        def construct_domain(df):
            columns = OrderedDict(df.dtypes)
            attributes = OrderedDict()
            metas = OrderedDict()
            for name, dtype in columns.items():

                if issubclass(dtype.type, np.number):
                    if len(df[name].unique()) >= 13 or issubclass(dtype.type, np.inexact) or (df[name].max() > len(df[name].unique())):
                        attributes[name] = Orange.data.ContinuousVariable(name)
                    else:
                        df[name] = df[name].astype(str)
                        attributes[name] = Orange.data.DiscreteVariable(name, values = sorted(df[name].unique().tolist()))
                else:
                    metas[name] = Orange.data.StringVariable(name)

            domain = Orange.data.Domain(attributes = attributes.values(), metas = metas.values())
            return domain, list(attributes.keys()), list(metas.keys())

    # add 'tcount' column to df --- top
    from collections import OrderedDict
    import Orange
    from Orange.data import Table, Domain, ContinuousVariable, DiscreteVariable, StringVariable
    from orangecontrib.text import Corpus
    from orangecontrib.text.vectorization.bagofwords import BowVectorizer
    
    def add_tcount(df):

        # convert pd to Orange table 
        def pandas_to_orange(df):
            domain, attributes, metas = construct_domain(df)
            orange_table = Orange.data.Table.from_numpy(domain = domain, X = df[attributes].values, Y = None, metas = df[metas].values, W = None)
            return orange_table

        def construct_domain(df):
            columns = OrderedDict(df.dtypes)
            attributes = OrderedDict()
            metas = OrderedDict()
            for name, dtype in columns.items():

                if issubclass(dtype.type, np.number):
                    if len(df[name].unique()) >= 13 or issubclass(dtype.type, np.inexact) or (df[name].max() > len(df[name].unique())):
                        attributes[name] = Orange.data.ContinuousVariable(name)
                    else:
                        df[name] = df[name].astype(str)
                        attributes[name] = Orange.data.DiscreteVariable(name, values = sorted(df[name].unique().tolist()))
                else:
                    metas[name] = Orange.data.StringVariable(name)

            domain = Orange.data.Domain(attributes = attributes.values(), metas = metas.values())
            return domain, list(attributes.keys()), list(metas.keys())

        # add tcount to df
        data = pandas_to_orange(df[['bug_id','Text','subject_weighted']])  # DataFrame 轉成 data table 
        # Data table 轉成 Corpus -- 也是一種 data table 
        corpus = Corpus.from_table(data.domain,data)
        corpus.set_text_features([StringVariable('Text'), StringVariable('subject_weighted')]) # 指定 text columns 否則他亂設的不對。
        bow = BowVectorizer()
        bow = BowVectorizer(norm=bow.L2, wlocal=bow.COUNT, wglobal=bow.SMOOTH)
        vectors = bow.transform(corpus)
        vectors = Table.from_table(Domain(vectors.domain.attributes[1:], None),vectors) 
            # 濾掉第一個 column which is named '?' that means NaN. 萬一沒有 '?' 就是 '000a' 被濾掉，不管它。
        X = vectors.X   # vectors.X_df if you want to view it
        tcount = [x.getnnz() for x in X] # vectors.X 橫向的 token 個數，越大字越多可能品質越好，太小鐵定不好。
        df['tcount'] = tcount # 增加一欄 
        # end of add_tcount(df)
        # it works !     
    # add 'tcount' column to df --- bottom

    # add 'sw' column to df --- top
    def add_sw(df):
        sw_set = {'sw', 'bios', 'application', 'software', 'driver', 'fw', 'f/w', 'firmware', 'whql', 'win10', 'win11'}
            # Examined from feature general_component 仔細核對過，確定沒有漏的，而且跨 OEM 就這一套有效。
        df['comp'] = df.affected_component + ' ' + df.component + ' ' + df.general_componentType
        df['sw'] = [bool(sum([token in row_r.comp.lower() for token in sw_set])) for i,row_r in df.iterrows()] 
            # about 0.8% incorrect according to DataRobot. Accept it at the moment, correct it someday. 
        df.drop('comp', axis='columns', inplace=True)    
    # add 'sw' column to df --- bottom 
    
    # --- main ---- 
    df0 = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\hubble2-2022-9-30.feather")
    subject = [ row.bug_subject if row.bug_subject else row.shortdescription
                     for row in df0.itertuples() ]
        df0['subject'] = subject
        
        selected_columns = ['id' ,'bug_id' ,'subject', 'Text' ,'component' ,'general_componentType' 
            ,'affected_component' ,'originator' ,'owner' ,'creator_employee_id' ,'owner_employee_id', 'subject_weighted', 'customersystem']
        sw_set = {'sw', 'bios', 'application', 'software', 'driver', 'fw', 'f/w', 'firmware', 'whql', 'win10', 'win11'}
        dfs = []
        for oem in df0.customersystem.unique():
            %f oem -->

            # add 'sw' column
            df = df0[df0.customersystem == oem][selected_columns].copy()
            df['comp'] = df.affected_component + ' ' + df.component + ' ' + df.general_componentType
            df['sw'] = [bool(sum([token in row_r.comp.lower() for token in sw_set])) for i,row_r in df.iterrows()]
            df.drop('comp', axis='columns', inplace=True)
            
            print(df.groupby(['sw'])['sw'].count()) # 查看 sw column 的 true,false 各幾個，與 mito 比較 --- 一致！
            print()
            dfs.append(df[df.sw ==  True].sample(100))
            dfs.append(df[df.sw == False].sample(100))
        df = pd.concat(dfs)
        
        # add tcount column and remove 
        add_tcount(df)
        df.index = list(range(df.shape[0])) # [i for i in range(df.shape[0])]
        lowq = [i for i,row in df.iterrows() if row.tcount < setup.tcount_threshold[row.customersystem]]
        df.drop(lowq, inplace=True)
        df.info()
        df.to_csv(r"c:\Users\8304018\Downloads\hubble2_sample_1200_sw_eval.csv", index=False)

    [ ] 從 DataRobot 取得 predict 的結果
        from mitosheet import *; register_analysis("id-jwczndrnpl");
        # Read in filepaths as dataframes
        df = pd.read_csv(f'{downloads}\\' +
             hubble2_sample_1200_sw.csv_eXtreme_Gradient_Boosted_Trees_Classifier_(48)_100_632175a131184fc17b684b46_ab156f5c-7c36-49b2-b5b8-4_hubble2_sample_1200_sw_eval.csv', 
             encoding='utf-8')

        # Set formula of diffs
        df['diffs'] = OR(AND(df['sw']==True,df['PredictedLabel']==0),AND(df['sw']==False,df['PredictedLabel']==1))
        df[df.diffs == True].to_csv(r"c:\Users\8304018\Downloads\diffs.csv", index=False)

\ Hubble2NN python design pattern of "Chain of responsibility" 

    # get input from input.json file 
    with open(r"%s\hubble2\input.json" % downloads) as f:
        ui = json.load(f)

    # Leverage Python design patterns
    # https://hyp.is/gdouNimzEe2s0tPjaqOlew/www.toptal.com/python/python-design-patterns
    class ContentFilter:

        def __init__(self, filters=None):
            self._filters = list()
            if filters is not None:
                self._filters += filters

        def filter(self, content):
            for filter in self._filters:
                content = filter(content)
            return content

    # 成功了
    class Chipset_filter():
        def __init__(self, df0):
            self.df0 = df0 # cooked 三年來的 TR 總表 DataFrame 約 9W 條 TR

        def filter(self, dfi):
            if ui['auto_chipset']: # input.py 中設定的開關
                # 套用 target TR 的 chipset 從 df0 而非 dfi 裡去找保險找得到
                ui['target_chipsets'] = set(self.df0.query('bug_id == "%s"' % ui['bug_id']).chipset.values[0]) 
            logging.info("target_chipsets %s" % ui['target_chipsets']); 
            print("target_chipsets %s" % ui['target_chipsets'])
            
            # 掃描 dfi 找出 chipsets 與 target_chipsets 有交集的 TR
            df_chipset_index = [] # 符合 chipset 條件的 TR index list 
            dfo = dfi
            if ui['target_chipsets']:
                for row in dfi.itertuples():
                    # 以下針對 row 這條 TR 檢查，看它的 chipset 屬不屬於 target chipsets 
                    tr_pjchips = set(row.chipset) if len(row.chipset) else set()
                    if ui['target_chipsets'] & tr_pjchips: 
                        df_chipset_index.append(row.Index)
                logging.info("符合 chipset 條件的 TR df_chipset_index count %d" % len(df_chipset_index)); 
                print("符合 chipset 條件的 TR df_chipset_index count %d" % len(df_chipset_index))

                # get wildcards index list, 當 TR 的 chipset 沒有指定時，與其「漏掉」不如給 user 選擇要不要包含進來。
                wildcards = [i for i,row in dfi.iterrows() if not len(row.chipset)] if ui['include_chipset_unknown'] else []
                dfo = dfi.loc[sorted(list({*wildcards,*df_chipset_index}))] # 合併 chipset screen 與 wildcards 的方法
                logging.info("Optionally include wildcards df.shape %s" % str(dfo.shape)); 
                print("Optionally include wildcards df.shape %s" % str(dfo.shape))
            return dfo

    '''        
        self test 
        df0.shape
        chipset_filter = Chipset_filter(df0)    
        chipset_filter.filter(df0).shape # 成功了
    '''
    
    class Keyword_filter():
        
        def filter(self, dfi):
            def keyword_screen(df, ands):
                bug_idxs = [] # row index list
                count = 0 
                for row in df.itertuples():
                    text = "%s %s %s" % (row.bug_subject.lower(), row.shortdescription.lower(), row.Text)
                    bow = set(text.split(' ')) # 這條 TR 的 bow 整理成集合
                    count_down = len(ands) # 有多少 ands 要滿足，到 0 表示全部滿足。
                    for ors in ands: # 祭出 synonym-expaded keywords
                        ors_flag = False # 先假設 ors 不成立
                        for kw in ors:
                            kw = kw.lower()
                            if (kw in bow) or (text.find(kw) != -1): # BoW 不成立才老實做 linear search 
                                count_down -= 1
                                ors_flag = True
                                break # bow 裡只要有一個 keyword 這整條 ors 就成立了，跳下ㄧ個 ands 
                        if not ors_flag :
                            break # 只要有一條 ors 不成立，後面都不用看了。跳下一條 TR。
                    if not count_down: # 到 0 表示成立了， row 為所求
                        bug_idxs.append(row.Index)
                        count += 1    
                return bug_idxs
            
            # Screen by synonyms expanded keywords
            logging.info("Screen by keywords with synonym expansion"); 
            print("Screen by keywords with synonym expansion")
            # 撈出 bug index list 
            idxs = keyword_screen(dfi, ui['ands'])
            df_kw = dfi.loc[idxs]
            logging.info("df_kw.shape %s" % str(df_kw.shape)); 
            print("df_kw.shape %s" % str(df_kw.shape))
            return df_kw
    '''
    keyword_filter = Keyword_filter()
    dfi = chipset_filter.filter(df0)
    dfi.shape
    df_kw = keyword_filter.filter(dfi)

    '''

    hubble2_filter = ContentFilter([
        Chipset_filter(df0).filter,
        Keyword_filter().filter
        ])

    df = hubble2_filter.filter(df0)
    df.shape

\ bug_id 撞號 1. 直接取最新的。df0 本來就是用 last_updated 排序的，因此用 bug_id loc[] 出來取 [0] 就對了！
    
    先掌握三年來的撞號表 
        # 輕易查出 bug_id top 208880 freq=2 (87701-87693)=8 有 8組 重複的，都是 2 個 TR 
        mitosheet.sheet(df0[['id','bug_id','subject_weighted']])

        # 進一步列出所有重複的 rows 
        # https://www.geeksforgeeks.org/find-duplicate-rows-in-a-dataframe-based-on-all-or-selected-columns/
        dups = df0[df0.duplicated('bug_id')]
        df0.query('bug_id in @dups.bug_id')

        df0[df0.bug_id == "213110"].last_updated #  
        df0[df0.bug_id == "205928"].last_updated # 
        df0[df0.bug_id == "212005"].last_updated #  都是新的在上面無誤
        df0[df0.bug_id == "212007"].last_updated #  都是新的在上面無誤
        df0[df0.bug_id == "208881"].last_updated #  都是新的在上面無誤
        df0[df0.bug_id == "208880"].last_updated #  都是新的在上面無誤
        df0[df0.bug_id == "201905"].last_updated # 
        df0[df0.bug_id == "206492"].last_updated # 

\ 改用 global target_tr object 

    class Target_TR:
        def __init__(self, bug_id):
            self.bug_id = str(bug_id)
            self.tr = df0[df0.bug_id == self.bug_id].iloc[0]

        # bug_id = '206492'
        # df0[df0.bug_id == bug_id][['customersystem','last_updated']]
        # target_tr = Target_TR(bug_id)
        # target_tr.bug_id
        # type(target_tr.tr)
        # target_tr.tr.id
        # target_tr.tr.last_updated

    target_tr.tr.chipset  
    target_tr.tr.subject_weighted
    set(target_tr.tr.subject_weighted.split(' '))
    set(target_tr.tr.subject_weighted.split(' '))

\ yb matches keyboard <--- fix this problem  --> done 
\ Hubble2 project quality study --> Bug curve by Matplotlib ok 
\ Hubble2 bug curve 改用 plotly 呈現
\ 整頓 hubble2NN runtime ... done
\ working on tcount AI model 

    # tcount 篩檢: 2022/9/30 10:54 用現成的 tcount threshold 先把 6 大 system 打散標好應用 tcount 按 customer threshold 標好的
    # quality label 訓練一個 model 用來分辨 6 大 systems 的樣本，與 tcount label 解答比對。 要故意多選 tcount 不及格的 TR 來參加訓練。
    
    quality = [0 if bug.tcount < setup.tcount_threshold[bug.customersystem] else 1 for bug in df0.itertuples()]
    df0['quality'] = quality
    
    %f quality count --> nip
    %f df0 :> shape -->
    %f quality py> sum(pop()) -->
    %f quality count nip quality py> sum(pop()) - --> # 低品質的有 7K 條，夠用來做訓練了。

    
    # 查看各 customer 的 tcount quality 好壞各有多少，以便決定抽樣條數。當前因為 UTS 只有 15 條 quality==0 的，故都只抽 15.
    # 15 * 2 * 6 = 180 總共。
    for customer in sw_df.customersystem.unique():
        customer_df = sw_df[sw_df.customersystem == customer]
        # ic(customer_df.quality.unique())
        ic(customer, customer_df.groupby(['quality'])['quality'].count());

    # 取樣好壞各 15
    dfs = []
    for customer in sw_df.customersystem.unique():
        for q in [0,1]:
            dfs.append(sw_df.query('customersystem == @customer & quality == @q').sample(n=15))
    training_df = pd.concat(dfs)
    training_df[['id' ,'sw' ,'bug_id' ,'subject_tokens' ,'customersystem' ,'Text' ,'tcount' ,'quality']]
    
\ 改寫 DataGrab 的 scan_begin_end(t0,t1) 方便抓本來沒抓到的 fields (如下 columns list 者). Kernel 直接用 DataGrab。 

    def scan_begin_end(begin, end, columns):
        client = Elasticsearch(setup.kibana_url)
        s = Search(using=client, index="hubble2") \
            .query("exists", field="bug_id") \
            .filter("range",last_updated = {"gte": begin,"lt": end}) \
            .extra(track_total_hits=True) \
            .source(columns)

        df = pd.DataFrame()
        for hit in s.scan():
            d = { k:hit[k] for k in hit if k!="meta"}
            d["id"] = hit.meta['id'] # 唯一不重複的 id 在 Kibana 上的 meta columns 內 
            df = df.append(d, ignore_index=True)
        fname = "status %s-%s-%s.feather" % (begin.year, begin.month, begin.day)  # 用開始日期當作檔名存檔 
        df.reset_index().to_feather(setup.path + fname )  # feather 有時候會要求先 reset_index
        return df

    %%time 
    t1 = setup.now
    t0 = datetime.datetime(t1.year,t1.month,1)
    month_count = setup.max_month
    columns = [
        'issue_status','status'
    ]

    while True:
        fname = "status %s-%s-%s.feather" % (t0.year, t0.month, 1)  # 該月一日當作檔名
        if os.path.exists(setup.path + fname) : # skip if exist
            logging.debug(" %s exists" % fname)
        else:
            scan_df = scan_begin_end(t0,t1, columns) # 留下 scan_df 方便 debug
            logging.info("%s ~ %s %s downloaded" % (t0, t1, fname))
        if month_count and t0 > setup.start_date:
            t1 = t0
            t0 = t1 - datetime.timedelta(days=1)
            t0 = datetime.datetime(t0.year,t0.month,1)
            month_count -= 1
        else:
            break
    # Wall time: 5.98 ms
    
    # Read back monthly data from .feathere files
    t1 = setup.now
    t0 = datetime.datetime(t1.year,t1.month,1)
    dfs = [] # 逐月分開放的 DataFrames. 注意順序是從本月往前推
    df = pd.DataFrame()
    month_count = setup.max_month
    while True:
        fname = "status %s-%s-%s.feather" % (t0.year, t0.month, 1)  # 該月一日當作檔名
        logging.debug("%s ~ %s - %s" % (t0,t1,fname))
        df = pd.read_feather(setup.path + fname)
        dfs.append(df)
        if month_count and t0 > setup.start_date:
            t1 = t0
            t0 = t1 - datetime.timedelta(days=1)
            t0 = datetime.datetime(t0.year,t0.month,1)
            month_count -= 1    
        else:
            break
    del df # fire an error, df should not be used here after
    df = pd.concat(dfs)
    df.shape
    df.info()

    
    # 要取 status 就用 issue_status 一個 column 就很完整, 另一個 status column 幾乎與 issue_status 一樣，只有 LilyTBG 有部分是 NaN. 
    df['a'] = df.issue_status == df.status  # 撈出兩者不同的
    df['b'] = df.id.str.slice(0,7)          # 查看，發現只有小部分 LilyTBG 且其值為 NaN
    mitosheet.sheet(df)

    # 以上取得的 issue_status 併回 df0 產生新檔 df3 
    temp_df = df.drop_duplicates(subset=['id']) # Remove duplicates so lookup merge only returns first match
    df_tmp = temp_df.drop(['eastatus', 'general_status', 'status', 'index', 'executivestatus', 'statusupdate',
                           'vendorissuestatus', 'a', 'b', 'processstatus'], axis=1)
    df3 = df0.merge(df_tmp, left_on=['id'], right_on=['id'], how='left', suffixes=['_df0', '_df'])

    # df3 存檔
    df3.to_feather(setup.path + "df3 with issue_status.feather") 
    df3 = pd.read_feather(setup.path + "df3 with issue_status.feather")  
    
\  研究 new draft 的 quality. ROSA, LilyLBG, LilyTBG, Annie 都有 new,draft 但 Astro 直接就是 open 
\  Indeed! new and draft TRs' 真的，當然，沒有甚麼內容。至於 tcount 已經打算畫圖 histogram 疊加各種 status 出來看了。
    df3['new'] = 0
    for index in df3[(df3.issue_status == "Draft") | (df3.issue_status == "New")].index:
        df3.at[index,'new'] = 1

    先只挑 sw 的，否則會有 field 的雜的。
    經查, new,draft 的共有 225 其中 tcount quality 0 的 54; 1 的 711 基本上有關聯。看看這些品質 1 的。。。
    
    334895 334887 334868 334910 334847 334853 213037 213499 213110 212177 212176 212794 PIMS-166312 213033 PIMS-165585 213225 213222 213109 DFT192954 331494 331482 331477 
    331476 331140 PIMS-155552 210900 PIMS-153717 PIMS-150638 PIMS-150637 209760 PIMS-149297 PIMS-149299 PIMS-148079 207458 207879 207882 207880 321447 207696 207466 PIMS-136244 
    PIMS-139519 PIMS-138966 PIMS-138523 PIMS-138914 PIMS-137086 PIMS-138536 319281 318976 318979 PIMS-135821 PIMS-134929 PIMS-134020 PIMS-132577 PIMS-129630 PIMS-129629 
    PIMS-129626 PIMS-129582 309286 PIMS-128109 PIMS-128091 PIMS-128191 PIMS-128190 PIMS-128188 PIMS-128187 PIMS-127149 PIMS-114846 PIMS-113807 295699 PIMS-111619 PIMS-110272 
    292007 PIMS-99588 PIMS-98307 PIMS-97676 286555 PIMS-90194 PIMS-90777 PIMS-91566 PIMS-89644 PIMS-89553 PIMS-87693 PIMS-87332 PIMS-86676 PIMS-86224 PIMS-85460 PIMS-84661 
    PIMS-83920 PIMS-82432 PIMS-82433 PIMS-82706 PIMS-82793 PIMS-80111 DFT161993 282009 PIMS-80780 PIMS-80269 196268 DFT160974 PIMS-76829 PIMS-76768 PIMS-73043 PIMS-72543 
    PIMS-69838 278239 278210 PIMS-69920 PIMS-67705 PIMS-66556 PIMS-66557 PIMS-65726 DFT157597 DFT157304 PIMS-57389 PIMS-52213 PIMS-52097 192227 PIMS-47208 PIMS-44014 PIMS-27989 
    256860 PIMS-13239 PIMS-27975 PIMS-34030 PIMS-31223 253731 PIMS-30471 PIMS-29782 PIMS-28047 PIMS-24127 PIMS-20836 248138 247549 PIMS-17555 PIMS-15627 PIMS-15606 183684 
    PIMS-7589 PIMS-1393 PIMS-5934 PIMS-5512 PIMS-5396 ECR565052 ECR565038 ECR565034 184496 233300 182362 182361 182360 182359 182358 182357 182356 DFT000959 DFT000857 181312 
    181545 181488 180659 181313 ECR545331 ECR545326 ECR545324 ECR545321 ECR545308 ECR545307 ECR545274 ECR545244 ECR545241 218471

\ 等級分 of tcount study already done! Now tcount threshold for each customer are automated. 

    https://we147121.pixnet.net/blog/post/291747236-%E4%B8%89%E5%88%86%E9%90%98%E6%90%9E%E6%87%82%E7%99%BE%E5%88%86%E7%AD%89%E7%B4%9A%E8%88%87%E7%99%BE%E5%88%86%E4%BD%8D%E6%95%B8#:~:text=1.-,%E7%94%B1%E5%8E%9F%E5%A7%8B%E5%88%86%E6%95%B8%E8%A8%88%E7%AE%97%E7%99%BE%E5%88%86%E7%AD%89%E7%B4%9A,-%E9%A6%96%E5%85%88%E5%B0%87%E5%8E%9F%E5%A7%8B


    score = [10,20,30,40,50,55,54,80,90,*([100]*91)]
    score = [35,34,30,37,86,88,83,84,96,96]
    students = range(1,11)
    students = range(1,101)
    df = pd.DataFrame(data={'student': students,
                           'score': score})
    df
    df['pct_rank'] = df['score'].rank(pct=True)
    df['default_rank'] = df['score'].rank(method="dense")
    df['max_rank'] = df['score'].rank(method='max')
    df['NA_bottom'] = df['score'].rank(na_option='bottom')
    df['pct_rank'] = df['score'].rank(pct=True,method="dense")

    35,56,62,71,76,78,83,84,90,96

    pr = (100/10*1)+(100/10*1/2)
    pr

    a = df.score.describe()
    %f a type -->
    %f a :> ['75%'] -->
    %f a :> ['max'] -->
    我們要看的不是 PR rank 等級分，而是分佈。df.describe() 所呈現者。先從 df.score.describe()['max'] 比較起， 然後依序
        df.score.describe()['max']
        df.score.describe()['75%'] # percentile 分位數，取得分位數值用 df.score.quantile(0.07)
        df.score.describe()['50%']
        df.score.describe()['25%']
        df.score.describe()['min']
    比等級大的就是屬於該等級，比 min 小的也會有，那就是破下限的；反之也有破上限的，我們不用太強調。

    # 查 tcount threshold 分屬該 customer 的分位數

    for customer in df3.customersystem.unique():
        print("\n"+customer)
        print(df3[(df3.customersystem == customer) & (df3.sw == True)].tcount.describe(percentiles=[.05, .1, .15, .2, .25, .5, .75]))

    df3.shape
    df3.info()
    for customer in df3.customersystem.unique():
        print("%s %.0f" % (customer, df3[(df3.customersystem == customer) & (df3.sw == True)].tcount.quantile(0.05)))

    %%time    
    tcount_threshold = {customer:int(df3[(df3.customersystem == customer) & (df3.sw == True)].tcount.quantile(0.05)) for customer in df3.customersystem.unique()}

\ plot tcount histogram of variant status. 
\ tcount is like entropy that grows one way so the X axis is tcount and the Y axis is count of TRs and the distribution will tell us the 
\ 先後關係 of any, probably unknown, statuses. 

    df = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\df3 with issue_status.feather")
    figsize = (24,10)
    n_bins  = 50
    customers = df.customersystem.unique()

    class Tcount_histogram:
        def __init__(self, df):
            self.df = df  # like df0 dataset 
            self.i = 0    # iteration count of each customer

        def plot(self):
            if self.i >= len(customers): 
                print("all over")
                return
            customer = customers[self.i]
            customer_df = self.df[self.df.customersystem == customer]
            fig, ax = plt.subplots(figsize=figsize)
            for status in customer_df.issue_status.unique():
                x = customer_df[(customer_df.issue_status == status)].tcount
                ax.hist(x, n_bins, density=False, label=status, log=True)

            ax.grid(True)
            ax.legend(loc='upper right')
            ax.set_title('tcount distribution - %s' % customer, weight='bold', size=32)
            ax.set_xlabel('tcount', labelpad=20, weight='bold', size=12)
            ax.set_ylabel('TR count', labelpad=20, weight='bold', size=12);
            self.i += 1
            

    tcount_histogram = Tcount_histogram(df)
    tcount_histogram.plot()

\ 搞懂 binary search for buckets         
    import bisect

    places = [
        (501, 'ASIA'),
        (1262, 'EUROPE'),
        (3389, 'LATAM'),
        (5409, 'US'),
        (float('inf'), 'US'),
    ]
    places = [
        (float('-inf'), 'ASIA'), # 塞進這個，兩頭完整，就不用管 bisect_right or bisect_left
        (501, 'ASIA'),
        (1262, 'EUROPE'),
        (3389, 'LATAM'),
        (5409, 'US'),
        (float('inf'), 'US'),  # 塞進這個，兩頭完整，就不用管 bisect_right or bisect_left
    ]
    places.sort() # list must be sorted

    for to_find in (389, 1300, 5400, 5500):
        pos = bisect.bisect_right(places, (to_find,))
        print ('%s -> %s' % (to_find, places[pos]))
    for to_find in (389, 1300, 5400, 5500, 900):
        pos = bisect.bisect_left(places, (to_find,))
        print ('%s -> %s' % (to_find, places[pos]))
    for to_find in (389, 1300, 5400, 5500, 900):
        pos = bisect.bisect_right(places, (to_find,))
        print ('%s -> %s' % (to_find, places[pos]))
    for to_find in (389, 1300, 5400, 5500, 900):
        pos = bisect.bisect(places, (to_find,))
        print ('%s -> %s' % (to_find, places[pos]))
    
\ 算出 all SW TR 的 tcount 在該 customer 該 state 當中的分位值 10% 20% ... 90% 落在哪一段？直接對 9W 近三年的 SW TR 全部算好填入。
\ 此後重算當月有更新的就好。有了所有 TR 的分位值就是它的品質了，進一步可以用來評斷 owner, project, component 的品質表現。當大家都用心寫了分佈還會漂移。
\ 分 bucket 的方法 : https://stackoverflow.com/questions/2899129/find-value-within-a-range-in-lookup-table  

    df = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\df3 with issue_status.feather")
    dfs = []
    for customer in df.customersystem.unique():
        customer_df = df[(df.customersystem == customer) & (df.sw == True)]
        for status in customer_df.issue_status.unique():
            status_df = customer_df[(customer_df.issue_status == status)]
            quantiles = status_df.tcount.describe(percentiles=[.1, .2, .3, .4, .5, .6, .7, .8, .9])
            buckets = [(float('-inf'),10)] # 塞進這個，兩頭完整，就不用管 bisect_right or bisect_left
            buckets += [(quantiles.iat[level],(i+1)*10) for i,level in enumerate(range(4,4+9))] # 取 10% 到 90% 的 tcount threshold 值
            buckets += [(float('inf'),90)]  # 塞進這個，兩頭完整，就不用管 bisect_right or bisect_left
            for i,row in status_df.iterrows():
                pos = bisect.bisect(buckets, (row.tcount,))
                status_df.at[i,'tcount_quantile'] = buckets[pos][1]
            dfs.append(status_df)
    cooked_df = pd.concat(dfs)
    # 驗算，條數一致無誤
    cooked_df.shape
    df[(df.sw == True)].shape

    # 按 customer account 列出所有 owner 的 closed TR median 
    s = "customer,owner,median\n"
    for customer in cooked_df.customersystem.unique():
        select_df = (cooked_df
                     .query('customersystem == @customer')
                     .query('issue_status in ["Closed","Complete","Fixed"]'))
        for owner in select_df.owner_employee_id.unique():
            median = select_df.query('owner_employee_id == @owner').tcount_quantile.median()
            s += f'{customer},"{owner}",{median}\n'
    %f s -->

    # 存檔
    %f s s" c:\Users\8304018\Downloads\Hubble2 owner quality.csv" writeTextFile

\ tcount 可能還不如用 tf-idf ? 

    euclidean_distances(vectors.X[:10])  # 兩兩之間的距離，我想知道 vector 的長度 norm 怎麼算？問 Google 吧

    dir(vectors.X[0])
    len(vectors.X[0])
    type(vectors.X[0])
    vectors.X_df.shape

\ Norm (length) of a spares vector https://stackoverflow.com/a/20169063/2179543
    Some simple fake data:

    a = np.arange(9.).reshape(3,3)
    s = sparse.csr_matrix(a)
    To get the norm of each row from the sparse, you can use:

    np.sqrt(s.multiply(s).sum(1))
    
    s = vectors.X[:5]
    np.sqrt(s.multiply(s).sum(1)) # <-------------- 一口氣算出所有 vectors' norm 用來與 tcount 比較     
    
    And the renormalized s would be

    s.multiply(1/np.sqrt(s.multiply(s).sum(1)))
    or to keep it sparse before renormalizing:

    s.multiply(sparse.csr_matrix(1/np.sqrt(s.multiply(s).sum(1))))
    To get ordinary matrix or array from it, use:

    m = s.todense()
    a = s.toarray()
    If you have enough memory for the dense version, you can get the norm of each row with:

    n = np.sqrt(np.einsum('ij,ij->i',a,a))
    or

    n = np.apply_along_axis(np.linalg.norm, 1, a)
    To normalize, you can do

    an = a / n[:, None]
    or, to normalize the original array in place:

    a /= n[:, None]
    The [:, None] thing basically transposes n to be a vertical array.

\ Generator, iterator, 

    # Generator function is returned by yield instruction 
    def my_range(start, end):
        while start <= end:
            yield start
            start += 1
    for i in my_range(1, 10):
        print(i)

    import random
    def rand_stream():
        while True:
            yield random.random()        

    # Iterator 可以在 for loop 之外應用，只跑一輪，多跑就是 StopIteration error 
    l = [1, 2, 3]
    it = iter(l)
    print(next(it))
    print(next(it))
    print(next(it))

\ Level up your Pandas skills with query() and eval()
    \ https://medium.com/@dreamferus/level-up-your-pandas-skills-with-query-and-eval-f065951162df

    # create fake data
    import pandas as pd
    import numpy as np

    data = []
    for _ in range(100):
        data.append({
            "gender": "Male",
            "height": np.random.normal(178, 10),
            "age": np.random.uniform(20, 70)
        })
    for _ in range(100):
        data.append({
            "gender": "Female",
            "height": np.random.normal(166, 8),
            "age": np.random.uniform(20, 70)
        })
    df = (pd.DataFrame(data)
        # sample to mix order
        .sample(frac=1.0, replace=False)
        .reset_index(drop=True)
     )    

    # the old way  df.loc[] 
    (
        df[(df["gender"] == "Female") & (df["age"] >= 20) & (df["age"] <= 30)]["height"]
        .pipe(lambda x: [x.max(), x.min()])
    )
    # [176.40488215460397, 150.4165879152207]

        (
            df[(df["gender"] == "Female") & (df["age"] >= 20) & (df["age"] <= 30)]["height"]
            .pipe(lambda x: peforth.push(x).bp(11,locals()))
        )
        # x 就是整個 df , 公式 df.pipe(lambda x: ... ) , 用 peforth 一下就查出來了，如上。

    # the new way of query() and eval() 
    (
        df.query("gender == 'Female' and 20 <= age <= 30")
        .eval("height.max(), height.min()")
    )
    # array([176.40488215460397, 150.4165879152207], dtype=object) 
    
    
    # The old way again 
    a = df[df["gender"] == "Male"].reset_index(drop=True).assign(age=df.age-10) # 錯誤的例子，the self 已經不是 df 了
    b = (df[df["gender"] == "Male"].reset_index(drop=True).pipe(lambda x: x.assign(age=x.age-10)))  # 必須用 pipe() 才對

    # let's add some filtering afterwards
    b = (df[df["gender"] == "Male"].reset_index(drop=True)
        .pipe(lambda x: x.assign(age=x.age-10))  # column operation
        .pipe(lambda x: x[x["age"] > 50]))       # filtering 
    
    # The new way 
    c = (df.query("gender == 'Male'")
        .reset_index(drop=True)
        .eval("age=age-10")
        .query("age > 50"))
    assert b.equals(c)    # 兩個 dataFrame 是否相等有這個方法！

\ Clean Up Your Python Code with Decorators
    \ https://medium.com/@bubbapora_76246/clean-up-your-python-code-with-decorators-613e7ad4444b

    import pandas as pd

    transaction_ids = [101, 102, 103, 104, 105]
    item_ids = ['shirts', 'socks', 'jeans', 'socks', 'shirts']
    sale_amts = [25, 12, 32, None, 20]

    df = pd.DataFrame({'trans_id': transaction_ids,
                       'item_id': item_ids,
                       'sale_amt': sale_amts})

    import functools

    def tracking_decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            df = kwargs['data']
            original_trans_set = set(df['trans_id'])
            df = func(*args, **kwargs)
            post_filtering_trans_set = set(df['trans_id'])
            filtered_trans = list(original_trans_set - post_filtering_trans_set)
            return df, filtered_trans
        return wrapper

    @tracking_decorator
    def remove_nulls(data):
        return data.dropna()

    remove_nulls(data=df)

\ hubble2 研究 LilyLBG component 分布 on each projects 

    {GitHub}/hubble2-nearest-neighbor/Hubble2 Bug Curve Study.ipynb
    
\ study status fields and list the union of open-working 

    df3 = pd.read_feather(setup.path + "hubble2-2022-10-13.feather")  

    # 研究所有的 status 
    for customer in df3.customersystem.unique():
        %f cr customer -->
        for status in df3.issue_status.unique():
            n = df3.query('issue_status == @status and customersystem == @customer').shape[0]
            if n: 
                ic(status, n)

    # 查證過了
    openworking = ['Analyze',  'Analyzing',  'Assigned',  'New',  'Open',  'Reopened',  
                   'Review',  'Submitted', 'WAIT',  'Verify',  'Wait',  'Working']
    
    df3['ow'] = [bool(sum([status == row.issue_status for status in openworking])) 
                 for row in df3.itertuples()]
    
    df3[['ow', 'issue_status']].groupby(['ow']).count()
    df3[df3.ow == True].issue_status.unique()
    df3[df3.ow == False].issue_status.unique()

\ Pandas tips

    df.equals(df0)  # 直接比不成功，不知道哪裡動過了。   
    df.bug_id.equals(df0.bug_id)  # 分開比成功  

\ To speed up hubble2NN     
    \ Modyfy hubble2NN() to cacheNN(df, trid) 

    done --> GitHub/hubble2-nearest-neighbor/Hubble2NN_with_synonyms_and_chipset.ipynb

\ Suppress FutureWarning or Escalate the warning to Error so as to locate the code to improve it 

    # C:\Users\8304018\AppData\Local\Temp\ipykernel_9792\2444753131.py:16: FutureWarning: 
    #     The frame.append method is deprecated and will be removed from pandas in a future version. Use pandas.concat instead.
    import warnings
    warnings.simplefilter(action='error', category=FutureWarning) # action='ignore' 或者 action='error'

\ df.append(d, ignore_index=True) Future warning 改用 pd.concat 
    
    # df.append(d, ignore_index=True) Future warning 
    row = pd.DataFrame([pd.Series(d)])  # 單一 row 的 DataFrame.
    df = pd.concat([df,row], ignore_index = True)  # 接上總表的最下面

\ Python my_function(*args, **kwargs) 用法 Understand *args And **kwargs in Python
    \ https://python.plainenglish.io/understand-args-and-kwargs-in-python-51d49454f9b8

    def my_function(*args, **kwargs):
        ic(args)
        ic(kwargs)

    my_function(1,2,3, k=11,q=22)    
    # ic| args: (1, 2, 3)
    # ic| kwargs: {'k': 11, 'q': 22}

\ 8 Stunning Python Dictionary Tricks That Make Your Code Elegant
    https://medium.com/techtofreedom/8-stunning-python-dictionaries-tricks-that-make-your-code-elegant-d8b3cb08bd15

    1. Using Union Operators To Merge Dictionaries 等效 unpacking 但可能更快。
        cities_us = {'New York City': 'US', 'Los Angeles': 'US'}
        cities_uk = {'London': 'UK', 'Birmingham': 'UK'}

        cities = cities_us | cities_uk; ic(cities)
        
        cities = {};         ic(cities)
        cities |= cities_us; ic(cities)
        cities |= cities_uk; ic(cities)

    2. Unpacking Dictionaries with Asterisks
    
        cities = {**cities_us, **cities_uk}; ic(cities)
    
    3. Using Dictionary Comprehension To Create Dictionaries 這個我知道
    
        cities = ['London', 'New York', 'Tokyo', 'Cambridge', 'Oxford']
        countries = ['UK', 'US', 'Japan', 'UK', 'UK']
        uk_cities = {city: country for city, country in zip(cities, countries) if country == 'UK'} # zip() 可以這樣用！
        ic(uk_cities)
        # {'London': 'UK', 'Cambridge': 'UK', 'Oxford': 'UK'}
        
    4. Reversing Keys and Values of a Dictionary
    
        cities = {'London': 'UK', 'Tokyo': 'Japan', 'New York': 'US'}
        # Method 1
        reversed_cities = {v: k for k, v in cities.items()}
        ic(reversed_cities)
        # {'UK': 'London', 'Japan': 'Tokyo', 'US': 'New York'}

        # Method 2
        reversed_cities = dict(zip(cities.values(), cities.keys()))
        ic(reversed_cities)

        # Method 3
        reversed_cities = dict(map(reversed, cities.items()))
        ic(reversed_cities)

    5. Converting Lists Into Dictionaries

        cities = [('London', 'UK'), ('New York', 'US'), ('Tokyo', 'Japan')]
        d_cities = dict(cities)
        ic(d_cities)
        # {'London': 'UK', 'New York': 'US', 'Tokyo': 'Japan'}
        
        cities = ['London', 'Leeds', 'Birmingham']
        d_cities = dict.fromkeys(cities,['UK']) # set the default value to 'UK' 
        print(d_cities)
        # {'London': 'UK', 'Leeds': 'UK', 'Birmingham': 'UK'}

    6. Sorting a Dictionary

        cities = {'London': '2', 'Tokyo': '3', 'New York': '1'}
        print(sorted(cities.items(),key=lambda d:d[1]))  # 用 lambda 傳回 item 的 key 
        # [('New York', '1'), ('London', '2'), ('Tokyo', '3')]    
        
    7. Using Defaultdict
    
        city = {'UK':'London','Japan':'Tokyo'}
        print(city['Italy'])
        # KeyError: 'Italy'    
        
        from collections import defaultdict
        city = defaultdict(str)
        city['UK'] = 'London'
        ic(city['Italy'])

    8. Using Counter
    
        from collections import Counter
        author = "Stella"
        chars = Counter(author)
        print(chars)
        # Counter({'l': 2, 'S': 1, 't': 1, 'e': 1, 'a': 1})
        
\ FastAPI Fundamentals — Getting Faster with FastAPI
    https://levelup.gitconnected.com/fastapi-fundamentals-getting-faster-with-fastapi-866545b841ca
    
    \ pip installs 
        (LBB2 base) conda create -n FastAPI python=3.10 
        activate FastAPI
        pip install fastapi[all]  <--- install all dependencies of FastAPI at once, this I didn't know!  
        conda install jupyterlab

    \ main.py
    
        # main.py
        from fastapi import FastAPI
        from pydantic import BaseModel
        from typing import Union

        class User(BaseModel):
            user_id: int
            user_age: Union[int, None] = None
            user_name : str
            user_nick_name = 'Oreo'

        app = FastAPI()

        @app.get('/')
        def base():
            return "Hello World!"

        # Example 1: path parameter without type validation
        @app.get('/student1/{student_id}')
        def base(student_id):
            return f"Student id is {student_id} & it's type is : {type(student_id)}"

        # Example 2: path parameter with type validation
        @app.get('/student2/{student_id}')
        def base(student_id:int): # student_id has to be an int
            return f"Student id is {student_id} & it's type is : {type(student_id)}"

            # http://localhost:8000/student2/4243242342    
            # "Student id is 4243242342 & it's type is : <class 'int'>"

            # http://localhost:8000/student2/4232aaa   指定 type 真的有檢查，故意放錯果然被逮！
            # {"detail":[{"loc":["path","student_id"],"msg":"value is not a valid integer","type":"type_error.integer"}]}

        @app.get('/add')    
        def base(x: int, y: int): # type validation
            return f"x = {x} and y = {y} and x + y = {x+y}"

        @app.get('/multiply')    
        def base(x: int = 4, y: int = 6): # type validation
            return f"x = {x} and y = {y} and x * y = {x*y}"

        @app.post('/login')    
        def base(user: User): # type validation
            return {"msg": "login successful"}

    \ run server on a terminal tab of jupyterlab 
    
        (FastAPI) D:\OneDrive\Documents\Jupyter Notebooks\FastAPI>uvicorn main:app         
        INFO:     Started server process [9364]
        INFO:     Waiting for application startup.
        INFO:     Application startup complete.
        INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)
        INFO:     127.0.0.1:55712 - "GET / HTTP/1.1" 200 OK
        INFO:     127.0.0.1:55712 - "GET /favicon.ico HTTP/1.1" 404 Not Found
    
        open web page http://localhost:8000/ and it shows "Hello World!" successfully 
    
    \ http://localhost:8000/docs     
    
        真的就有 student1 student2 等 endpoint/rout/pai 的 playground on the web page !!

\ hubble2NN API server building up 
    
    1. make main.py from "d:\GitHub\hubble2-nearest-neighbor\Hubble2NN_with_synonyms_and_chipset.ipynb"
    2. modify main.py to add above things 
    3. at hubbl2NN working directory run:  
        uvicorn main:app 

\ add tcount and also add norm 
    for i,df in enumerate(dfs):
        break

    def add_tcount(df):

        # convert pd to Orange table 
        def pandas_to_orange(df):
            domain, attributes, metas = construct_domain(df)
            orange_table = Orange.data.Table.from_numpy(
                domain = domain, 
                X = df[attributes].values, 
                Y = None, 
                metas = df[metas].values, 
                W = None
            )
            return orange_table

        def construct_domain(df):
            columns = OrderedDict(df.dtypes)
            attributes = OrderedDict()
            metas = OrderedDict()
            for name, dtype in columns.items():

                if issubclass(dtype.type, np.number):
                    if len(df[name].unique()) >= 13 or \
                        issubclass(dtype.type, np.inexact) \
                        or (df[name].max() > len(df[name].unique())):
                        attributes[name] = Orange.data.ContinuousVariable(name)
                    else:
                        df[name] = df[name].astype(str)
                        attributes[name] = Orange.data.DiscreteVariable(
                            name, 
                            values = sorted(df[name].unique().tolist())
                        )
                else:
                    metas[name] = Orange.data.StringVariable(name)

            domain = Orange.data.Domain(attributes = attributes.values(), metas = metas.values())
            return domain, list(attributes.keys()), list(metas.keys())

        # add tcount to df
        data = pandas_to_orange(df[['bug_id','Text','subject_tokens']])  # DataFrame 轉成 data table 
        # Data table 轉成 Corpus -- 也是一種 data table 
        corpus = Corpus.from_table(data.domain,data)
        corpus.set_text_features([StringVariable('Text'), StringVariable('subject_tokens')]) # 指定 text columns 否則他亂設的不對。

        # add tcount 
        bow = BowVectorizer() # 取得 bow 接著要用到 bow.L2, bow.COUNT 等等
        bow = BowVectorizer(norm=bow.NONE, wlocal=bow.COUNT, wglobal=bow.SMOOTH)
        vectors = bow.transform(corpus)
        vectors = Table.from_table(Domain(vectors.domain.attributes[1:], None),vectors) 
            # 濾掉第一個 column which is named '?' that means NaN. 萬一沒有 '?' 就是 '000a' 被濾掉，不管它。
        X = vectors.X   # vectors.X_df if you want to view it
        tcount = [x.getnnz() for x in X] # vectors.X 橫向的 token 個數，越大字越多可能品質越好，太小鐵定不好。
        norms = np.sqrt(X.multiply(X).sum(1)) # 一口氣算出所有 vectors' norm 

        df['tcountNONE'] = tcount # 增加一欄 
        df['tcountL2'] = tcount # 增加一欄 

        # norm 
        s = vectors.X[:10]
        np.sqrt(s.multiply(s).sum(1)) # <-------------- 一口氣算出所有 vectors' norm 用來與 tcount 比較     
            matrix([[194.88243086],
                    [418.61212044],
                    [ 37.22287324],
                    [103.17103212],
                    [220.57036555],
                    [130.66777038],
                    [112.818593  ],
                    [347.30023549],
                    [123.66946816],
                    [159.11735175]])        
                    df['tcountL2'].head(10)
                    # end of add_tcount(df)

\ Class DictObj()
    https://joelmccune.com/python-dictionary-as-object/

    class DictObj:
        def __init__(self, in_dict:dict):
            assert isinstance(in_dict, dict)
            for key, val in in_dict.items():
                if isinstance(val, (list, tuple)):
                   setattr(self, key, [DictObj(x) if isinstance(x, dict) else x for x in val])
                else:
                   setattr(self, key, DictObj(val) if isinstance(val, dict) else val)

    d = {'a':11, 'b':22}
    d

    dd = DictObj(d)
    dd.a --> 
    dd.b --> 
    dd.__dict__ -->

\ decorator 

    def log_result(func):
        def inner(*args, **kwargs):
            res = func(*args, **kwargs)
            print("The result is ", res)
            return res
        return inner

    @log_result
    def square(a):
        return a*a

    # 這個可就厲害了！上面的的 decorator 結構再外加一層，如此就可以帶上 arg 參數！
    def log_with_name(name):
        def log_result(f):
            def inner(*args, **kwargs):
                res = f(*args, **kwargs)
                print(name + ": " + str(res))
                return res
            return inner
        return log_result    
    
    # 這個可就厲害了！上面的的 decorator 結構再外加一層，如此就可以帶上 arg 參數！
    def log_with_name(name):
        def log_result(func):
            def inner(*args, **kwargs):
                res = func(*args, **kwargs)
                print(f"Function is {name} result is ", res)
                return res
            return inner
        return log_result    
    
    @log_with_name('三次方')
    def triple(a):
        return a*a*a

    # 作者原來的例子有誤，訂正好了如下：
    def log_with_name(name):
        def log_result(f):
            def inner(*args, **kwargs):
                res = f(*args, **kwargs)
                print(name + ": " + str(res))
                return res
            return inner
        return log_result

    @log_with_name("Sum")
    def sum(a, b):
        return a+b
    @log_with_name("Difference")
    def subtract(a, b):
        return a-b

    \ decorator 記錄 filter 濾掉了甚麼

    import pandas as pd
    import io, functools

    def tracking_decorator(func):
        @functools.wraps(func) # 可有可無，只是避免將來 <func>.__name__ 都變成不是 <func> 而是 'wrapper'
        def wrapper(*args, **kwargs):
            df = kwargs['data']
            original_trans_set = set(df['trans_id'])
            df = func(*args, **kwargs)
            post_filtering_trans_set = set(df['trans_id'])
            filtered_trans = list(original_trans_set - post_filtering_trans_set)
            return df, filtered_trans
        return wrapper
    
    @tracking_decorator
    def filtering_func(data):
        # apply some filtering #
        filtered_dataframe = data[data['col 2'] > 100]
        return filtered_dataframe
    
    csv_in_memory = """
             trans_id  ;  col 1  ;    col 2       
        a1; 4.4 ;99
        b2; 4.5 ; 200
        c3;   4.7 ; 65    
          d4 ;3.2 ; 140
        """
    csv_in_memory = '\n'.join([s.strip() for s in csv_in_memory.split('\n') if s.strip()])  # 去蕪存菁
    df = pd.read_csv(io.StringIO(csv_in_memory)  , sep=";")
    df.columns = [str(c).strip() for c in df.columns]  # (來自 D-tale) update columns to strings in case they are numbers also trim leading and tailing white spaces.
    %f df :> columns -->
    # 表中的 index column 是個普通 column 不是 index  (row 的 label)
    df    
    
    filtering_func(data=df)
    filtering_func.__name__

\ 用 Path() 讀文字檔

    from pathlib import Path
    Path(r"d:\hcchen\Downloads\DevTools.py").read_text(encoding="utf-8")

\ python tricks

    \ 擷取前後與中間
    a, b, c, *mid, d  = [1, 2, 3, 4, 5, 6,7,8,9]
    print(a, b, c, mid, d)

    \ filter : Print all odd numbers
    numbers = [1, 12, 37, 43, 51, 62, 83, 43, 90, 2020]
    print(list(filter(lambda x: x % 2 == 1, numbers)))  # [1, 37, 43, 51, 83, 43]

\ request http <-- peforth remote debugger 

    \ GOM 置換 peforth accept command 的例子。看來咱既要置換 accept 也要置換 print() 從 projectk 裡面下手？ 

        # Dialog is ready, use it for the user input function
        def peforth_accept ():
            try:
                s = str(gom.script.sys.show_user_defined_dialog(dialog=DIALOG).input)
            except:
                s = "exit" # click cancel or press ESC
            return s

        def bp(id=None,locals=None):
            if id==None: 
                id = 0
                prompt='bp> '
            else:
                prompt="{}>".format(id)
            if id in peforth.bps: peforth.push(locals).ok(prompt, cmd="to _locals_")
        peforth.bp = bp
        peforth.push(bp).dictate("py: vm.bp=pop()")
        peforth.bps = [i for i in range(1000)]

        # initialization
        peforth.push(peforth_accept).dictate(
            """
            \ 設定適合 Gom 環境的 User Interface
                constant peforth_accept // ( -- str ) python function get command line from a text dialog
                : accept peforth_accept :> () dup . cr ; // ( -- str ) A text dialog that reads a command line
                import time \ time :> sleep(1) 暫停一秒鐘

            \ 把 gom 變成 peforth 裡的 python global
                gom	 py: vm.gom=pop()
            """

    \ request 的 get 等於 readline, post 等於 printline 成功! 用 session 速度就快了。
    
        import requests
        
        r = requests.get("http://localhost:8000") # 直接用 requests 固定有 2 秒鐘的 overhead. 改用 session 就快了。 
        
        session = requests.Session()
        r = session.get("http://localhost:8000")
        
        r = requests.get("http://localhost:8000/readtext")
        r = requests.post(
            "http://localhost:8000/printline", 
            headers = {"accept": "application/json", "Content-Type": "application/json"},
            data='{"text": "aabbccccdddd"}'
            )
    
        %f r dir -->
        %f r :> close                  -->
        %f r :> connection             -->
        %f r :> content                -->
        %f r :> cookies                -->
        %f r :> elapsed                -->
        %f r :> encoding               -->
        %f r :> headers                -->
        %f r :> history                -->
        %f r :> is_permanent_redirect  -->
        %f r :> is_redirect            -->
        %f r :> iter_content           -->
        %f r :> iter_lines             -->
        %f r :> json()                 -->
        %f r :> links                  -->
        %f r :> next                   -->
        %f r :> ok                     -->
        %f r :> raise_for_status       -->
        %f r :> raw                    -->
        %f r :> reason                 -->
        %f r :> request                -->
        %f r :> status_code            -->
        %f r :> text                   -->
        %f r :> url                    -->
    
        %%time
        for i in range(3): # 每個 request 有 2 sec 的 overhead 
            r = requests.post(
                "http://localhost:8000/printline", 
                headers = {"accept": "application/json", "Content-Type": "application/json"},
                data='{"text": "111xxx"}'
                )
            # Wall time 3min23s 實在有夠慢。 改用 curl 就很快，FastAPI 的 /docs 網頁也很快，
            # 所有只有用 python 的 requests 超慢，每個 request 都得花 2 sec overhead time <-- excellent inference 成仔好棒！
            
        \ 答案是用 "Session" 見本文: https://stackoverflow.com/questions/62599036/python-requests-is-slow-and-takes-very-long-to-complete-http-or-https-request
            
            requests.head("http://localhost:8000") # --> 405 沒有 support 'head' command 故改用
            requests.get("http://localhost:8000", stream=True)
            
            %%time
            session = requests.Session()
            for i in range(1000):
                s = "%d" % i
                r = session.post(
                    "http://localhost:8000/printline", 
                    data='{"text": "%s"}' % s
                    )
            # 改用 session 就快了。寫錯帶 bug 會沒反應。 100 圈 2.27 s; 1000 圈 4.04 sec  
        
        \ 這篇文章也在問， FastAPI 怎麼這麼慢？ 咱的問題只是沒有用 session 而已，他的不知啥事。
            https://github.com/tiangolo/fastapi/issues/2690 
            
        \ 改試 flask 看看。。。。。 不用了，用上 session 好了。

        \ 用 curl 就很快，以為差在 header ... 結果不是 header 的問題。咱的 FastAPI application 不需要 header. 
          問題出在 requests 有很長的 overhead time, 它有 session 可解。
        
            curl -X 'POST' \
              'http://localhost:8000/printline' \
              -H 'accept: application/json' \
              -H 'Content-Type: application/json' \
              -d '{
              "text": "string swer 6677 2234234"
            }'        

    \ 可以來寫 repl 了   # https://dev.to/amal/building-the-python-repl-3468

        \ 不要用奇怪的 '" 套疊，弄到好很複雜難懂。
        
            r = console.printline("bbb")
            r = console.printline("'bbb'") <---- 這個可以。 
            r = console.printline('"bbb"') <---- 這個印不出來，沒反應！其實 r 有傳回 error 422 and messages. 
                                                 改用 /docs 試試看。。。。。果然看到 error 422 原因是裡面變成 "text": '\''"bbb"'\''
                                                 不要用奇怪的 '" 套疊，弄到好很複雜難懂。
            r = console.printline('\"bbb\"')                                         
            r :> close                  --> <bound method Response.close of <Response [422]>> (<class 'method'>)
            r :> connection             --> <requests.adapters.HTTPAdapter object at 0x000001EA795D5ED0> (<class 'requests.adapters.HTTPAdapter'>)
            r :> content                --> b'{"detail":[{"loc":["body",11],"msg":"Expecting \',\' delimiter: line 1 column 12 (char 11)","type":"value_error.jsondecode","ctx":{"msg":"Expecting \',\' delimiter","doc":"{\\"text\\": \\"\\"bbb\\"\\"}","pos":11,"lineno":1,"colno":12}}]}' (<class 'bytes'>)
            r :> cookies                --> <RequestsCookieJar[]> (<class 'requests.cookies.RequestsCookieJar'>)
            r :> elapsed                --> 0:00:02.018007 (<class 'datetime.timedelta'>)
            r :> encoding               --> utf-8 (<class 'str'>)
            r :> headers                --> {'date': 'Tue, 01 Nov 2022 01:52:19 GMT', 'server': 'uvicorn', 'content-length': '227', 'content-type': 'application/json'} (<class 'requests.structures.CaseInsensitiveDict'>)
            r :> history                --> [] (<class 'list'>)
            r :> is_permanent_redirect  --> False (<class 'bool'>)
            r :> is_redirect            --> False (<class 'bool'>)
            r :> iter_content           --> <bound method Response.iter_content of <Response [422]>> (<class 'method'>)
            r :> iter_lines             --> <bound method Response.iter_lines of <Response [422]>> (<class 'method'>)
            r :> json                   --> <bound method Response.json of <Response [422]>> (<class 'method'>)
            r :> links                  --> {} (<class 'dict'>)
            r :> next                   --> None (<class 'NoneType'>)
            r :> ok                     --> False (<class 'bool'>)
            r :> raise_for_status       --> <bound method Response.raise_for_status of <Response [422]>> (<class 'method'>)
            r :> raw                    --> <urllib3.response.HTTPResponse object at 0x000001EA797D8D60> (<class 'urllib3.response.HTTPResponse'>)
            r :> reason                 --> Unprocessable Entity (<class 'str'>)
            r :> request                --> <PreparedRequest [POST]> (<class 'requests.models.PreparedRequest'>)
            r :> status_code            --> 422 (<class 'int'>)
            r :> text                   --> {"detail":[{"loc":["body",11],"msg":"Expecting ',' delimiter: line 1 column 12 (char 11)","type":"value_error.jsondecode","ctx":{"msg":"Expecting ',' delimiter","doc":"{\"text\": \"\"bbb\"\"}","pos":11,"lineno":1,"colno":12}}]} (<class 'str'>)
            r :> url                    --> http://localhost:8000/printline (<class 'str'>)    
    
        \ 先做個 echo 來試驗，光這樣就發現不少 bug 了！
            1. get 不怕中文, post 送出去的 string 要正確指定 data=string.encode('utf-8')
            2. 用 post 送 emitting string 上 web server 不要用單引號括 string 會很複雜難搞, 用雙引號沒問題。 
            3. [X] repl 一開始怪怪的。。。原因是 server 端收到 readline() command 執行 sys.stdin.readline()  
                   這時候 client 端程式停掉了，這個 stdin 的 readline() 乃成無主的等待，下回 client 端再跑起來，
                   這個無主的 readline() 就會形成一個無效的 keyin line 卡在前面。只能在 server 端敲一下 enter 
                   把它結束掉，別無他法？ 
                   
            # 這個 echo repl 已經很成功了，輸入任何東西都 echo；輸入 exit 結束 client 端的 repl echo loop 但是 server 端仍繼續執行。
            import requests
            import glob, os
            for path_ in ["./", "../", "../../", "../../../", "../../../../", "../../../../../", "../../../../../../"]:
                pathname_ = glob.glob(
                    os.path.join(path_, "forth.py"),
                    recursive=True
                    )
                if pathname_ : 
                    # get_ipython().magic("run %s" % pathname_[0])
                    get_ipython().run_line_magic("run", pathname_[0]) # python 3.8.1 之後建議新寫法
                    break
            from icecream import ic
            ic.configureOutput(outputFunction=print)
            ic('----- 實驗開始 -----');
            session = requests.Session()

            class Debug_console:

                def __init__(self, url):
                    self.url = url
                    self.session = requests.Session()
                
                def readline(self):
                    r = session.get(f"{self.url}/readline")
                    return r
                    
                def printline(self,s):
                    r = session.post(
                        f"{self.url}/printline", 
                        data=('{"text": "%s"}' % s).encode('utf-8') # get() 到中文沒事，post() 帶中文要處理。 
                        )
                    return r 
            
                def repl(self) -> None:
                    while True:
                        self.printline("------ Echo ------")
                        r = self.readline()
                        ic(r.text, r.ok, r.status_code)
                        if r.ok:
                            s = r.json()
                            self.printline(s) # r.json() 直接就是 input string 如果用 r.text 就有括號的麻煩。
                            if s.lower() == 'exit':
                                break # stop the repl loop 
                        else:
                            ic("console.readline() failed: %s" % str(r.json()))
                            break
                    ic("repl loop terminated")
                            
            console = Debug_console('http://localhost:8000')
            r = console.repl()

        \ 上面 echo 成功了，嘗試 python repl 
        
            from builtins import print
            import requests, sys, glob, os, io, logging
            for path_ in ["./", "../", "../../", "../../../", "../../../../", "../../../../../", "../../../../../../"]:
                pathname_ = glob.glob(
                    os.path.join(path_, "forth.py"),
                    recursive=True
                    )
                if pathname_ : 
                    # get_ipython().magic("run %s" % pathname_[0])
                    get_ipython().run_line_magic("run", pathname_[0]) # python 3.8.1 之後建議新寫法
                    break
            from icecream import ic
            
            # These two lines enable debugging at httplib level (requests->urllib3->http.client)
            # You will see the REQUEST, including HEADERS and DATA, and RESPONSE with HEADERS but without DATA.
            # The only thing missing will be the response.body which is not logged.
            try:
                import http.client as http_client
            except ImportError:
                # Python 2
                import httplib as http_client
            http_client.HTTPConnection.debuglevel = 1

            # You must initialize logging, otherwise you'll not see debug output.
            logging.basicConfig()
            logging.getLogger().setLevel(logging.DEBUG)
            requests_log = logging.getLogger("requests.packages.urllib3")
            requests_log.setLevel(logging.DEBUG)
            requests_log.propagate = True

            # requests.get('https://httpbin.org/headers')
            
            ic.configureOutput(outputFunction=print)
            ic('----- 實驗開始 -----');
            
            # redirect stdout to IO buffer so as to grab exec() results 

            class Debug_console:

                def __init__(self, url):
                    self.url = url
                    self.session = requests.Session()
                    self.stdout_save = sys.stdout  # keep a handle on the real standard output
                
                def readline(self):
                    r = self.session.get(f"{self.url}/readline")
                    return r
                    
                def printline(self,s):
                    data = '{"text": "%s"}' % s
                    ic("printline data:",data)
                    r = self.session.post(
                        f"{self.url}/printline", 
                        headers = {
                            "Accept-Language": "en-US,en;q=0.9,zh-TW;q=0.8,zh;q=0.7,zh-CN;q=0.6",
                            "Content-Type": "application/json",
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/106.0.0.0 Safari/537.36"
                            },
                        data=data.encode("utf-8")
                        )
                    return r 
            
                def repl_eval(self) -> None:
                    while True:
                        self.printline("------ python eval repl ------")
                        r = self.readline()
                        ic(r.text, r.ok, r.status_code)
                        if r.ok:
                            s = r.json()
                            # -----
                            result = eval(s)
                            if result:
                                self.printline(result)
                            if s.lower() == 'exit':
                                break # stop the repl loop 
                        else:
                            ic("console.readline() failed: %s" % str(r.json()))
                            break
                    ic("repl loop terminated")
                    # sys.stdout = self.stdout_save # 恢復原來的 stdout 
                            
                def repl_exec(self) -> None:
                    while True:
                        self.printline("------ python exec repl ------")
                        r = self.readline()
                        ic(r.text, r.ok, r.status_code)
                        if r.ok:
                            s = r.json()
                            if s.lower() == 'exit':
                                break # stop the repl loop 
                                
                            sys.stdout = _buffer = io.StringIO() # Choose a file-like object to be new stdout 
                            exec(s) # stdout 都跑進了 _buffer
                            _buffer.seek(0)
                            display_buffer = _buffer.read()
                            _buffer.truncate(0) # 清除 stdout buffer 
                            sys.stdout = self.stdout_save # 恢復 stdout 
                            ic(display_buffer)
                            if display_buffer:
                                r = self.printline(display_buffer)
                                ic(r.json())
                            
                        else:
                            ic("console.readline() failed: %s" % str(r.json()))
                            break
                    ic("repl loop terminated")
                    # sys.stdout = self.stdout_save # 恢復原來的 stdout 
                            
            console = Debug_console('http://localhost:8000')
            console.repl_exec()
            

        \ redirect stdout 試試看，先成功再說 --> 成功！
            目的是把 client 端的 stdout 都導向 FastAPI server 端，這樣 peforth remote console 才看得到畫面。 
            東西都在 client 端，console 只是要看到畫面而已。
        
            import sys,io,peforth 
            stdout_save = sys.stdout       # keep a handle on the real standard output
            sys.stdout = s = io.StringIO() # Choose a file-like object to be new stdout 
            peforth.execute('help');       #  
            sys.stdout = stdout_save
            s.seek(0)
            print(s.read())
            s.truncate(0) # 清除 stdout buffer 
            
            \ redirection 很成功，但是 print 東西會出問題的現象有了更深入的了解：從 server 端輸入 print("abc") 會出問題 (如下) 改成 print("abc",end="") 就 OK 
                root cause 呼之欲出。。。。。。  
                return is from requests module so that error was from server or FastAPI. That means printline() method does not accept CR, right?
                simply try console.printline() see see . . . yeah that's true.  
              
                    r = console.printline("bbb\n") <---- \n is an invalid control character 
                    r.json()
                    {'detail': [{'loc': ['body', 13],
                       'msg': 'Invalid control character at: line 1 column 14 (char 13)',
                       'type': 'value_error.jsondecode',
                       'ctx': {'msg': 'Invalid control character at',
                        'doc': '{"text": "bbb\n"}',
                        'pos': 13,
                        'lineno': 1,
                        'colno': 14}}]}

                    r = console.printline("bbb\t") <---- \t is an invalid control character as well 注意看，整片 error message 都一樣。
                    r.json()
                    {'detail': [{'loc': ['body', 13],
                       'msg': 'Invalid control character at: line 1 column 14 (char 13)',
                       'type': 'value_error.jsondecode',
                       'ctx': {'msg': 'Invalid control character at',
                        'doc': '{"text": "bbb\t"}',
                        'pos': 13,
                        'lineno': 1,
                        'colno': 14}}]}
                            
                    \ 這是正常的樣子 
                    ic| r.text: '"print(\\"abc \\",end=\\"\\")"'
                        r.ok: True
                        r.status_code: 200
                    ic| display_buffer: 'abc '
                    ic| r.json(): {'status': 0}
                
                printline() method is very simple, just accept the text from FastAPI and print it to the server console. When in trouble
                the text never came. Blocked before my function when in FastAPI. So what if using curl or /docs? can they accept control
                characters like \n or \t ? Ha! both curl and /docs work fine with: 

                    curl -X 'POST' \
                      'http://localhost:8000/printline' \
                      -H 'accept: application/json' \
                      -H 'Content-Type: application/json' \
                      -d '{
                      "text": "aabbcc\nddeeff\nggg"
                    }'

                so 
                
        \ --- 研究為何 /docs request printline 甚麼都 ok 而 cline 用 requests 自己送就連 'aaa\nbbb' 都不行？ ----------
        
            跑 localhost:8000/docs 用 F12 查看, body 真的就是 
            
                {
                  "text": "aabbcc\nddeeff\nggg 試試看中文"
                }            

            懷疑 header 可能有 trick. 如下確實有些重要的東西可能有關。應該是 
        
            /docs 
                POST /printline HTTP/1.1
                Accept-Encoding: gzip, deflate, br
                Accept-Language: en-US,en;q=0.9,zh-TW;q=0.8,zh;q=0.7,zh-CN;q=0.6
                Connection: keep-alive
                Content-Length: 51
                Content-Type: application/json
                Cookie: _xsrf=2|108d27ef|aaa11c6c61ae37260ff99755a50fba29|1665122394; username-localhost-8889="2|1:0|10:1666246261|23:username-localhost-8889|44:MWNhYzIyNGE3NDg2NGU1YWIxMzU1NDZkYjM4YTkzYjY=|2d5acf677b0ee2dad1e1bf3b42ab6d56649a0cb84b5d132a70ae29254cd16825"; username-localhost-8888="2|1:0|10:1667357998|23:username-localhost-8888|44:MzFlOTIyNGE4MjQzNDg4Y2ExYTYzNzZlZmYyMGQ2NTE=|3b97224916e752fc7cd6d971de744fd14c818af45dba08adfbe4344143dba204"
                Host: localhost:8000
                Origin: http://localhost:8000
                Referer: http://localhost:8000/docs
                Sec-Fetch-Dest: empty
                Sec-Fetch-Mode: cors
                Sec-Fetch-Site: same-origin
                User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/106.0.0.0 Safari/537.36
                accept: application/json
                sec-ch-ua: "Chromium";v="106", "Google Chrome";v="106", "Not;A=Brand";v="99"
                sec-ch-ua-mobile: ?0
                sec-ch-ua-platform: "Windows"

            想要對照 client requests 送出甚麼東西，用 chrome F12 devtools 看到的是 jupyterlab 的活動，看不到 requests 的。把 url 改到
            網路上的 __https://httpbin.org__ service is a way. Stackoverflow https://stackoverflow.com/a/16630836/2179543 suggests logging seems to be the best way. 

            r = console.printline("qqq\\nkkk")  執行這行，以下是 logging 到的結果：
            
                DEBUG:urllib3.connectionpool:Starting new HTTP connection (1): localhost:8000
                ic| "printline data:": 'printline data:'
                    data: '{"text": "qqq\
                           kkk"}'
                DEBUG:urllib3.connectionpool:http://localhost:8000 "POST /printline HTTP/1.1" 200 12
                send: b'POST /printline HTTP/1.1\r\n
                Host: localhost:8000\r\nUser-Agent: python-requests/2.28.1\r\n
                      Accept-Encoding: gzip, deflate, br\r\n
                      Accept: */*\r\n
                      Connection: keep-alive\r\n
                      Content-Length: 20\r\n\r\n'
                send: b'{"text": "qqq\\nkkk"}'
                reply: 'HTTP/1.1 200 OK\r\n'
                header: date: Wed, 02 Nov 2022 04:02:04 GMT
                header: server: uvicorn
                header: content-length: 12
                header: content-type: application/json

            找到的 requests 例子有 headers 的 (本檔上面就有)
                payload = {"input": "Machu Picchu was believed (by Richard L. Burger, professor of anthropology at Yale University) to have been built in the 1450s.However, a 2021 study led by Burger used radiocarbon dating (specifically, AMS) to reveal that Machu Picchu may have been occupied from around 1420-1530 AD. Construction appears to date from two great Inca rulers, Pachacutec Inca Yupanqui (1438–1471) and Túpac Inca Yupanqui (1472–1493).  There is a consensus among archeologists that Pachacutec ordered the construction of the royal estate for his use as a retreat, most likely after a successful military campaign. Although Machu Picchu is considered to be a royal estate, it would not have been passed down in the line of succession. Rather it was used for 80 years before being abandoned, seemingly because of the Spanish conquests in other parts of the Inca Empire. It is possible that most of its inhabitants died from smallpox introduced by travelers before the Spanish conquistadors arrived in the area."}
                headers = {
                    "Accept-Language": "en-US,en;q=0.9,zh-TW;q=0.8,zh;q=0.7,zh-CN;q=0.6"
                }
            
        \ -------------

\ huble2 by customer word cloud 研究。希望分 customer 列出各自的 stopwords 
    https://stackoverflow.com/questions/27488446/how-do-i-get-word-frequency-in-a-corpus-using-scikit-learn-countvectorizer
    
    import numpy as np    
    from sklearn.feature_extraction.text import CountVectorizer
    texts = ["dog cat fish", "dog cat cat", "fish bird", "bird"]    
    cv = CountVectorizer()   
    cv_fit = cv.fit_transform(texts)    
    word_list = cv.get_feature_names_out()
    \ 接下來就有兩種方法取 word frequency 
        # 原文第一種： Added [0] here to get a 1d-array for iteration by the zip function.  
        count_list = np.asarray(cv_fit.sum(axis=0))[0]
        word_freq = dict(zip(word_list, count_list))  # Output: {'bird': 2, 'cat': 3, 'dog': 2, 'fish': 2} 成功了
        # 第二種： 改用 Counter
        from collections import Counter
        ss = []
        for s in texts:
            ss += s.split() # texts 拆成個別 tokens 一長串 list 
        counter = Counter()
        counter.update(ss)
        most_common = counter.most_common(20)
        most_common[-10:]
        
    # 分 customer 個別的 word frequency list 

        df0 = pd.read_feather(r"c:\Users\8304018\Downloads\hubble2\hubble2-2022-11-3.feather")

        %%time
        dfs = {}
        texts = {}
        word_freq = {}
        for customer in df0.customersystem.unique():
            dfs[customer] = df0[(df0.customersystem == customer) & (df0.sw == True)][['customersystem','sw','bug_id','subject','Text']]
            texts[customer] = dfs[customer].Text.values
            cv = CountVectorizer()   
            cv_fit = cv.fit_transform(texts[customer])    
            word_list = cv.get_feature_names_out()
            count_list = np.asarray(cv_fit.sum(axis=0))[0]
                # Added [0] here to get a 1d-array for iteration by the zip function. 
            word_freq[customer] = dict(zip(word_list, count_list))  # Output: {'bird': 2, 'cat': 3, 'dog': 2, 'fish': 2} 成功了
            sorted(word_freq[customer].items(),key=lambda x:x[1])[-100:] # 結果看到 word count 大的離譜
        # Wall time: 24.5 s

    # 整個 df0 一起的 word frequency list 
            
        # max(count_list) --> 竟然有 1245317 可能有錯，驗算看看。。。沒錯，那就是要取 log 下圖就是了。  
        cv_fit = cv.fit_transform(df0.Text.values)    
        word_list = cv.get_feature_names_out()
        count_list = np.asarray(cv_fit.sum(axis=0))[0]
            # Added [0] here to get a 1d-array for iteration by the zip function. 
        word_freq = dict(zip(word_list, count_list))  # Output: {'bird': 2, 'cat': 3, 'dog': 2, 'fish': 2} 成功了
        max(count_list)    
        min(count_list)    
        pd.Series(count_list).describe()
        %f count_list count --> nip
        %f word_freq count --> nip
        sorted(word_freq.items(),key=lambda x:x[1])[-100:] # 結果看到 word count 大得離譜
    
    # Normalization or standardization 研究
        # 原文 https://zhuanlan.zhihu.com/p/104695889
        # 這個例子一跑成功，畫出 histogram 比較不同 normalization transform 的結果圖形。 真的是 hubble2 的三年 word frequency !
        # Note! 下圖已經取 log 嘍！ Hubble2 三年的 word count 大小差很多： 1 ~ 1.2M 的差距，取 log  整個都還能呈現。
        #     evernote:///view/2472143/s22/f098c6ee-84ed-b728-0d20-f1c3ec6b77f6/e5135377-040a-eb4b-a719-8e1a50369a67

        import numpy as np
        from sklearn.preprocessing import StandardScaler,MinMaxScaler, minmax_scale
        import matplotlib.pyplot as plt
        # test_arr = np.array(list(word_freq.values()))  # np.random.randn( 10000) * 100 + 500
        test_arr = np.random.randn( 10000) * 100 + 500            
        std_sca = StandardScaler()
        mmx_sca = MinMaxScaler()
        test_std = std_sca.fit_transform(test_arr.reshape(-1, 1)) # 橫的 list 改成豎的
        test_mmx = mmx_sca.fit_transform(test_arr.reshape(-1, 1)) # 橫的 list 改成豎的
        fig,axes = plt.subplots( 1, 3,figsize=( 20, 5))
        plt.suptitle(' Original Array ===> After Standard Transform ===> After MinMax TransForm',fontsize= 20)
        for idx, array in enumerate([test_arr,test_std,test_mmx]):
            axes[idx].hist(array,log=True)
            
    mmx_sca = minmax_scale(test_arr)
    fig,axes = plt.subplots(figsize=( 20, 5))
    axes.hist(mmx_sca)

    
    from sklearn.preprocessing import minmax_scale 
    x = [0,1,2,3,4,5] 
    minmax_scale(x) # array([0. , 0.2, 0.4, 0.6, 0.8, 1. ]) 
    y = [[0,0,0],[1,1,1],[2,2,2]]
    minmax_scale(y) # array([[0. , 0. , 0. ], [0.5, 0.5, 0.5], [1. , 1. , 1. ]]) 
    minmax_scale(y, axis=1) # array([[0., 0., 0.], [0., 0., 0.], [0., 0., 0.]]) 

    y = [[0,1,2],[1,2,3],[2,3,4]] 
    minmax_scale(y) # array([[0. , 0. , 0. ], [0.5, 0.5, 0.5], [1. , 1. , 1. ]]) 

    minmax_scale(y, axis=1) # array([[0. , 0.5, 1. ], [0. , 0.5, 1. ], [0. , 0.5, 1. ]])

    sorted(word_freq['Astro'].items(),key=lambda x:x[1])[-10:] # 結果看到 word count 大的離譜
    
    
    
    \ 改用 NLTK , because the aboves 出現巨大的 word count 值未解，可能 stackoverflow 抄來的 sample code 有問題。
    import nltk
    nltk.download('punkt')    
    
    from collections import Counter
    from nltk.tokenize import RegexpTokenizer
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize

    text='''aa bbb cc dd cc cc bbb cc
    '''
    text = np.array(list(word_freq['Astro'].values()))  # np.random.randn( 10000) * 100 + 500

    # tokenize
    raw = ' '.join(word_tokenize(text.lower()))

    tokenizer = RegexpTokenizer(r'[A-Za-z]{2,}')
    words = tokenizer.tokenize(raw)

    # remove stopwords
    stop_words = set(stopwords.words('english'))
    words = [word for word in words if word not in stop_words]

\ Dump 三年來的 TR to the dump/ directory 每千條一個 .txt file 
            
    \ 所有 TR 全 dump 成一大 'dump' folder 當天更新
        \ 這個 dump 供 Notepad++ 用 RegEx 做 linear search 以便於與 Nearest Neighbor 方法比較。
        \ 如果想要逐月更新以圖只更新最後一個月也不行，因為前面的 TR 會逐漸往後跑。想想很頭痛，算了吧！
        \ 2019 ~ 2021 一年一分鐘共約三分鐘。

        %%time
        if not PROD:
            # 每次重做，先清空 dump directory 
            dump_directory = setup.path + 'dump\\'  
            for f in os.listdir(dump_directory):
                os.remove(os.path.join(dump_directory, f))

            i = 1 
            s = ""
            for bug in df.iloc:
                if s=="": # 每千條 bug 之第一條的 last_updated 時間組成檔名，df 最上面已經是最新的。
                    t = pd.to_datetime(bug['last_updated'])
                    fname = "%02d-%02d-%02d-%02d%02d%02d.txt" % (t.year, t.month, t.day, t.hour, t.minute, t.second)
                s += "%4s %12s   %s BoW: %s   tcount: %d\n" % (
                        i, 
                        bug.bug_id, 
                        bug.subject.replace("\r\n","").replace("\n",""), 
                        bug.Text, 
                        bug.tcount)
                if i % 1000 == 0:
                    %f i -->
                    with open(dump_directory + fname, 'w', encoding='utf-8') as f:
                        f.write(s)
                    s = ""
                i += 1    
        # Wall time: 9m 2.6s

\ 依據 500 * 6 條 NN 評判 algorithm 的好壞，評判方式：計數 keyword AndsOrs 吻合的條樹，目前就這個方法。

    \ dump 成 <bug_id>.txt 列出 500*6 條 NN，使用 hubble2NN 的 jupyterlab kernel 就有 buglist 可用。
        results = buglist.get_customers_neighbors("1936505") # -> {'ROSA':df[:500], 'UTS':df[:500], ... } 一秒完成！
        %%time
        s = ""
        for oem,df in results.items():
            i = 0
            for bug in df.iloc:
                s += "%4s %8s %12s   %s BoW: %s tcount: %d  dist: %f\n" % (
                        i, 
                        bug.customersystem,
                        bug.bug_id, 
                        bug.subject.replace("\r\n","").replace("\n",""), 
                        bug.Text, 
                        bug.tcount,
                        bug.dist,
                        )
                if i % 100 == 0:
                    %f i -->
                i += 1    
        with open(r"c:\Users\8304018\Downloads\1936505.txt", 'w', encoding='utf-8') as f:
            f.write(s)
        # Wall time: 7.6s

    \ 2022/11/4 16:38 DevTools.py regular expression Ands Ors search of a file (dump of 500*6 NN TRs)

        \ regular expression Ands Ors search from a file
          要做精細： subject 用 linear search, Bow 用 in set() --> 好了，試試看。。。  

        from pathlib import Path
        def AndsOrs(lines, ands):
            count = 0 
            founds = []
            for line in lines:
                subject = line[:line.find("BoW:")] # 一行搞定取得 subject string 
                bow = set(line.split(' ')) # 這條 TR 的 bow 整理成集合
                count_down = len(ands) # 有多少 ands 要滿足，到 0 表示全部滿足。
                for ors in ands: # 祭出 synonym-expaded keywords
                    ors_flag = False # 先假設 ors 不成立
                    for kw in ors:
                        kw = kw.lower()
                        if (kw in bow) or (subject.find(kw) != -1): # BoW 不成立才老實做 linear search 
                            count_down -= 1
                            ors_flag = True
                            break # bow 裡只要有一個 keyword 這整條 ors 就成立了，跳下ㄧ個 ands 
                    if not ors_flag :
                        break # 只要有一條 ors 不成立，後面都不用看了。跳下一條 TR。
                if not count_down: # 到 0 表示成立了， line 為所求
                    founds.append(line)
                    count += 1    
            return count, founds

        '''
        c,t = AndsOrs(
            Path(r"c:\Users\8304018\Downloads\328505b.txt").read_text(encoding="utf-8").split('\n')[3:], 
            [
            # ["monitor"] # c --> 281 (新版 262) 這是察看 word frequency 極高的 monitor 證實根本沒關係的 TR 裡都有這個字眼
            # ["display"] # c --> 305 (新版 253) display 亦然。所以評鑑 NN algorithm 好壞不是合用，鑑別力不好。
            # ["monitor","display"] # c --> 496  (新版 441)
            # ["bsod","bluescreen", "blue screen"],["usb"] # c --> 21 --> 17 --> 20 新版變多，好不好？
            # ["bsod","bluescreen", "blue screen"] # c --> 110 --> 108 新版又變少。。
            # ["driver"],["wireless","wifi","wlan"],["uninstall"] # c --> 101 --> 89 新版變少，這題本來就不太好。
            ])
        c

        \ 把 AndsOrs() -> founds list 轉成 .txt file 以便手工檢查 AndsOrs() 是否正確。
            s = ""
            for l in t:
                s += "%s\n" % l
            %f s char c:\Users\8304018\Downloads\2.txt writeTextFile
            
        \ 好了，算出 AndsOrs() synonym-expanded keyword count 值當作 NN algorithm 的成績已完成。
        '''

        '''
        c,t = AndsOrs(
            Path(r"c:\Users\8304018\Downloads\1936505.txt").read_text(encoding="utf-8").split('\n')[3:], 
            [
              ['0x7e'], ['BSOD','bluescreen','blue screen','blue-screen'],
            ])
        c

        \ 把 AndsOrs() -> founds list 轉成 .txt file 以便手工檢查 AndsOrs() 是否正確。
            s = ""
            for l in t:
                s += "%s\n" % l
            %f s char c:\Users\8304018\Downloads\3.txt writeTextFile
            
        \ 好了，算出 AndsOrs() synonym-expanded keyword count 值當作 NN algorithm 的成績已完成。
        '''
    from pregex.core.classes import AnyButWhitespace
    from pregex.core.quantifiers import OneOrMore
    from pregex.core.operators import Either

\ [第 24 天] 機器學習（4）分群演算法 clustering - kmeans, heirachical, DBSCAN
    https://ithelp.ithome.com.tw/articles/10187314

    \ Kmeans 
    
        from sklearn import cluster, datasets

        # 讀入鳶尾花資料
        iris = datasets.load_iris()
        iris_X = iris.data

        # KMeans 演算法
        kmeans_fit = cluster.KMeans(n_clusters = 3).fit(iris_X)

        # 印出分群結果
        cluster_labels = kmeans_fit.labels_
        print("分群結果：")
        print(cluster_labels)
        print("---")

        # 印出品種看看
        iris_y = iris.target
        print("真實品種：")
        print(iris_y)

        \ check performance 
        
        from sklearn import cluster, datasets, metrics

        # 讀入鳶尾花資料
        iris = datasets.load_iris()

        # KMeans 演算法
        kmeans_fit = cluster.KMeans(n_clusters = 3).fit(iris_X)
        cluster_labels = kmeans_fit.labels_

        # 印出績效
        silhouette_avg = metrics.silhouette_score(iris_X, cluster_labels)
        print(silhouette_avg)

    \ Hierarchical Clustering

        from sklearn import cluster, datasets

        # 讀入鳶尾花資料
        iris = datasets.load_iris()
        iris_X = iris.data

        # Hierarchical Clustering 演算法
        hclust = cluster.AgglomerativeClustering(linkage = 'ward', affinity = 'euclidean', n_clusters = 3)

        # 印出分群結果
        hclust.fit(iris_X)
        cluster_labels = hclust.labels_
        print(cluster_labels)
        print("---")

        # 印出品種看看
        iris_y = iris.target
        print(iris_y)

        \ check performance 
        
        from sklearn import cluster, datasets, metrics

        # 讀入鳶尾花資料
        iris = datasets.load_iris()
        iris_X = iris.data

        # Hierarchical Clustering 演算法
        hclust = cluster.AgglomerativeClustering(linkage = 'ward', affinity = 'euclidean', n_clusters = 3)

        # 印出績效
        hclust.fit(iris_X)
        cluster_labels = hclust.labels_
        silhouette_avg = metrics.silhouette_score(iris_X, cluster_labels)
        print(silhouette_avg)

    \ 不要再用K-means！ 超實用分群法 DBSCAN 詳解 sklearn DBSCAN使用介紹 --> 直接跳 HDBSCAN 
        https://axk51013.medium.com/%E4%B8%8D%E8%A6%81%E5%86%8D%E7%94%A8k-means-%E8%B6%85%E5%AF%A6%E7%94%A8%E5%88%86%E7%BE%A4%E6%B3%95dbscan%E8%A9%B3%E8%A7%A3-a33fa287c0e
    \ 2021年資料科學家必備分群法（Clustering）：HDBSCAN簡介
        https://axk51013.medium.com/不要再用k-means-超實用分群法dbscan詳解-a33fa287c0e
        
        conda install -c conda-forge hdbscan 
        
        from sklearn.datasets import make_blobs
        import hdbscan
        import matplotlib.pyplot as plt

        N_SAMPLES = 1000
        RANDOM_STATE = 42
        X, y = make_blobs(n_samples=N_SAMPLES,
                          cluster_std=[2.0, 0.5],
                          centers=[(0, 0), (5, 5)],
                          random_state=RANDOM_STATE)

        plt.figure(figsize = (10, 10))
        plt.scatter(X[:, 0], X[:, 1])
        plt.show()

        hclusterer = hdbscan.HDBSCAN(min_cluster_size=5).fit(X)

        plt.figure(figsize = (10, 10))
        plt.scatter(X[:, 0], X[:, 1], c = hclusterer.labels_)
        plt.show()

\ OpenAI 寫的程式！
            
import winsound
import time

while True:
    winsound.Beep(440, 1000)
    time.sleep(10)            